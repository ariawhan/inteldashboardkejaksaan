import os
import re
import stat
import uuid
import base64
import hashlib
import json
import shutil
import threading
import time
import zipfile
from datetime import date, datetime
from functools import wraps
from io import BytesIO
from pathlib import Path
from urllib.parse import quote, urlparse
from xml.sax.saxutils import escape as xml_escape

import pymysql
import requests
import click
from bs4 import BeautifulSoup, NavigableString, Tag
from cryptography.fernet import Fernet, InvalidToken
from dotenv import load_dotenv
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from flask import Flask, abort, flash, g, jsonify, redirect, render_template, request, send_file, send_from_directory, session, url_for
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas as reportlab_canvas
from reportlab.platypus import (
    Image,
    KeepTogether,
    LongTable,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from PIL import Image as PILImage, ImageOps
from playwright.sync_api import sync_playwright
from xml.sax.saxutils import escape
from werkzeug.exceptions import RequestEntityTooLarge
from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash, generate_password_hash


load_dotenv()

app = Flask(__name__)
app.config.update(
    SECRET_KEY=os.getenv("SECRET_KEY", "dev-secret-ganti-saat-production"),
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    PERMANENT_SESSION_LIFETIME=60 * 60 * 12,
    MAX_CONTENT_LENGTH=80 * 1024 * 1024,
)

DB_CONFIG = {
    "host": os.getenv("DB_HOST", "127.0.0.1"),
    "port": int(os.getenv("DB_PORT", "3306")),
    "user": os.getenv("DB_USER", "root"),
    "password": os.getenv("DB_PASSWORD", ""),
    "charset": "utf8mb4",
    "cursorclass": pymysql.cursors.DictCursor,
}
DB_NAME = os.getenv("DB_NAME", "intel_dashboard")
MIGRATIONS_DIR = Path(__file__).resolve().parent / "migrations"
STATIC_IMG_DIR = Path(__file__).resolve().parent / "static" / "img"
UPLOAD_DIR = Path(__file__).resolve().parent / "uploads" / "reports"
SIGNATORY_UPLOAD_DIR = Path(__file__).resolve().parent / "uploads" / "signatories"
ORGANIZATION_UPLOAD_DIR = Path(__file__).resolve().parent / "uploads" / "organization"
USER_AVATAR_UPLOAD_DIR = Path(__file__).resolve().parent / "uploads" / "users"
MAX_ATTACHMENT_SIZE_BYTES = 100 * 1024
MAX_ATTACHMENT_DIMENSION = 2200
ALLOWED_IMAGE_TYPES = {"image/jpeg", "image/png", "image/webp"}
ISSUE_CODES = {"Ds.1", "Ds.2", "Ds.3", "Dip.1", "Dip.2", "Dip.3", "Dip.4",
               "Dsb.1", "Dsb.2", "Dsb.3", "Dsb.4", "Dek.1", "Dek.2", "Dek.3", "Dek.4",
               "Dpp.1", "Dpp.2", "Dpp.3", "Dpp.4", "Dti.1", "Dti.2", "Dti.3", "Dti.4"}
ISSUE_CODE_LABELS = {
    "Ds.1": "Penyusunan Program, Laporan dan Penilaian",
    "Ds.2": "Tata Usaha",
    "Ds.3": "Keuangan",
    "Dip.1": "Ideologi",
    "Dip.2": "Politik",
    "Dip.3": "Pertahanan dan Keamanan",
    "Dip.4": "Cegah Tangkal, Pengawasan Orang Asing, Pengamanan Sumber Daya Organisasi Kejaksaan dan Pengamanan Penanganan Perkara",
    "Dsb.1": "Peredaran Barang Cetakan dan Media Komunikasi",
    "Dsb.2": "Aliran Kepercayaan Masyarakat dan Aliran Keagamaan serta Pencegahan Penyalahgunaan dan Penodaan Agama",
    "Dsb.3": "Budaya dan Kemasyarakatan",
    "Dsb.4": "Sosial, Ketertiban dan Ketentraman Umum, Pembinaan Masyarakat Taat Hukum (Binmatkum)",
    "Dek.1": "Keuangan dan Kekayaan Negara",
    "Dek.2": "Investasi dan Penerimaan Negara",
    "Dek.3": "Perdagangan, Perindustrian dan Ketenagakerjaan",
    "Dek.4": "Sumber Daya Alam dan Agraria atau Tata Ruang",
    "Dpp.1": "Pengamanan Pembangunan Infrastruktur Transportasi dan Telekomunikasi",
    "Dpp.2": "Pengamanan Pembangunan Infrastruktur Pengairan, Pertanian dan Kelautan",
    "Dpp.3": "Pengamanan Pembangunan Infrastruktur Energi, Sumber Daya Alam dan Ilmu Pengetahuan dan Teknologi",
    "Dpp.4": "Pengamanan Pembangunan Infrastruktur Kawasan dan Sektor Strategis Lainnya",
    "Dti.1": "Produksi Intelijen",
    "Dti.2": "Pemantauan",
    "Dti.3": "Pengamanan Informasi",
    "Dti.4": "Pengembangan Sumber Daya Teknologi Informasi",
}
REGISTER_INFORMATION_VALUES = [f"{letter}{number}" for letter in "ABCDEF" for number in range(1, 7)]
CHROME_EXECUTABLE = Path(os.getenv("CHROME_EXECUTABLE", r"C:\Program Files\Google\Chrome\Application\chrome.exe"))
INTELIZ_LOGIN_URL = "https://inteliz.kejaksaan.go.id/login"
INTELIZ_2FA_PATH = "/2fa/challenge"
INTELIZ_LAPINHAR_CREATE_URL = "https://inteliz.kejaksaan.go.id/lapinhar/create"
INTELIZ_LAPINSUS_CREATE_URL = "https://inteliz.kejaksaan.go.id/lapinsus/create"
INTELIZ_AUTH_SESSIONS = {}
INTELIZ_AUTH_LOCK = threading.Lock()
SIPEDE_LOGIN_URL = "https://sipede.kejaksaan.go.id/login"
SIPEDE_BASE_URL = "https://sipede.kejaksaan.go.id"
SIPEDE_SURATKELUAR_CREATE_URL = "https://sipede.kejaksaan.go.id/suratkeluar/create?idSurat=125"
SIPEDE_AUTH_SESSIONS = {}
SIPEDE_AUTH_LOCK = threading.Lock()


class IntelizAuthenticationRequired(RuntimeError):
    pass


def browser_launch_args(extra_args=None, browser_key="BROWSER", use_proxy=True):
    args = [
        "--disable-blink-features=AutomationControlled",
        "--disable-features=BlockInsecurePrivateNetworkRequests",
    ]
    if use_proxy:
        proxy_server = (
            os.getenv(f"{browser_key}_PROXY_SERVER", "").strip()
            or os.getenv("BROWSER_PROXY_SERVER", "").strip()
        )
        proxy_bypass = (
            os.getenv(f"{browser_key}_PROXY_BYPASS", "").strip()
            or os.getenv("BROWSER_PROXY_BYPASS", "").strip()
        )
        if proxy_server:
            args.append(f"--proxy-server={proxy_server}")
        elif os.getenv(f"{browser_key}_USE_SYSTEM_PROXY", os.getenv("BROWSER_USE_SYSTEM_PROXY", "1")) == "1":
            args.append("--proxy-server=system")
        if proxy_bypass:
            args.append(f"--proxy-bypass-list={proxy_bypass}")
    if extra_args:
        args.extend(extra_args)
    return args


@app.errorhandler(RequestEntityTooLarge)
def handle_request_too_large(_error):
    message = (
        "Ukuran upload terlalu besar. Silakan kurangi jumlah foto sekali upload "
        "atau kecilkan ukuran file sebelum dikirim."
    )
    if request.path.startswith("/lapinhar") or request.path.startswith("/lapinsus"):
        flash(message, "error")
        return redirect(request.url)
    return jsonify(message=message), 413

PDF_FONT_NAME = "Times-Roman"
PDF_FONT_BOLD = "Times-Bold"
PDF_FONT_ITALIC = "Times-Italic"
PDF_FONT_BOLD_ITALIC = "Times-BoldItalic"
WINDOWS_FONT_DIR = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
try:
    pdfmetrics.registerFont(TTFont("TimesNewRoman", str(WINDOWS_FONT_DIR / "times.ttf")))
    pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", str(WINDOWS_FONT_DIR / "timesbd.ttf")))
    pdfmetrics.registerFont(TTFont("TimesNewRoman-Italic", str(WINDOWS_FONT_DIR / "timesi.ttf")))
    pdfmetrics.registerFont(TTFont("TimesNewRoman-BoldItalic", str(WINDOWS_FONT_DIR / "timesbi.ttf")))
    pdfmetrics.registerFontFamily("TimesNewRoman", normal="TimesNewRoman", bold="TimesNewRoman-Bold",
                                  italic="TimesNewRoman-Italic", boldItalic="TimesNewRoman-BoldItalic")
    PDF_FONT_NAME, PDF_FONT_BOLD = "TimesNewRoman", "TimesNewRoman-Bold"
    PDF_FONT_ITALIC, PDF_FONT_BOLD_ITALIC = "TimesNewRoman-Italic", "TimesNewRoman-BoldItalic"
except (OSError, ValueError):
    pass


def initialize_database():
    connection = pymysql.connect(**DB_CONFIG)
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                f"CREATE DATABASE IF NOT EXISTS `{DB_NAME}` "
                "CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci"
            )
            cursor.execute(f"USE `{DB_NAME}`")
            cursor.execute(
                """CREATE TABLE IF NOT EXISTS schema_migrations (
                    version VARCHAR(100) PRIMARY KEY,
                    applied_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
                ) ENGINE=InnoDB"""
            )
            cursor.execute("SELECT version FROM schema_migrations")
            applied = {row["version"] for row in cursor.fetchall()}
            for migration_file in sorted(MIGRATIONS_DIR.glob("*.sql")):
                if migration_file.name in applied:
                    continue
                statements = [
                    statement.strip()
                    for statement in migration_file.read_text(encoding="utf-8").split(";")
                    if statement.strip()
                ]
                for statement in statements:
                    cursor.execute(statement)
                cursor.execute(
                    "INSERT INTO schema_migrations (version) VALUES (%s)",
                    (migration_file.name,),
                )
            admin_username = os.getenv("ADMIN_USERNAME", "admin")
            cursor.execute("SELECT id FROM users WHERE username = %s", (admin_username,))
            if cursor.fetchone() is None:
                cursor.execute(
                    "INSERT INTO users (username, full_name, password_hash, role) VALUES (%s, %s, %s, 'admin')",
                    (
                        admin_username,
                        os.getenv("ADMIN_FULL_NAME", "Administrator"),
                        generate_password_hash(os.getenv("ADMIN_PASSWORD", "admin123")),
                    ),
                )
        connection.commit()
    finally:
        connection.close()


@app.cli.command("migrate")
def migrate_command():
    """Jalankan seluruh migration database yang belum diterapkan."""
    initialize_database()
    print(f"Migration database '{DB_NAME}' selesai.")


def existing_max_document_number(document_type, document_year):
    table_name = "lapinhar_reports" if document_type == "lapinhar" else "lapinsus_reports"
    prefix = "R-LIH" if document_type == "lapinhar" else "R-LIK"
    pattern = re.compile(rf"^{prefix}-(\d+)[A-Z]*/.+/{document_year}$", re.IGNORECASE)
    max_number = 0
    connection = pymysql.connect(database=DB_NAME, **DB_CONFIG)
    try:
        with connection.cursor() as cursor:
            like_pattern = f"{prefix}-%/{document_year}".replace("\\/", "/")
            cursor.execute(
                f"SELECT report_number FROM {table_name} "
                "WHERE report_number LIKE %s",
                (like_pattern,),
            )
            for row in cursor.fetchall():
                match = pattern.match(row.get("report_number") or "")
                if match:
                    max_number = max(max_number, int(match.group(1)))
    finally:
        connection.close()
    return max_number


@app.cli.command("set-last-number")
@click.argument("document_type", type=click.Choice(["lapinhar", "lapinsus"], case_sensitive=False))
@click.argument("last_number", type=int)
@click.option("--year", "document_year", default=lambda: date.today().year, show_default="tahun berjalan", type=int)
def set_last_number_command(document_type, last_number, document_year):
    """Set nomor surat terakhir manual. Sistem akan memakai nomor berikutnya."""
    initialize_database()
    document_type = document_type.lower()
    existing_max = existing_max_document_number(document_type, document_year)
    requested_next = last_number + 1
    safe_next = max(requested_next, existing_max + 1)
    connection = pymysql.connect(database=DB_NAME, autocommit=True, **DB_CONFIG)
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """INSERT INTO document_counters (document_type, document_year, next_number)
                   VALUES (%s, %s, %s)
                   ON DUPLICATE KEY UPDATE next_number = VALUES(next_number)""",
                (document_type, document_year, safe_next),
            )
    finally:
        connection.close()
    prefix = "R-LIH" if document_type == "lapinhar" else "R-LIK"
    if safe_next != requested_next:
        print(
            f"Nomor {prefix} tahun {document_year} diset ke {safe_next}. "
            f"Database sudah punya nomor sampai {existing_max}, jadi angka dibuat aman."
        )
    else:
        print(f"Nomor {prefix} tahun {document_year} berhasil diset. Nomor berikutnya: {safe_next}.")


def get_db():
    if "db" not in g:
        g.db = pymysql.connect(database=DB_NAME, autocommit=True, **DB_CONFIG)
    return g.db


@app.teardown_appcontext
def close_db(_error=None):
    db = g.pop("db", None)
    if db is not None:
        db.close()


def fetch_one(query, params=()):
    with get_db().cursor() as cursor:
        cursor.execute(query, params)
        return cursor.fetchone()


def fetch_all(query, params=()):
    with get_db().cursor() as cursor:
        cursor.execute(query, params)
        return cursor.fetchall()


def login_required(view):
    @wraps(view)
    def wrapped_view(**kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return view(**kwargs)
    return wrapped_view


def admin_required(view):
    @wraps(view)
    @login_required
    def wrapped_view(**kwargs):
        if session.get("role") != "admin":
            flash("Menu tersebut hanya dapat diakses oleh admin.", "error")
            return redirect(url_for("dashboard"))
        return view(**kwargs)
    return wrapped_view


MAHASISWA_ALLOWED_ENDPOINTS = {
    "static", "logout", "lapinhar", "lapinsus",
    "create_lapinhar", "edit_lapinhar", "delete_lapinhar",
    "create_lapinsus", "edit_lapinsus", "delete_lapinsus",
    "check_lapinhar_number", "reload_lapinhar_number",
    "check_lapinsus_number", "reload_lapinsus_number",
    "reserve_backdated_number_endpoint", "cancel_backdated_number_endpoint",
    "lapinhar_attachment_file", "lapinsus_attachment_file",
    "export_lapinhar_preview_pdf", "user_avatar",
    "signatory_signature_file", "organization_digital_stamp_file",
}


@app.before_request
def restrict_mahasiswa_access():
    if session.get("role") != "mahasiswa":
        return None
    endpoint = request.endpoint
    if endpoint in MAHASISWA_ALLOWED_ENDPOINTS:
        return None
    if endpoint == "dashboard":
        return redirect(url_for("lapinhar"))
    flash("Akun mahasiswa hanya dapat mengakses dan mengelola laporan milik sendiri.", "error")
    return redirect(url_for("lapinhar"))


@app.route("/", methods=["GET", "POST"])
def login():
    if session.get("user_id"):
        return redirect(url_for("dashboard"))
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        user = fetch_one("SELECT * FROM users WHERE username = %s", (username,))
        if user is None or not check_password_hash(user["password_hash"], request.form.get("password", "")):
            flash("Username atau kata sandi tidak sesuai.", "error")
        elif not user.get("is_active", 1):
            flash("Akun Anda sedang dinonaktifkan. Hubungi administrator.", "error")
        else:
            session.clear()
            session.permanent = request.form.get("remember") == "on"
            session.update(
                user_id=user["id"], full_name=user["full_name"],
                username=user["username"], role=user["role"],
                profile_photo=user.get("profile_photo")
            )
            return redirect(url_for("lapinhar") if user["role"] == "mahasiswa" else url_for("dashboard"))
    return render_template("login.html")


@app.route("/dashboard")
@login_required
def dashboard():
    if session.get("role") == "mahasiswa":
        return redirect(url_for("lapinhar"))
    report_filter = "" if session["role"] == "admin" else " WHERE created_by = %s"
    params = () if session["role"] == "admin" else (session["user_id"],)
    rows = fetch_all(
        "SELECT report_type, COUNT(*) total, SUM(status = 'draft') draft, "
        "SUM(status <> 'draft' AND NOT EXISTS("
        "SELECT 1 FROM register_intelijen_entries register_entries "
        "WHERE register_entries.source_report_type=reports.report_type "
        "AND register_entries.source_report_id=reports.id"
        ")) register_belum "
        f"FROM reports{report_filter} GROUP BY report_type", params
    )
    counts = {
        "lapinhar": {"total": 0, "draft": 0, "register_belum": 0},
        "lapinsus": {"total": 0, "draft": 0, "register_belum": 0},
    }
    for row in rows:
        counts[row["report_type"]] = {
            "total": row["total"],
            "draft": int(row["draft"] or 0),
            "register_belum": int(row["register_belum"] or 0),
        }
    user_count = fetch_one("SELECT COUNT(*) total FROM users")["total"] if session["role"] == "admin" else None
    return render_template("dashboard.html", counts=counts, user_count=user_count, active="dashboard")


MONTH_NAMES_ID = [
    "", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
    "Juli", "Agustus", "September", "Oktober", "November", "Desember",
]


def format_indonesian_date(value):
    if not value:
        return "-"
    if isinstance(value, datetime):
        value = value.date()
    return f"{value.day} {MONTH_NAMES_ID[value.month]} {value.year}"


def issue_code_label(code):
    return ISSUE_CODE_LABELS.get(code or "", code or "-")


GENERIC_REGISTER_CONFIGS = {
    "rin2": {
        "code": "R.IN.2",
        "title": "Register Surat Keluar",
        "description": "Register nomor surat keluar, tujuan, perihal, lampiran, dan keterangan.",
        "date_field": "letter_date",
        "chart_field": "recipient",
        "chart_title": "Grafik Kepada",
        "fields": [
            {"name": "letter_number", "label": "Nomor Surat", "type": "text", "required": True},
            {"name": "letter_date", "label": "Tanggal Surat", "type": "date", "required": True},
            {"name": "recipient", "label": "Kepada", "type": "text", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "attachment", "label": "Lampiran", "type": "text"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
    "rin4": {
        "code": "R.IN.4",
        "title": "Register Ekspedisi Surat",
        "description": "Register ekspedisi surat berdasarkan satker, laporan, tujuan, penerima, dan waktu diterima.",
        "date_field": "report_date",
        "chart_field": "recipient",
        "chart_title": "Grafik Kepada",
        "fields": [
            {"name": "satker_name", "label": "Nama Satker", "type": "satker", "required": True},
            {"name": "report_date", "label": "Tanggal Laporan", "type": "date", "required": True},
            {"name": "letter_number_date", "label": "No & Tanggal Surat", "type": "textarea", "required": True},
            {"name": "recipient", "label": "Kepada", "type": "text", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "attachment", "label": "Lampiran", "type": "text"},
            {"name": "received_time", "label": "Waktu Diterima", "type": "datetime"},
            {"name": "receiver_name", "label": "Nama Penerima", "type": "text"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
    "rin6": {
        "code": "R.IN.6",
        "title": "Register Arsip",
        "description": "Register arsip masuk, kode penyimpanan, lampiran, dan keterangan.",
        "date_field": "report_date",
        "chart_field": "storage_code",
        "chart_title": "Grafik Kode Penyimpanan",
        "fields": [
            {"name": "satker_name", "label": "Nama Satker", "type": "satker", "required": True},
            {"name": "report_date", "label": "Tanggal Laporan", "type": "date", "required": True},
            {"name": "received_time", "label": "Waktu Terima", "type": "date", "required": True},
            {"name": "received_from", "label": "Diterima Dari", "type": "text", "required": True},
            {"name": "letter_number_date", "label": "No & Tanggal Surat", "type": "textarea", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "attachment", "label": "Lampiran", "type": "text"},
            {"name": "storage_code", "label": "Kode Penyimpanan", "type": "text", "default": "Ds.1"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
    "rin7": {
        "code": "R.IN.7",
        "title": "Kegiatan Bidang Ideologi, Politik, Pertahanan dan Keamanan, Cegah Tangkal dan Pengawasan Orang Asing, Pengamanan Sumber Daya Organisasi Kejaksaan dan Pengamanan Penanganan Perkara",
        "description": "Register kegiatan bidang ideologi, politik, pertahanan keamanan, cegah tangkal, pengawasan orang asing, dan pengamanan.",
        "date_field": "activity_date",
        "chart_field": "sector",
        "chart_title": "Grafik Sektor",
        "fields": [
            {"name": "sector", "label": "Sektor", "type": "text", "required": True},
            {"name": "number", "label": "Nomor", "type": "text", "required": True},
            {"name": "activity_date", "label": "Tanggal", "type": "date", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "executor", "label": "Petugas Pelaksana", "type": "text", "required": True},
            {"name": "result", "label": "Hasil Pelaksanaan Kegiatan", "type": "textarea"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
    "rin8": {
        "code": "R.IN.8",
        "title": "Register Kegiatan Bidang Sosial Budaya dan Kemasyarakatan",
        "description": "Register kegiatan bidang sosial budaya dan kemasyarakatan.",
        "date_field": "activity_date",
        "chart_field": "sector",
        "chart_title": "Grafik Sektor",
        "fields": [
            {"name": "sector", "label": "Sektor", "type": "text", "required": True},
            {"name": "number", "label": "Nomor", "type": "text", "required": True},
            {"name": "activity_date", "label": "Tanggal", "type": "date", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "executor", "label": "Petugas Pelaksana", "type": "text", "required": True},
            {"name": "result", "label": "Hasil Pelaksanaan Kegiatan", "type": "textarea"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
    "rin9": {
        "code": "R.IN.9",
        "title": "Register Kegiatan Bidang Ekonomi dan Keuangan",
        "description": "Register kegiatan bidang ekonomi dan keuangan.",
        "date_field": "activity_date",
        "chart_field": "sector",
        "chart_title": "Grafik Sektor",
        "fields": [
            {"name": "sector", "label": "Sektor", "type": "text", "required": True},
            {"name": "number", "label": "Nomor", "type": "text", "required": True},
            {"name": "activity_date", "label": "Tanggal", "type": "date", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "executor", "label": "Petugas Pelaksana", "type": "text", "required": True},
            {"name": "result", "label": "Hasil Pelaksanaan Kegiatan", "type": "textarea"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
    "rin10": {
        "code": "R.IN.10",
        "title": "Register Kegiatan Bidang Pengamanan Pembangunan Strategis",
        "description": "Register kegiatan bidang pengamanan pembangunan strategis.",
        "date_field": "presentation_date",
        "chart_field": "sector_activity_fund",
        "chart_title": "Grafik Sektor/Kegiatan",
        "fields": [
            {"name": "sector_activity_fund", "label": "Sektor, Nama Kegiatan & Sumber Dana", "type": "textarea", "required": True},
            {"name": "agency", "label": "K/L/D/I", "type": "text"},
            {"name": "budget_ceiling", "label": "Pagu Anggaran", "type": "text"},
            {"name": "request_letter", "label": "Nomor dan Tanggal Surat Permohonan", "type": "textarea"},
            {"name": "presentation_place", "label": "Tempat Paparan Pemohon - Tempat", "type": "text"},
            {"name": "presentation_date", "label": "Tempat Paparan Pemohon - Tanggal", "type": "date", "required": True},
            {"name": "intelligence_study", "label": "Telaahan Intelijen", "type": "textarea"},
            {"name": "accepted_reason", "label": "Tindak Lanjut Permohonan - Diterima", "type": "textarea"},
            {"name": "rejected_reason", "label": "Tindak Lanjut Permohonan - Ditolak", "type": "textarea"},
            {"name": "walpam_order_number_date", "label": "Surat Perintah Walpam - No/Tanggal", "type": "textarea"},
            {"name": "walpam_officers", "label": "Surat Perintah Walpam - Nama Petugas", "type": "textarea"},
            {"name": "contract_value", "label": "Nilai Kontrak", "type": "text"},
            {"name": "budget_efficiency", "label": "Efisiensi Anggaran", "type": "text"},
            {"name": "project_completed", "label": "Hasil Pelaksanaan Pengawalan dan Pengamanan - Proyek Selesai", "type": "textarea"},
            {"name": "project_stopped", "label": "Hasil Pelaksanaan Pengawalan dan Pengamanan - Penghentian", "type": "textarea"},
            {"name": "working_paper_number_date", "label": "Nomor dan Tanggal Kertas Kerja", "type": "textarea"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
    "rin11": {
        "code": "R.IN.11",
        "title": "Register Kegiatan Bidang Teknologi Informasi dan Produksi Intelijen",
        "description": "Register kegiatan bidang teknologi informasi dan produksi intelijen.",
        "date_field": "letter_date",
        "chart_field": "sector",
        "chart_title": "Grafik Sektor",
        "fields": [
            {"name": "sector", "label": "Sektor", "type": "text", "required": True},
            {"name": "sprint_disposition", "label": "Nomor/Tgl/Perihal Sprint/Disposisi", "type": "textarea", "required": True},
            {"name": "executor", "label": "Petugas Pelaksana", "type": "text", "required": True},
            {"name": "result", "label": "Hasil Pelaksanaan Kegiatan", "type": "textarea"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
            {"name": "letter_date", "label": "Tanggal Surat", "type": "date", "required": True},
        ],
    },
    "rin12": {
        "code": "R.IN.12",
        "title": "Register Operasi Intelijen Bidang Ideologi, Politik, Pertahanan dan Keamanan, Cegah Tangkal dan Pengawasan Orang Asing, Pengamanan Sumber Daya Organisasi Kejaksaan dan Pengamanan Perkara",
        "description": "Register operasi intelijen bidang ideologi, politik, pertahanan keamanan, cegah tangkal, pengawasan orang asing, dan pengamanan perkara.",
        "date_field": "letter_date",
        "chart_field": "sector",
        "chart_title": "Grafik Sektor",
        "fields": [
            {"name": "sector", "label": "Sektor", "type": "text", "required": True},
            {"name": "sprint_disposition", "label": "Nomor/Tgl/Perihal Sprint/Disposisi", "type": "textarea", "required": True},
            {"name": "executor", "label": "Petugas Pelaksana", "type": "text", "required": True},
            {"name": "result", "label": "Hasil Pelaksanaan Kegiatan", "type": "textarea"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
            {"name": "letter_date", "label": "Tanggal Surat", "type": "date", "required": True},
        ],
    },
    "rin13": {
        "code": "R.IN.13",
        "title": "Register Operasi Intelijen Bidang Sosial Budaya dan Kemasyarakatan",
        "description": "Register operasi intelijen bidang sosial budaya dan kemasyarakatan.",
        "date_field": "letter_date",
        "chart_field": "sector",
        "chart_title": "Grafik Sektor",
        "fields": [
            {"name": "sector", "label": "Sektor", "type": "text", "required": True},
            {"name": "sprint_disposition", "label": "Nomor/Tanggal/Perihal Sprint/Disposisi", "type": "textarea", "required": True},
            {"name": "executor", "label": "Petugas Pelaksana", "type": "text", "required": True},
            {"name": "result", "label": "Hasil Pelaksanaan Kegiatan", "type": "textarea"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
            {"name": "letter_date", "label": "Tanggal Surat", "type": "date", "required": True},
        ],
    },
    "rin14": {
        "code": "R.IN.14",
        "title": "Register Operasi Intelijen Bidang Ekonomi dan Keuangan",
        "description": "Register operasi intelijen bidang ekonomi dan keuangan.",
        "date_field": "letter_date",
        "chart_field": "sector",
        "chart_title": "Grafik Sektor",
        "fields": [
            {"name": "satker_name", "label": "Nama Satker", "type": "satker", "required": True},
            {"name": "sector", "label": "Sektor", "type": "text", "required": True},
            {"name": "letter_number_date", "label": "Nomor Surat/Tanggal", "type": "textarea", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "officer_name", "label": "Nama Petugas", "type": "text", "required": True},
            {"name": "result", "label": "Hasil", "type": "textarea"},
            {"name": "letter_date", "label": "Tanggal Surat", "type": "date", "required": True},
        ],
    },
    "rin15": {
        "code": "R.IN.15",
        "title": "Register Operasi Intelijen Bidang Pengamanan Pembangunan Strategis",
        "description": "Register operasi intelijen bidang pengamanan pembangunan strategis.",
        "date_field": "letter_date",
        "chart_field": "sector",
        "chart_title": "Grafik Sektor",
        "fields": [
            {"name": "sector", "label": "Sektor", "type": "text", "required": True},
            {"name": "letter_number", "label": "Nomor Surat", "type": "text", "required": True},
            {"name": "letter_date", "label": "Tanggal", "type": "date", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "officer_name", "label": "Nama Petugas", "type": "text", "required": True},
            {"name": "result", "label": "Hasil", "type": "textarea"},
        ],
    },
    "rin16": {
        "code": "R.IN.16",
        "title": "Register Operasi Intelijen Bidang Teknologi Informasi dan Produksi Intelijen",
        "description": "Register operasi intelijen bidang teknologi informasi dan produksi intelijen.",
        "date_field": "letter_date",
        "chart_field": "sector",
        "chart_title": "Grafik Sektor",
        "fields": [
            {"name": "sector", "label": "Sektor", "type": "text", "required": True},
            {"name": "letter_number", "label": "Nomor Surat", "type": "text", "required": True},
            {"name": "letter_date", "label": "Tanggal", "type": "date", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "officer_name", "label": "Nama Petugas", "type": "text", "required": True},
            {"name": "result", "label": "Hasil", "type": "textarea"},
        ],
    },
    "rin17": {
        "code": "R.IN.17",
        "title": "Register Berita Masuk",
        "description": "Register berita masuk, pengirim, tujuan, waktu diterima, jumlah halaman, dan petugas.",
        "date_field": "entry_date",
        "chart_field": "sender",
        "chart_title": "Grafik Dari",
        "fields": [
            {"name": "entry_date", "label": "Tanggal", "type": "date", "required": True},
            {"name": "news_number", "label": "Nomor", "type": "text", "required": True},
            {"name": "news_date", "label": "Tanggal Berita", "type": "date", "required": True},
            {"name": "sender", "label": "Dari", "type": "text", "required": True},
            {"name": "recipient", "label": "Kepada", "type": "text", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "received_datetime", "label": "Tanggal/Jam Diterima", "type": "datetime"},
            {"name": "page_count", "label": "Jumlah Halaman", "type": "number"},
            {"name": "officer_name", "label": "Nama Petugas", "type": "text"},
        ],
    },
    "rin18": {
        "code": "R.IN.18",
        "title": "Register Berita Keluar",
        "description": "Register berita keluar, pengirim, tujuan, waktu dikirim, jumlah halaman, dan petugas.",
        "date_field": "entry_date",
        "chart_field": "recipient",
        "chart_title": "Grafik Kepada",
        "fields": [
            {"name": "entry_date", "label": "Tanggal", "type": "date", "required": True},
            {"name": "news_number", "label": "Nomor", "type": "text", "required": True},
            {"name": "news_date", "label": "Tanggal Berita", "type": "date", "required": True},
            {"name": "sender", "label": "Dari", "type": "text", "required": True},
            {"name": "recipient", "label": "Kepada", "type": "text", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "sent_datetime", "label": "Tanggal/Jam Dikirim", "type": "datetime"},
            {"name": "page_count", "label": "Jumlah Halaman", "type": "number"},
            {"name": "officer_name", "label": "Nama Petugas", "type": "text"},
        ],
    },
    "rin19": {
        "code": "R.IN.19",
        "title": "Register Telaahan Intelijen",
        "description": "Register telaahan intelijen berdasarkan pembuat, perihal, tindak lanjut, dan keterangan.",
        "date_field": "letter_date",
        "chart_field": "maker",
        "chart_title": "Grafik Pembuat",
        "fields": [
            {"name": "satker_name", "label": "Nama Satker", "type": "satker", "required": True},
            {"name": "letter_date", "label": "Tanggal Surat", "type": "date", "required": True},
            {"name": "maker", "label": "Pembuat", "type": "text", "required": True},
            {"name": "subject", "label": "Perihal", "type": "textarea", "required": True},
            {"name": "follow_up", "label": "Tindak Lanjut", "type": "textarea"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
    "rin20": {
        "code": "R.IN.20",
        "title": "Register Ekspedisi Berita",
        "description": "Register ekspedisi berita, tujuan, penerima, dan keterangan.",
        "date_field": "letter_date",
        "chart_field": "recipient",
        "chart_title": "Grafik Kepada",
        "fields": [
            {"name": "satker_name", "label": "Nama Satker", "type": "satker", "required": True},
            {"name": "news_number", "label": "Nomor Berita", "type": "text", "required": True},
            {"name": "letter_date", "label": "Tanggal Surat", "type": "date", "required": True},
            {"name": "recipient", "label": "Kepada", "type": "text", "required": True},
            {"name": "receiver_signature", "label": "Nama & TTD Penerima", "type": "textarea"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
    "rin21": {
        "code": "R.IN.21",
        "title": "Register Tamu Pos Pelayanan Hukum dan Penerimaan Pengaduan Masyarakat (PPH & PPM)",
        "description": "Register tamu Pos Pelayanan Hukum dan Penerimaan Pengaduan Masyarakat.",
        "date_field": "visit_time",
        "chart_field": "organization_name",
        "chart_title": "Grafik Organisasi",
        "fields": [
            {"name": "receiver_officer", "label": "Nama Petugas Penerima Laporan", "type": "text", "required": True},
            {"name": "visit_time", "label": "Waktu", "type": "datetime", "required": True},
            {"name": "identity", "label": "Identitas", "type": "textarea", "required": True},
            {"name": "organization_name", "label": "Nama Organisasi", "type": "text"},
            {"name": "information", "label": "Informasi yang Disampaikan", "type": "textarea", "required": True},
            {"name": "submitted_document", "label": "Surat/Dokumen yang Disampaikan", "type": "textarea"},
            {"name": "signature", "label": "Tanda Tangan", "type": "text"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
    "rin22": {
        "code": "R.IN.22",
        "title": "Register Pelaksanaan Kegiatan Penerangan Hukum/Penyuluhan Hukum",
        "description": "Register kegiatan penerangan hukum/penyuluhan hukum, sasaran, waktu, tempat, materi, dan peserta.",
        "date_field": "activity_time",
        "chart_field": "activity_target",
        "chart_title": "Grafik Sasaran Kegiatan",
        "fields": [
            {"name": "satker_name", "label": "Nama Satker", "type": "satker", "required": True},
            {"name": "warrant_letter", "label": "Surat Perintah", "type": "text", "required": True},
            {"name": "activity_target", "label": "Sasaran Kegiatan", "type": "text", "required": True},
            {"name": "activity_time", "label": "Waktu", "type": "date", "required": True},
            {"name": "place", "label": "Tempat", "type": "text", "required": True},
            {"name": "material", "label": "Materi", "type": "textarea", "required": True},
            {"name": "participant_count", "label": "Jumlah Peserta", "type": "number"},
            {"name": "remarks", "label": "Keterangan", "type": "text", "default": "Arsip"},
        ],
    },
}


def format_time_value(value):
    if not value:
        return "-"
    if hasattr(value, "strftime"):
        return value.strftime("%H.%M")
    if hasattr(value, "total_seconds"):
        total_seconds = int(value.total_seconds())
        hours = (total_seconds // 3600) % 24
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02d}.{minutes:02d}"
    return str(value)


def save_user_profile_photo(uploaded_file, old_filename=None):
    if not uploaded_file or not uploaded_file.filename:
        return old_filename
    extension = Path(secure_filename(uploaded_file.filename)).suffix.lower()
    if extension not in {".png", ".jpg", ".jpeg", ".webp"}:
        raise ValueError("Format foto harus PNG, JPG, JPEG, atau WEBP.")
    USER_AVATAR_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"avatar_{uuid.uuid4().hex[:16]}.jpg"
    with PILImage.open(uploaded_file) as image:
        image = ImageOps.exif_transpose(image).convert("RGB")
        image.thumbnail((512, 512), PILImage.LANCZOS)
        image.save(USER_AVATAR_UPLOAD_DIR / filename, "JPEG", quality=88, optimize=True)
    if old_filename and old_filename != filename:
        old_path = (USER_AVATAR_UPLOAD_DIR / old_filename).resolve()
        try:
            old_path.relative_to(USER_AVATAR_UPLOAD_DIR.resolve())
            if old_path.is_file():
                old_path.unlink()
        except ValueError:
            pass
    return filename


def excel_column_name(index):
    name = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        name = chr(65 + remainder) + name
    return name


def nihil_row_values(column_count):
    row = [""] * max(1, int(column_count or 1))
    letters = list("NIHIL")
    start = max(0, (len(row) - len(letters)) // 2)
    for offset, letter in enumerate(letters):
        if start + offset < len(row):
            row[start + offset] = letter
    return row


def nihil_table_rows_values(column_count, row_count=7):
    rows = [[""] * max(1, int(column_count or 1)) for _ in range(max(3, row_count))]
    rows[len(rows) // 2] = nihil_row_values(column_count)
    return rows


def make_xlsx(headers, rows, sheet_name="Sheet1", title=None, period_text=None):
    def cell_xml(row_number, col_number, value, style_id=0):
        cell_ref = f"{excel_column_name(col_number)}{row_number}"
        text = "" if value is None else str(value)
        style_attr = f' s="{style_id}"' if style_id else ""
        return (
            f'<c r="{cell_ref}" t="inlineStr"{style_attr}>'
            f"<is><t>{xml_escape(text)}</t></is></c>"
        )

    sheet_rows = []
    data_total = len(rows or [])
    current_row = 1
    if title:
        sheet_rows.append(f'<row r="{current_row}">' + cell_xml(current_row, 1, title, 1) + "</row>")
        current_row += 1
    if period_text:
        sheet_rows.append(f'<row r="{current_row}">' + cell_xml(current_row, 1, f"Bulan : {period_text}", 1) + "</row>")
        current_row += 1
    header_row = current_row
    sheet_rows.append(
        f'<row r="{header_row}">' + "".join(
            cell_xml(header_row, index, header, 1) for index, header in enumerate(headers, 1)
        ) + "</row>"
    )
    if not rows:
        rows = nihil_table_rows_values(len(headers))
    for row_index, row in enumerate(rows, header_row + 1):
        sheet_rows.append(
            f'<row r="{row_index}">' + "".join(
                cell_xml(row_index, col_index, value)
                for col_index, value in enumerate(row, 1)
            ) + "</row>"
        )
    signature_start_row = len(sheet_rows) + 3
    try:
        kajari = fetch_one("SELECT * FROM signatories WHERE position_code='kajari'") or {}
        kasi = fetch_one("SELECT * FROM signatories WHERE position_code='kasi_intel'") or {}
    except Exception:
        kajari, kasi = {}, {}
    def excel_acting_signer(base_signer, prefix):
        if request.args.get(f"use_{prefix}") != "1":
            return base_signer
        acting_type = str(request.args.get(f"{prefix}_type") or "").strip().lower()
        acting_label = "Plt." if acting_type == "plt" else "Plh." if acting_type == "plh" else ""
        base_position = str(base_signer.get("position_name") or "").strip()
        signer_name = str(request.args.get(f"{prefix}_name") or base_signer.get("full_name") or "").strip()
        signer_position_detail = str(request.args.get(f"{prefix}_position") or "").strip()
        signer_nip = str(request.args.get(f"{prefix}_nip") or "").strip()
        rank_nip = signer_position_detail
        if signer_nip:
            rank_nip = f"{rank_nip} NIP. {signer_nip}".strip()
        result = dict(base_signer)
        result.update(
            full_name=signer_name or "-",
            position_name=f"{acting_label} {base_position}".strip() if acting_label else base_position or "-",
            rank_nip=rank_nip or "-",
        )
        return result
    kajari = excel_acting_signer(kajari, "acting_kajari")
    kasi = excel_acting_signer(kasi, "acting_kasi")
    signature_date_raw = request.args.get("signature_date", "").strip()
    try:
        signature_date = datetime.strptime(signature_date_raw, "%Y-%m-%d").date() if signature_date_raw else date.today()
    except ValueError:
        signature_date = date.today()
    signature_date_text = f"Singaraja, {format_indonesian_date(signature_date)}"
    col_count = max(len(headers), 9)
    left_col = 2
    middle_col = max(3, col_count // 2)
    right_col = max(middle_col + 2, col_count - 2)

    signature_rows = [
        (signature_start_row, {
            left_col: "Mengetahui",
            middle_col: "Rekapitulasi",
            right_col: signature_date_text,
        }, 1),
        (signature_start_row + 1, {
            left_col: kajari.get("position_name") or "Kepala Kejaksaan Negeri Buleleng",
            middle_col: "Sisa bulan Lalu :",
            right_col: kasi.get("position_name") or "Kasi Intelijen",
        }, 0),
        (signature_start_row + 2, {
            middle_col: f"Masuk Bulan laporan : {data_total}",
        }, 0),
        (signature_start_row + 3, {
            middle_col: f"Jumlah : {data_total}",
        }, 0),
        (signature_start_row + 4, {
            middle_col: "Diselesaikan :",
        }, 0),
        (signature_start_row + 5, {
            middle_col: "Sisa Bulan Laporan :",
        }, 0),
        (signature_start_row + 8, {
            left_col: kajari.get("full_name") or "-",
            right_col: kasi.get("full_name") or "-",
        }, 1),
        (signature_start_row + 9, {
            left_col: kajari.get("rank_nip") or "-",
            right_col: kasi.get("rank_nip") or "-",
        }, 0),
    ]
    for row_number, values, default_style in signature_rows:
        sheet_rows.append(
            f'<row r="{row_number}">' + "".join(
                cell_xml(row_number, col_number, value, default_style)
                for col_number, value in sorted(values.items())
            ) + "</row>"
        )

    column_widths = "".join(
        f'<col min="{index}" max="{index}" width="{width}" customWidth="1"/>'
        for index, width in enumerate(([7, 28, 26, 28, 18, 35, 42, 42, 20] + [24] * max(0, len(headers) - 9)), 1)
    )
    sheet_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f"<cols>{column_widths}</cols><sheetData>{''.join(sheet_rows)}</sheetData></worksheet>"
    )
    workbook_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<sheets><sheet name="{xml_escape(sheet_name[:31])}" sheetId="1" r:id="rId1"/></sheets></workbook>'
    )
    styles_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        '<fonts count="2"><font><sz val="11"/><name val="Calibri"/></font>'
        '<font><b/><sz val="11"/><name val="Calibri"/></font></fonts>'
        '<fills count="2"><fill><patternFill patternType="none"/></fill>'
        '<fill><patternFill patternType="solid"><fgColor rgb="FFE8F6F5"/>'
        '<bgColor indexed="64"/></patternFill></fill></fills>'
        '<borders count="1"><border><left/><right/><top/><bottom/><diagonal/></border></borders>'
        '<cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>'
        '<cellXfs count="2"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/>'
        '<xf numFmtId="0" fontId="1" fillId="1" borderId="0" xfId="0" applyFont="1" applyFill="1"/></cellXfs>'
        '</styleSheet>'
    )

    output = BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
            '<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
            '<Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>'
            '</Types>'
        ))
        archive.writestr("_rels/.rels", (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
            '</Relationships>'
        ))
        archive.writestr("xl/workbook.xml", workbook_xml)
        archive.writestr("xl/_rels/workbook.xml.rels", (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>'
            '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
            '</Relationships>'
        ))
        archive.writestr("xl/styles.xml", styles_xml)
        archive.writestr("xl/worksheets/sheet1.xml", sheet_xml)
    output.seek(0)
    return output


def generic_register_config(slug):
    config = GENERIC_REGISTER_CONFIGS.get(slug)
    if not config:
        abort(404)
    return config


def generic_register_payload(config, source=None):
    organization = fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}
    organization_name = organization.get("organization_name") or "Kejaksaan Negeri Buleleng"
    source = source or request.form
    data = {}
    for field in config["fields"]:
        name = field["name"]
        if field.get("type") == "satker":
            value = (source.get(name) or organization_name).strip()
        else:
            value = (source.get(name) or field.get("default") or "").strip()
        if field.get("type") in {"date", "datetime"} and not value and field.get("required"):
            value = date.today().isoformat()
        data[name] = value
    return data


def generic_register_entry_date(config, payload):
    raw_value = payload.get(config.get("date_field") or "")
    if not raw_value:
        return None
    try:
        return datetime.strptime(raw_value[:10], "%Y-%m-%d").date().isoformat()
    except ValueError:
        return None


def generic_register_display_value(field, payload):
    value = payload.get(field["name"]) or "-"
    if field.get("type") == "date" and value != "-":
        try:
            return datetime.strptime(value[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
        except ValueError:
            return value
    if field.get("type") == "datetime" and value != "-":
        try:
            return datetime.strptime(value[:16], "%Y-%m-%dT%H:%M").strftime("%d/%m/%Y %H.%M")
        except ValueError:
            return value.replace("T", " ")
    return value


def generic_register_visible_fields(config):
    if config.get("code") != "R.IN.10":
        return config["fields"]
    hidden_names = {
        "presentation_date",
        "rejected_reason",
        "walpam_officers",
        "project_stopped",
    }
    return [field for field in config["fields"] if field["name"] not in hidden_names]


def generic_register_display_value_by_name(config, name, payload):
    field = next((item for item in config["fields"] if item["name"] == name), {"name": name, "type": "text"})
    return generic_register_display_value(field, payload)


def rin10_grouped_headers():
    return [
        {"label": "Sektor, Nama Kegiatan & Sumber Dana", "name": "sector_activity_fund", "rowspan": 2},
        {"label": "K/L/D/I", "name": "agency", "rowspan": 2},
        {"label": "Pagu Anggaran", "name": "budget_ceiling", "rowspan": 2},
        {"label": "Nomor dan Tanggal Surat Permohonan", "name": "request_letter", "rowspan": 2},
        {"label": "Tempat Paparan Pemohon", "children": [
            {"label": "Tempat", "name": "presentation_place"},
            {"label": "Tanggal", "name": "presentation_date"},
        ]},
        {"label": "Telaahan Intelijen", "name": "intelligence_study", "rowspan": 2},
        {"label": "Tindak Lanjut Permohonan", "children": [
            {"label": "Diterima", "name": "accepted_reason"},
            {"label": "Ditolak", "name": "rejected_reason"},
        ]},
        {"label": "Surat Perintah Walpam", "children": [
            {"label": "No/Tanggal", "name": "walpam_order_number_date"},
            {"label": "Nama Petugas", "name": "walpam_officers"},
        ]},
        {"label": "Nilai Kontrak", "name": "contract_value", "rowspan": 2},
        {"label": "Efisiensi Anggaran", "name": "budget_efficiency", "rowspan": 2},
        {"label": "Hasil Pelaksanaan Pengawalan dan Pengamanan", "children": [
            {"label": "Proyek Selesai", "name": "project_completed"},
            {"label": "Penghentian", "name": "project_stopped"},
        ]},
        {"label": "Nomor dan Tanggal Kertas Kerja", "name": "working_paper_number_date", "rowspan": 2},
        {"label": "Keterangan", "name": "remarks", "rowspan": 2},
    ]


def rin10_leaf_fields():
    leaves = []
    for item in rin10_grouped_headers():
        if item.get("name"):
            leaves.append(item["name"])
        else:
            leaves.extend(child["name"] for child in item.get("children", []))
    return leaves


def register_export_period_text(month=None, year=None, fallback="-"):
    month = (month if month is not None else request.args.get("month", "")).strip()
    year = (year if year is not None else request.args.get("year", "")).strip()
    if month.isdigit() and year.isdigit() and 1 <= int(month) <= 12:
        return f"{MONTH_NAMES_ID[int(month)]} {year}"
    return year if year.isdigit() else fallback


def generic_register_filters(config):
    params = [config["code"]]
    conditions = ["entries.register_code=%s"]
    search = request.args.get("q", "").strip()[:150]
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    if search:
        conditions.append("(entries.payload LIKE %s OR entries.chart_label LIKE %s OR users.full_name LIKE %s)")
        params.extend([f"%{search}%"] * 3)
    if month.isdigit() and 1 <= int(month) <= 12:
        conditions.append("MONTH(entries.entry_date)=%s")
        params.append(int(month))
    if year.isdigit():
        conditions.append("YEAR(entries.entry_date)=%s")
        params.append(int(year))
    return "WHERE " + " AND ".join(conditions), tuple(params), search, month, year


@app.route("/register-intelijen/<slug>")
@login_required
def register_intelijen_generic(slug):
    config = generic_register_config(slug)
    where_clause, params, search, month, year = generic_register_filters(config)
    chart_month = request.args.get("chart_month", "all").strip() or "all"
    try:
        page = max(1, int(request.args.get("page", "1")))
    except ValueError:
        page = 1
    per_page = 10
    current_year = date.today().year
    current_month = date.today().month
    chart_conditions = ["register_code=%s", "YEAR(entry_date)=%s"]
    chart_params = [config["code"], current_year]
    if chart_month.isdigit() and 1 <= int(chart_month) <= 12:
        chart_conditions.append("MONTH(entry_date)=%s")
        chart_params.append(int(chart_month))
        chart_month_name = MONTH_NAMES_ID[int(chart_month)]
    else:
        chart_month = "all"
        chart_month_name = "Semua bulan"
    chart_rows = fetch_all(
        f"""SELECT COALESCE(NULLIF(TRIM(chart_label), ''), 'Tanpa Data') AS label, COUNT(*) AS total
            FROM register_intelijen_generic_entries
            WHERE {' AND '.join(chart_conditions)}
            GROUP BY label
            ORDER BY total DESC, label ASC
            LIMIT 12""",
        tuple(chart_params),
    )
    max_chart_total = max([int(row["total"] or 0) for row in chart_rows] or [0])
    total_current_year = int(fetch_one(
        "SELECT COUNT(*) AS total FROM register_intelijen_generic_entries WHERE register_code=%s AND YEAR(entry_date)=%s",
        (config["code"], current_year),
    )["total"] or 0)
    total_current_month = int(fetch_one(
        """SELECT COUNT(*) AS total FROM register_intelijen_generic_entries
           WHERE register_code=%s AND YEAR(entry_date)=%s AND MONTH(entry_date)=%s""",
        (config["code"], current_year, current_month),
    )["total"] or 0)
    filtered_total = int(fetch_one(
        f"""SELECT COUNT(*) AS total
            FROM register_intelijen_generic_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}""",
        params,
    )["total"] or 0)
    total_pages = max(1, (filtered_total + per_page - 1) // per_page)
    page = min(page, total_pages)
    entries = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_generic_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.entry_date DESC, entries.id DESC
            LIMIT %s OFFSET %s""",
        tuple(list(params) + [per_page, (page - 1) * per_page]),
    )
    for entry in entries:
        try:
            entry["payload_data"] = json.loads(entry.get("payload") or "{}")
        except json.JSONDecodeError:
            entry["payload_data"] = {}
    years = fetch_all(
        """SELECT DISTINCT YEAR(entry_date) AS report_year
           FROM register_intelijen_generic_entries
           WHERE register_code=%s AND entry_date IS NOT NULL
           ORDER BY report_year DESC""",
        (config["code"],),
    )
    organization_name = (fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}).get("organization_name") or "Kejaksaan Negeri Buleleng"
    return render_template(
        "register_intelijen_generic_list.html",
        active=f"register_intelijen_{slug}",
        slug=slug,
        config=config,
        entries=entries,
        search=search,
        filter_month=month,
        filter_year=year,
        page=page,
        per_page=per_page,
        total_pages=total_pages,
        filtered_total=filtered_total,
        month_names=MONTH_NAMES_ID,
        years=sorted({int(row["report_year"]) for row in years if row["report_year"]} | {current_year}, reverse=True),
        current_year=current_year,
        current_month=current_month,
        current_month_name=MONTH_NAMES_ID[current_month],
        chart_month=chart_month,
        chart_month_name=chart_month_name,
        chart_rows=chart_rows,
        max_chart_total=max_chart_total,
        total_current_year=total_current_year,
        total_current_month=total_current_month,
        organization_name=organization_name,
        display_value=generic_register_display_value,
        visible_fields=generic_register_visible_fields(config),
        display_value_by_name=generic_register_display_value_by_name,
        rin10_grouped_headers=rin10_grouped_headers(),
        rin10_leaf_fields=rin10_leaf_fields(),
        today=date.today().isoformat(),
    )


@app.route("/register-intelijen/<slug>/create", methods=["GET", "POST"])
@login_required
def create_register_intelijen_generic(slug):
    config = generic_register_config(slug)
    if request.method == "POST":
        data = generic_register_payload(config)
        missing = [field["label"] for field in config["fields"] if field.get("required") and not data.get(field["name"])]
        if missing:
            flash("Field wajib diisi: " + ", ".join(missing), "error")
        else:
            entry_date = generic_register_entry_date(config, data)
            chart_label = data.get(config.get("chart_field") or "") or ""
            with get_db().cursor() as cursor:
                cursor.execute(
                    """INSERT INTO register_intelijen_generic_entries
                       (register_code, entry_date, chart_label, payload, created_by)
                       VALUES (%s,%s,%s,%s,%s)""",
                    (config["code"], entry_date, chart_label, json.dumps(data, ensure_ascii=False), session["user_id"]),
                )
            flash(f"Data {config['code']} berhasil disimpan.", "success")
            return redirect(url_for("register_intelijen_generic", slug=slug))
    else:
        data = generic_register_payload(config, {})
    return render_template(
        "register_intelijen_generic_form.html",
        active=f"register_intelijen_{slug}",
        slug=slug,
        config=config,
        data=data,
        entry=None,
    )


@app.route("/register-intelijen/<slug>/<int:entry_id>/edit", methods=["GET", "POST"])
@login_required
def edit_register_intelijen_generic(slug, entry_id):
    config = generic_register_config(slug)
    entry = fetch_one(
        "SELECT * FROM register_intelijen_generic_entries WHERE id=%s AND register_code=%s",
        (entry_id, config["code"]),
    )
    if not entry:
        flash(f"Data {config['code']} tidak ditemukan.", "error")
        return redirect(url_for("register_intelijen_generic", slug=slug))
    if session.get("role") != "admin" and entry["created_by"] != session["user_id"]:
        flash(f"Anda tidak dapat mengubah data {config['code']} ini.", "error")
        return redirect(url_for("register_intelijen_generic", slug=slug))
    if request.method == "POST":
        data = generic_register_payload(config)
        missing = [field["label"] for field in config["fields"] if field.get("required") and not data.get(field["name"])]
        if missing:
            flash("Field wajib diisi: " + ", ".join(missing), "error")
        else:
            entry_date = generic_register_entry_date(config, data)
            chart_label = data.get(config.get("chart_field") or "") or ""
            with get_db().cursor() as cursor:
                cursor.execute(
                    """UPDATE register_intelijen_generic_entries
                       SET entry_date=%s, chart_label=%s, payload=%s
                       WHERE id=%s AND register_code=%s""",
                    (entry_date, chart_label, json.dumps(data, ensure_ascii=False), entry_id, config["code"]),
                )
            flash(f"Data {config['code']} berhasil diperbarui.", "success")
            return redirect(url_for("register_intelijen_generic", slug=slug))
    else:
        try:
            data = json.loads(entry.get("payload") or "{}")
        except json.JSONDecodeError:
            data = {}
        data = generic_register_payload(config, data)
    return render_template(
        "register_intelijen_generic_form.html",
        active=f"register_intelijen_{slug}",
        slug=slug,
        config=config,
        data=data,
        entry=entry,
    )


@app.post("/register-intelijen/<slug>/<int:entry_id>/delete")
@login_required
def delete_register_intelijen_generic(slug, entry_id):
    config = generic_register_config(slug)
    entry = fetch_one(
        "SELECT id,created_by FROM register_intelijen_generic_entries WHERE id=%s AND register_code=%s",
        (entry_id, config["code"]),
    )
    if not entry:
        flash(f"Data {config['code']} tidak ditemukan.", "error")
    elif session.get("role") != "admin" and entry["created_by"] != session["user_id"]:
        flash(f"Anda tidak dapat menghapus data {config['code']} ini.", "error")
    else:
        with get_db().cursor() as cursor:
            cursor.execute("DELETE FROM register_intelijen_generic_entries WHERE id=%s AND register_code=%s", (entry_id, config["code"]))
        flash(f"Data {config['code']} berhasil dihapus.", "success")
    return redirect(url_for("register_intelijen_generic", slug=slug))


@app.get("/register-intelijen/<slug>/export-excel")
@login_required
def export_register_intelijen_generic_excel(slug):
    config = generic_register_config(slug)
    where_clause, params, *_ = generic_register_filters(config)
    entries = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_generic_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.entry_date ASC, entries.id ASC""",
        params,
    )
    if config["code"] == "R.IN.10":
        headers = ["No"]
        for item in rin10_grouped_headers():
            if item.get("children"):
                headers.extend([f"{item['label']} - {child['label']}" for child in item["children"]])
            else:
                headers.append(item["label"])
        excel_field_names = rin10_leaf_fields()
    else:
        visible_fields = generic_register_visible_fields(config)
        headers = ["No"] + [field["label"] for field in visible_fields]
        excel_field_names = [field["name"] for field in visible_fields]
    rows = []
    for index, entry in enumerate(entries, 1):
        try:
            payload = json.loads(entry.get("payload") or "{}")
        except json.JSONDecodeError:
            payload = {}
        rows.append([index] + [generic_register_display_value_by_name(config, field_name, payload) for field_name in excel_field_names])
    workbook = make_xlsx(
        headers,
        rows,
        config["code"],
        title=f"{config['code']} {config['title'].upper()}",
        period_text=register_export_period_text(),
    )
    return send_file(
        workbook,
        as_attachment=True,
        download_name=f"{config['code'].replace('.', '')}-{date.today().strftime('%Y%m%d')}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.get("/register-intelijen/<slug>/export-pdf")
@login_required
def export_register_intelijen_generic_pdf(slug):
    config = generic_register_config(slug)
    where_clause, params, *_ = generic_register_filters(config)
    entries = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_generic_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.entry_date ASC, entries.id ASC""",
        params,
    )
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    period_text = f"{MONTH_NAMES_ID[int(month)]} {year}" if month.isdigit() and year.isdigit() and 1 <= int(month) <= 12 else (year if year.isdigit() else "-")
    use_scan_signature = request.args.get("use_scan_signature") == "1"
    use_digital_stamp = request.args.get("use_digital_stamp") == "1"
    use_acting_kajari = request.args.get("use_acting_kajari") == "1"
    use_acting_kasi = request.args.get("use_acting_kasi") == "1"
    signature_date_raw = request.args.get("signature_date", "").strip()
    try:
        signature_date = datetime.strptime(signature_date_raw, "%Y-%m-%d").date() if signature_date_raw else date.today()
    except ValueError:
        signature_date = date.today()
    organization = fetch_one("SELECT * FROM organization_settings WHERE id=1") or {}
    kajari = fetch_one("SELECT * FROM signatories WHERE position_code='kajari'") or {}
    kasi = fetch_one("SELECT * FROM signatories WHERE position_code='kasi_intel'") or {}

    def acting_signer(base_signer, prefix, enabled):
        if not enabled:
            return base_signer
        acting_type = str(request.args.get(f"{prefix}_type") or "").strip().lower()
        acting_label = "Plt." if acting_type == "plt" else "Plh." if acting_type == "plh" else ""
        base_position = str(base_signer.get("position_name") or "").strip()
        signer_name = str(request.args.get(f"{prefix}_name") or base_signer.get("full_name") or "").strip()
        signer_position_detail = str(request.args.get(f"{prefix}_position") or "").strip()
        signer_nip = str(request.args.get(f"{prefix}_nip") or "").strip()
        rank_nip = signer_position_detail
        if signer_nip:
            rank_nip = f"{rank_nip} NIP. {signer_nip}".strip()
        result = dict(base_signer)
        result.update(
            full_name=signer_name or "-",
            position_name=f"{acting_label} {base_position}".strip() if acting_label else base_position or "-",
            rank_nip=rank_nip or "-",
            signature_image=None,
        )
        return result

    kajari = acting_signer(kajari, "acting_kajari", use_acting_kajari)
    kasi = acting_signer(kasi, "acting_kasi", use_acting_kasi)

    def safe_upload_path(base_dir, filename):
        if not filename:
            return None
        path = (base_dir / filename).resolve()
        try:
            path.relative_to(base_dir.resolve())
        except ValueError:
            return None
        return path if path.is_file() else None

    def pdf_image(path, max_width, max_height):
        if not path:
            return Spacer(1, max_height)
        try:
            with PILImage.open(path) as image:
                width, height = image.size
        except Exception:
            return Spacer(1, max_height)
        if not width or not height:
            return Spacer(1, max_height)
        ratio = min(max_width / width, max_height / height)
        return Image(str(path), width=width * ratio, height=height * ratio)

    def signature_image_block(scan_path, stamp_path=None):
        flowables = []
        if stamp_path:
            flowables.append(pdf_image(stamp_path, 2.35 * cm, 2.35 * cm))
        if scan_path:
            flowables.append(pdf_image(scan_path, 6.25 * cm, 2.75 * cm))
        if not flowables:
            flowables = [Spacer(1, 2.75 * cm)]
        image_block = Table([flowables], colWidths=([2.65 * cm] if stamp_path else []) + ([6.9 * cm] if scan_path else [6.9 * cm]), hAlign="CENTER")
        image_block.setStyle(TableStyle([
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return image_block

    output = BytesIO()
    page_size = landscape((8.5 * inch, 13 * inch))
    page_width, page_height = page_size
    title_text = f"{config['code']} {config['title'].upper()}"
    title_canvas_style = ParagraphStyle(
        "GenericRegisterCanvasTitle",
        fontName=PDF_FONT_BOLD,
        fontSize=12,
        leading=13.5,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=0,
    )
    header_width = page_width - 2 * cm
    title_height = Paragraph(escape(title_text), title_canvas_style).wrap(header_width, 2.2 * cm)[1]
    dynamic_top_margin = max(2.75 * cm, 1.45 * cm + title_height + 0.86 * cm)
    document = SimpleDocTemplate(output, pagesize=page_size, rightMargin=1*cm, leftMargin=1*cm, topMargin=dynamic_top_margin, bottomMargin=1.1*cm, allowSplitting=True)
    styles = getSampleStyleSheet()
    cell_style = ParagraphStyle("GenericCell", parent=styles["Normal"], fontName=PDF_FONT_NAME, fontSize=8, leading=9.4, splitLongWords=True)
    header_style = ParagraphStyle("GenericHeader", parent=cell_style, fontName=PDF_FONT_BOLD, alignment=TA_CENTER)
    signature_style = ParagraphStyle("GenericSignature", parent=styles["Normal"], fontName=PDF_FONT_NAME, fontSize=10, leading=12, alignment=TA_CENTER, spaceBefore=0, spaceAfter=0)
    signature_name_style = ParagraphStyle("GenericSignatureName", parent=signature_style, fontName=PDF_FONT_BOLD)
    recap_title_style = ParagraphStyle("GenericRecapTitle", parent=signature_style, fontName=PDF_FONT_BOLD, alignment=TA_LEFT)
    recap_style = ParagraphStyle("GenericRecap", parent=signature_style, alignment=TA_LEFT)
    def pdf_cell(value, style=cell_style):
        return Paragraph(escape(str(value if value is not None else "-")).replace("\n", "<br/>"), style)
    if config["code"] == "R.IN.10":
        top_header = [pdf_cell("No", header_style)]
        sub_header = [pdf_cell("", header_style)]
        for item in rin10_grouped_headers():
            if item.get("children"):
                top_header.append(pdf_cell(item["label"], header_style))
                top_header.extend([pdf_cell("", header_style) for _ in item["children"][1:]])
                sub_header.extend([pdf_cell(child["label"], header_style) for child in item["children"]])
            else:
                top_header.append(pdf_cell(item["label"], header_style))
                sub_header.append(pdf_cell("", header_style))
        table_data = [top_header, sub_header]
        leaf_names = rin10_leaf_fields()
        for index, entry in enumerate(entries, 1):
            try:
                payload = json.loads(entry.get("payload") or "{}")
            except json.JSONDecodeError:
                payload = {}
            table_data.append([pdf_cell(index, header_style)] + [
                pdf_cell(generic_register_display_value_by_name(config, field_name, payload))
                for field_name in leaf_names
            ])
        if len(table_data) == 2:
            for nihil_row in nihil_table_rows_values(len(table_data[0])):
                table_data.append([pdf_cell(value, header_style if value else cell_style) for value in nihil_row])
    else:
        visible_fields = generic_register_visible_fields(config)
        table_data = [[pdf_cell("No", header_style)] + [pdf_cell(field["label"], header_style) for field in visible_fields]]
        for index, entry in enumerate(entries, 1):
            try:
                payload = json.loads(entry.get("payload") or "{}")
            except json.JSONDecodeError:
                payload = {}
            table_data.append([pdf_cell(index, header_style)] + [pdf_cell(generic_register_display_value(field, payload)) for field in visible_fields])
        if len(table_data) == 1:
            for nihil_row in nihil_table_rows_values(len(table_data[0])):
                table_data.append([pdf_cell(value, header_style if value else cell_style) for value in nihil_row])
    usable_width = page_width - document.leftMargin - document.rightMargin
    first_width = 0.8 * cm
    data_col_count = len(table_data[0]) - 1
    other_width = (usable_width - first_width) / max(1, data_col_count)
    table = LongTable(table_data, repeatRows=2 if config["code"] == "R.IN.10" else 1, colWidths=[first_width] + [other_width] * data_col_count, hAlign="LEFT", splitByRow=True)
    table_style_commands = [
        ("GRID", (0, 0), (-1, -1), 0.5, "#000000"),
        ("BACKGROUND", (0, 0), (-1, 1 if config["code"] == "R.IN.10" else 0), "#E8EDF3"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("VALIGN", (0, 0), (-1, 1 if config["code"] == "R.IN.10" else 0), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, 1 if config["code"] == "R.IN.10" else 0), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2.5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2.5),
        ("TOPPADDING", (0, 0), (-1, -1), 3.2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.2),
    ]
    if config["code"] == "R.IN.10":
        table_style_commands.append(("SPAN", (0, 0), (0, 1)))
        current_col = 1
        for item in rin10_grouped_headers():
            if item.get("children"):
                span_to = current_col + len(item["children"]) - 1
                table_style_commands.append(("SPAN", (current_col, 0), (span_to, 0)))
                current_col = span_to + 1
            else:
                table_style_commands.append(("SPAN", (current_col, 0), (current_col, 1)))
                current_col += 1
    table.setStyle(TableStyle(table_style_commands))
    def draw_register_header(pdf_canvas, _document):
        pdf_canvas.saveState()
        header_x = _document.leftMargin
        header_width = page_width - _document.leftMargin - _document.rightMargin
        pdf_canvas.setFillColorRGB(0, 0, 0)
        title_paragraph = Paragraph(escape(title_text), title_canvas_style)
        title_top_y = page_height - 1.45 * cm
        title_height = title_paragraph.wrap(header_width, 2.05 * cm)[1]
        title_bottom_y = title_top_y - title_height
        title_paragraph.drawOn(pdf_canvas, header_x, title_bottom_y)
        pdf_canvas.setFont(PDF_FONT_BOLD, 12)
        pdf_canvas.drawString(header_x, title_bottom_y - 0.42 * cm, f"Bulan : {period_text}")
        pdf_canvas.restoreState()

    def signature_flowables(signer, prefix_text, date_text=None, include_stamp=False):
        scan_path = safe_upload_path(SIGNATORY_UPLOAD_DIR, signer.get("signature_image")) if use_scan_signature else None
        stamp_path = safe_upload_path(ORGANIZATION_UPLOAD_DIR, organization.get("digital_stamp")) if include_stamp else None
        lines = []
        if date_text:
            lines.append(Paragraph(escape(date_text), signature_style))
        if prefix_text:
            lines.append(Paragraph(prefix_text, signature_style))
        lines.extend([
            Paragraph(escape(str(signer.get("position_name") or "-")), signature_style),
            signature_image_block(scan_path, stamp_path),
            Paragraph(f"<u>{escape(str(signer.get('full_name') or '-'))}</u>", signature_name_style),
            Paragraph(escape(str(signer.get("rank_nip") or "-")), signature_style),
        ])
        return lines

    report_total = len(entries)
    recap_block = [
        Paragraph("Rekapitulasi", recap_title_style),
        Paragraph("Sisa bulan Lalu :", recap_style),
        Paragraph(f"Masuk Bulan laporan : {report_total}", recap_style),
        Paragraph(f"Jumlah : {report_total}", recap_style),
        Paragraph("Diselesaikan :", recap_style),
        Paragraph("Sisa Bulan Laporan :", recap_style),
    ]
    signature_table = Table(
        [[
            signature_flowables(kajari, "Mengetahui", include_stamp=use_digital_stamp),
            recap_block,
            signature_flowables(kasi, "", date_text=f"Singaraja, {format_indonesian_date(signature_date)}"),
        ]],
        colWidths=[10.2 * cm, 5.0 * cm, 10.2 * cm],
        hAlign="CENTER",
    )
    signature_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    document.build([table, Spacer(1, 0.7 * cm), KeepTogether(signature_table)], onFirstPage=draw_register_header, onLaterPages=draw_register_header)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f"{config['code'].replace('.', '')}-{date.today().strftime('%Y%m%d')}.pdf", mimetype="application/pdf")


def register_intelijen_rin1_form_data():
    form = request.form
    today = date.today().isoformat()
    return {
        "received_date": (form.get("received_date") or today).strip(),
        "received_time": (form.get("received_time") or "").strip(),
        "incoming_letter_number": (form.get("incoming_letter_number") or "").strip(),
        "incoming_letter_date": (form.get("incoming_letter_date") or today).strip(),
        "sender_name": (form.get("sender_name") or "").strip(),
        "subject": (form.get("subject") or "").strip(),
        "disposition_date": (form.get("disposition_date") or today).strip(),
        "disposition_content": (form.get("disposition_content") or "").strip(),
        "follow_up": (form.get("follow_up") or "").strip(),
        "remarks": (form.get("remarks") or "Arsip").strip() or "Arsip",
    }


@app.route("/register-intelijen/rin1")
@login_required
def register_intelijen_rin1():
    params = []
    conditions = []
    search = request.args.get("q", "").strip()[:150]
    sender = request.args.get("sender", "").strip()[:150]
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    chart_month = request.args.get("chart_month", "all").strip() or "all"
    try:
        page = max(1, int(request.args.get("page", "1")))
    except ValueError:
        page = 1
    per_page = 10
    if search:
        search_value = f"%{search}%"
        conditions.append(
            """(entries.incoming_letter_number LIKE %s OR entries.sender_name LIKE %s
                OR entries.subject LIKE %s OR entries.disposition_content LIKE %s
                OR entries.follow_up LIKE %s OR entries.remarks LIKE %s
                OR users.full_name LIKE %s)"""
        )
        params.extend([search_value] * 7)
    if sender:
        conditions.append("entries.sender_name LIKE %s")
        params.append(f"%{sender}%")
    if month.isdigit() and 1 <= int(month) <= 12:
        conditions.append("MONTH(entries.received_date)=%s")
        params.append(int(month))
    if year.isdigit():
        conditions.append("YEAR(entries.received_date)=%s")
        params.append(int(year))
    where_clause = f"WHERE {' AND '.join(conditions)}" if conditions else ""
    current_year = date.today().year
    current_month = date.today().month
    chart_conditions = ["YEAR(received_date)=%s"]
    chart_params = [current_year]
    if chart_month.isdigit() and 1 <= int(chart_month) <= 12:
        chart_conditions.append("MONTH(received_date)=%s")
        chart_params.append(int(chart_month))
        chart_month_name = MONTH_NAMES_ID[int(chart_month)]
    else:
        chart_month = "all"
        chart_month_name = "Semua bulan"
    sender_chart_rows = fetch_all(
        f"""SELECT COALESCE(NULLIF(TRIM(sender_name), ''), 'Tanpa Asal Surat') AS sender_label,
                   COUNT(*) AS total
            FROM register_intelijen_rin1_entries
            WHERE {' AND '.join(chart_conditions)}
            GROUP BY sender_label
            ORDER BY total DESC, sender_label ASC
            LIMIT 12""",
        tuple(chart_params),
    )
    max_sender_total = max([int(row["total"] or 0) for row in sender_chart_rows] or [0])
    total_current_year = int(fetch_one(
        "SELECT COUNT(*) AS total FROM register_intelijen_rin1_entries WHERE YEAR(received_date)=%s",
        (current_year,),
    )["total"] or 0)
    total_current_month = int(fetch_one(
        """SELECT COUNT(*) AS total FROM register_intelijen_rin1_entries
           WHERE YEAR(received_date)=%s AND MONTH(received_date)=%s""",
        (current_year, current_month),
    )["total"] or 0)
    filtered_total = int(fetch_one(
        f"""SELECT COUNT(*) AS total
            FROM register_intelijen_rin1_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}""",
        tuple(params),
    )["total"] or 0)
    total_pages = max(1, (filtered_total + per_page - 1) // per_page)
    page = min(page, total_pages)
    rows = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_rin1_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.received_date DESC, entries.received_time DESC, entries.id DESC
            LIMIT %s OFFSET %s""",
        tuple(params + [per_page, (page - 1) * per_page]),
    )
    for row in rows:
        row["received_time_text"] = format_time_value(row.get("received_time"))
    total = fetch_one("SELECT COUNT(*) AS total FROM register_intelijen_rin1_entries")["total"]
    years = fetch_all(
        """SELECT DISTINCT YEAR(received_date) AS report_year
           FROM register_intelijen_rin1_entries
           WHERE received_date IS NOT NULL
           ORDER BY report_year DESC"""
    )
    organization_name = (fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}).get("organization_name") or "Kejaksaan Negeri Buleleng"
    return render_template(
        "register_intelijen_rin1_list.html",
        active="register_intelijen_rin1",
        entries=rows,
        total_entries=int(total or 0),
        filtered_total=filtered_total,
        search=search,
        filter_sender=sender,
        page=page,
        per_page=per_page,
        total_pages=total_pages,
        filter_month=month,
        filter_year=year,
        month_names=MONTH_NAMES_ID,
        organization_name=organization_name,
        current_year=current_year,
        current_month=current_month,
        current_month_name=MONTH_NAMES_ID[current_month],
        chart_month=chart_month,
        chart_month_name=chart_month_name,
        sender_chart_rows=sender_chart_rows,
        max_sender_total=max_sender_total,
        total_current_year=total_current_year,
        total_current_month=total_current_month,
        today=date.today().isoformat(),
        years=sorted(
            {int(row["report_year"]) for row in years if row["report_year"]} | {date.today().year},
            reverse=True,
        ),
    )


@app.route("/register-intelijen/rin1/create", methods=["GET", "POST"])
@login_required
def create_register_intelijen_rin1():
    if request.method == "POST":
        data = register_intelijen_rin1_form_data()
        if not data["incoming_letter_number"] or not data["sender_name"] or not data["subject"]:
            flash("Nomor surat, asal surat, dan perihal wajib diisi.", "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    """INSERT INTO register_intelijen_rin1_entries
                       (register_code, received_date, received_time, incoming_letter_number,
                        incoming_letter_date, sender_name, subject, disposition_date,
                        disposition_content, follow_up, remarks, created_by)
                       VALUES ('R.IN.1', %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                    (
                        data["received_date"], data["received_time"] or None,
                        data["incoming_letter_number"], data["incoming_letter_date"],
                        data["sender_name"], data["subject"], data["disposition_date"] or None,
                        data["disposition_content"], data["follow_up"], data["remarks"],
                        session["user_id"],
                    ),
                )
            flash("Data R.IN.1 berhasil disimpan.", "success")
            return redirect(url_for("register_intelijen_rin1"))
    else:
        data = {
            "received_date": date.today().isoformat(),
            "received_time": "",
            "incoming_letter_number": "",
            "incoming_letter_date": date.today().isoformat(),
            "sender_name": "",
            "subject": "",
            "disposition_date": date.today().isoformat(),
            "disposition_content": "",
            "follow_up": "",
            "remarks": "Arsip",
        }
    return render_template(
        "register_intelijen_rin1_form.html",
        active="register_intelijen_rin1",
        form_title="Buat Register Surat Masuk R.IN.1",
        form_description="Input data surat masuk intelijen beserta disposisi dan tindak lanjut.",
        data=data,
        entry=None,
    )


@app.route("/register-intelijen/rin1/<int:entry_id>/edit", methods=["GET", "POST"])
@login_required
def edit_register_intelijen_rin1(entry_id):
    entry = fetch_one("SELECT * FROM register_intelijen_rin1_entries WHERE id=%s", (entry_id,))
    if not entry:
        flash("Data R.IN.1 tidak ditemukan.", "error")
        return redirect(url_for("register_intelijen_rin1"))
    if session.get("role") != "admin" and entry["created_by"] != session["user_id"]:
        flash("Anda tidak dapat mengubah data R.IN.1 ini.", "error")
        return redirect(url_for("register_intelijen_rin1"))
    if request.method == "POST":
        data = register_intelijen_rin1_form_data()
        if not data["incoming_letter_number"] or not data["sender_name"] or not data["subject"]:
            flash("Nomor surat, asal surat, dan perihal wajib diisi.", "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    """UPDATE register_intelijen_rin1_entries
                       SET received_date=%s, received_time=%s, incoming_letter_number=%s,
                           incoming_letter_date=%s, sender_name=%s, subject=%s,
                           disposition_date=%s, disposition_content=%s, follow_up=%s, remarks=%s
                       WHERE id=%s""",
                    (
                        data["received_date"], data["received_time"] or None,
                        data["incoming_letter_number"], data["incoming_letter_date"],
                        data["sender_name"], data["subject"], data["disposition_date"] or None,
                        data["disposition_content"], data["follow_up"], data["remarks"], entry_id,
                    ),
                )
            flash("Data R.IN.1 berhasil diperbarui.", "success")
            return redirect(url_for("register_intelijen_rin1"))
    else:
        data = {
            "received_date": entry["received_date"].isoformat() if entry.get("received_date") else date.today().isoformat(),
            "received_time": format_time_value(entry.get("received_time")).replace(".", ":") if entry.get("received_time") else "",
            "incoming_letter_number": entry.get("incoming_letter_number") or "",
            "incoming_letter_date": entry["incoming_letter_date"].isoformat() if entry.get("incoming_letter_date") else date.today().isoformat(),
            "sender_name": entry.get("sender_name") or "",
            "subject": entry.get("subject") or "",
            "disposition_date": entry["disposition_date"].isoformat() if entry.get("disposition_date") else date.today().isoformat(),
            "disposition_content": entry.get("disposition_content") or "",
            "follow_up": entry.get("follow_up") or "",
            "remarks": entry.get("remarks") or "Arsip",
        }
    return render_template(
        "register_intelijen_rin1_form.html",
        active="register_intelijen_rin1",
        form_title="Edit Register Surat Masuk R.IN.1",
        form_description="Perbarui data surat masuk intelijen.",
        data=data,
        entry=entry,
    )


@app.post("/register-intelijen/rin1/<int:entry_id>/delete")
@login_required
def delete_register_intelijen_rin1(entry_id):
    entry = fetch_one("SELECT id,created_by FROM register_intelijen_rin1_entries WHERE id=%s", (entry_id,))
    if not entry:
        flash("Data R.IN.1 tidak ditemukan.", "error")
    elif session.get("role") != "admin" and entry["created_by"] != session["user_id"]:
        flash("Anda tidak dapat menghapus data R.IN.1 ini.", "error")
    else:
        with get_db().cursor() as cursor:
            cursor.execute("DELETE FROM register_intelijen_rin1_entries WHERE id=%s", (entry_id,))
        flash("Data R.IN.1 berhasil dihapus.", "success")
    return redirect(url_for("register_intelijen_rin1"))


def register_rin1_filter_params():
    params = []
    conditions = []
    search = request.args.get("q", "").strip()[:150]
    sender = request.args.get("sender", "").strip()[:150]
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    if search:
        search_value = f"%{search}%"
        conditions.append(
            """(entries.incoming_letter_number LIKE %s OR entries.sender_name LIKE %s
                OR entries.subject LIKE %s OR entries.disposition_content LIKE %s
                OR entries.follow_up LIKE %s OR entries.remarks LIKE %s
                OR users.full_name LIKE %s)"""
        )
        params.extend([search_value] * 7)
    if sender:
        conditions.append("entries.sender_name LIKE %s")
        params.append(f"%{sender}%")
    if month.isdigit() and 1 <= int(month) <= 12:
        conditions.append("MONTH(entries.received_date)=%s")
        params.append(int(month))
    if year.isdigit():
        conditions.append("YEAR(entries.received_date)=%s")
        params.append(int(year))
    return f"WHERE {' AND '.join(conditions)}" if conditions else "", tuple(params)


@app.get("/register-intelijen/rin1/export-excel")
@login_required
def export_register_intelijen_rin1_excel():
    where_clause, params = register_rin1_filter_params()
    entries = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_rin1_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.received_date ASC, entries.received_time ASC, entries.id ASC""",
        params,
    )
    headers = [
        "No", "Tanggal Penerimaan", "Jam Penerimaan", "Nomor Surat Masuk",
        "Tanggal Surat Masuk", "Asal Surat", "Perihal", "Tanggal Disposisi",
        "Isi Disposisi", "Tindak Lanjut", "Ket",
    ]
    rows = []
    for index, entry in enumerate(entries, 1):
        rows.append([
            index,
            entry["received_date"].strftime("%d/%m/%Y") if entry.get("received_date") else "",
            f"{format_time_value(entry.get('received_time'))} Wita" if entry.get("received_time") else "",
            entry.get("incoming_letter_number") or "",
            entry["incoming_letter_date"].strftime("%d/%m/%Y") if entry.get("incoming_letter_date") else "",
            entry.get("sender_name") or "",
            entry.get("subject") or "",
            entry["disposition_date"].strftime("%d/%m/%Y") if entry.get("disposition_date") else "",
            entry.get("disposition_content") or "",
            entry.get("follow_up") or "",
            entry.get("remarks") or "",
        ])
    workbook = make_xlsx(
        headers,
        rows,
        "R.IN.1",
        title="R.IN.1 REGISTER SURAT MASUK",
        period_text=register_export_period_text(),
    )
    return send_file(
        workbook,
        as_attachment=True,
        download_name=f"RIN1-SURAT-MASUK-{date.today().strftime('%Y%m%d')}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.get("/register-intelijen/rin1/export-pdf")
@login_required
def export_register_intelijen_rin1_pdf():
    where_clause, params = register_rin1_filter_params()
    entries = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_rin1_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.received_date ASC, entries.received_time ASC, entries.id ASC""",
        params,
    )
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    use_scan_signature = request.args.get("use_scan_signature") == "1"
    use_digital_stamp = request.args.get("use_digital_stamp") == "1"
    use_acting_kajari = request.args.get("use_acting_kajari") == "1"
    use_acting_kasi = request.args.get("use_acting_kasi") == "1"
    signature_date_raw = request.args.get("signature_date", "").strip()
    try:
        signature_date = datetime.strptime(signature_date_raw, "%Y-%m-%d").date() if signature_date_raw else date.today()
    except ValueError:
        signature_date = date.today()
    if month.isdigit() and year.isdigit() and 1 <= int(month) <= 12:
        period_text = f"{MONTH_NAMES_ID[int(month)]} {year}"
    elif entries:
        received_dates = [row.get("received_date") for row in entries if row.get("received_date")]
        first_date, last_date = min(received_dates), max(received_dates)
        if first_date and last_date and first_date.month == last_date.month and first_date.year == last_date.year:
            period_text = f"{MONTH_NAMES_ID[first_date.month]} {first_date.year}"
        elif first_date and last_date and first_date.year == last_date.year:
            period_text = f"{MONTH_NAMES_ID[first_date.month]} - {MONTH_NAMES_ID[last_date.month]} {first_date.year}"
        else:
            period_text = f"{MONTH_NAMES_ID[first_date.month]} {first_date.year} - {MONTH_NAMES_ID[last_date.month]} {last_date.year}"
    else:
        period_text = year if year.isdigit() else "-"

    organization = fetch_one("SELECT * FROM organization_settings WHERE id=1") or {}
    kajari = fetch_one("SELECT * FROM signatories WHERE position_code='kajari'") or {}
    kasi = fetch_one("SELECT * FROM signatories WHERE position_code='kasi_intel'") or {}

    def acting_signer(base_signer, prefix, enabled):
        if not enabled:
            return base_signer
        acting_type = str(request.args.get(f"{prefix}_type") or "").strip().lower()
        acting_label = "Plt." if acting_type == "plt" else "Plh." if acting_type == "plh" else ""
        base_position = str(base_signer.get("position_name") or "").strip()
        signer_name = str(request.args.get(f"{prefix}_name") or base_signer.get("full_name") or "").strip()
        signer_position_detail = str(request.args.get(f"{prefix}_position") or "").strip()
        signer_nip = str(request.args.get(f"{prefix}_nip") or "").strip()
        rank_nip = signer_position_detail
        if signer_nip:
            rank_nip = f"{rank_nip} NIP. {signer_nip}".strip()
        result = dict(base_signer)
        result.update(
            full_name=signer_name or "-",
            position_name=f"{acting_label} {base_position}".strip() if acting_label else base_position or "-",
            rank_nip=rank_nip or "-",
            signature_image=None,
        )
        return result

    kajari = acting_signer(kajari, "acting_kajari", use_acting_kajari)
    kasi = acting_signer(kasi, "acting_kasi", use_acting_kasi)

    def safe_upload_path(base_dir, filename):
        if not filename:
            return None
        path = (base_dir / filename).resolve()
        try:
            path.relative_to(base_dir.resolve())
        except ValueError:
            return None
        return path if path.is_file() else None

    def pdf_image(path, max_width, max_height):
        if not path:
            return Spacer(1, max_height)
        try:
            with PILImage.open(path) as image:
                width, height = image.size
        except Exception:
            return Spacer(1, max_height)
        if not width or not height:
            return Spacer(1, max_height)
        ratio = min(max_width / width, max_height / height)
        return Image(str(path), width=width * ratio, height=height * ratio)

    def signature_image_block(scan_path, stamp_path=None):
        flowables = []
        if stamp_path:
            flowables.append(pdf_image(stamp_path, 2.35 * cm, 2.35 * cm))
        if scan_path:
            flowables.append(pdf_image(scan_path, 6.25 * cm, 2.75 * cm))
        if not flowables:
            flowables = [Spacer(1, 2.75 * cm)]
        image_block = Table([flowables], colWidths=([2.65 * cm] if stamp_path else []) + ([6.9 * cm] if scan_path else [6.9 * cm]), hAlign="CENTER")
        image_block.setStyle(TableStyle([
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return image_block

    output = BytesIO()
    page_size = landscape((8.5 * inch, 13 * inch))
    page_width, page_height = page_size
    document = SimpleDocTemplate(
        output,
        pagesize=page_size,
        rightMargin=1 * cm,
        leftMargin=1 * cm,
        topMargin=3 * cm,
        bottomMargin=1.1 * cm,
        allowSplitting=True,
    )
    styles = getSampleStyleSheet()
    cell_style = ParagraphStyle("Rin1Cell", parent=styles["Normal"], fontName=PDF_FONT_NAME, fontSize=8, leading=9.4, spaceBefore=0, spaceAfter=0, splitLongWords=True)
    center_style = ParagraphStyle("Rin1CellCenter", parent=cell_style, alignment=TA_CENTER)
    header_style = ParagraphStyle("Rin1Header", parent=cell_style, fontName=PDF_FONT_BOLD, alignment=TA_CENTER, leading=9.6)
    signature_style = ParagraphStyle("Rin1Signature", parent=styles["Normal"], fontName=PDF_FONT_NAME, fontSize=10, leading=12, alignment=TA_CENTER, spaceBefore=0, spaceAfter=0)
    signature_name_style = ParagraphStyle("Rin1SignatureName", parent=signature_style, fontName=PDF_FONT_BOLD)
    recap_title_style = ParagraphStyle("Rin1RecapTitle", parent=signature_style, fontName=PDF_FONT_BOLD, alignment=TA_LEFT)
    recap_style = ParagraphStyle("Rin1Recap", parent=signature_style, alignment=TA_LEFT)

    def rin1_cell(value, style=cell_style):
        return Paragraph(escape(str(value if value is not None else "-")).replace("\n", "<br/>"), style)

    data = [
        [
            rin1_cell("No", header_style),
            rin1_cell("WAKTU PENERIMAAN SURAT", header_style), "",
            rin1_cell("SURAT MASUK", header_style), "",
            rin1_cell("ASAL SURAT", header_style),
            rin1_cell("PERIHAL", header_style),
            rin1_cell("DISPOSISI", header_style), "", "",
            rin1_cell("KETERANGAN", header_style),
        ],
        [
            "",
            rin1_cell("Tanggal", header_style),
            rin1_cell("Jam", header_style),
            rin1_cell("Nomor", header_style),
            rin1_cell("Tanggal", header_style),
            "",
            "",
            rin1_cell("Tanggal", header_style),
            rin1_cell("Isi", header_style),
            rin1_cell("Tindak Lanjut", header_style),
            "",
        ],
    ]
    for index, entry in enumerate(entries, 1):
        data.append([
            rin1_cell(index, center_style),
            rin1_cell(entry["received_date"].strftime("%d/%m/%Y") if entry.get("received_date") else "-", center_style),
            rin1_cell(f"{format_time_value(entry.get('received_time'))} Wita" if entry.get("received_time") else "-", center_style),
            rin1_cell(entry.get("incoming_letter_number") or "-"),
            rin1_cell(entry["incoming_letter_date"].strftime("%d/%m/%Y") if entry.get("incoming_letter_date") else "-", center_style),
            rin1_cell(entry.get("sender_name") or "-"),
            rin1_cell(entry.get("subject") or "-"),
            rin1_cell(entry["disposition_date"].strftime("%d/%m/%Y") if entry.get("disposition_date") else "-", center_style),
            rin1_cell(entry.get("disposition_content") or "-"),
            rin1_cell(entry.get("follow_up") or "-"),
            rin1_cell(entry.get("remarks") or "-", center_style),
        ])
    if len(data) == 2:
        for nihil_row in nihil_table_rows_values(11):
            data.append([rin1_cell(value, header_style if value else cell_style) for value in nihil_row])
    base_widths = [0.7*cm, 1.55*cm, 1.45*cm, 3.15*cm, 1.55*cm, 3.0*cm, 3.55*cm, 1.55*cm, 2.35*cm, 3.7*cm, 1.35*cm]
    usable_width = page_width - document.leftMargin - document.rightMargin
    width_scale = usable_width / sum(base_widths)
    table = LongTable(
        data,
        repeatRows=2,
        colWidths=[item * width_scale for item in base_widths],
        hAlign="LEFT",
        splitByRow=True,
        splitInRow=False,
    )
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, "#000000"),
        ("SPAN", (0, 0), (0, 1)),
        ("SPAN", (1, 0), (2, 0)),
        ("SPAN", (3, 0), (4, 0)),
        ("SPAN", (5, 0), (5, 1)),
        ("SPAN", (6, 0), (6, 1)),
        ("SPAN", (7, 0), (9, 0)),
        ("SPAN", (10, 0), (10, 1)),
        ("BACKGROUND", (0, 0), (-1, 1), "#E8EDF3"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("VALIGN", (0, 0), (-1, 1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, 1), "CENTER"),
        ("ALIGN", (0, 2), (2, -1), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2.5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2.5),
        ("TOPPADDING", (0, 0), (-1, -1), 3.2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.2),
    ]))

    def signature_flowables(signer, prefix_text, date_text=None, include_stamp=False):
        scan_path = safe_upload_path(SIGNATORY_UPLOAD_DIR, signer.get("signature_image")) if use_scan_signature else None
        stamp_path = safe_upload_path(ORGANIZATION_UPLOAD_DIR, organization.get("digital_stamp")) if include_stamp else None
        lines = []
        if date_text:
            lines.append(Paragraph(escape(date_text), signature_style))
        if prefix_text:
            lines.append(Paragraph(prefix_text, signature_style))
        lines.extend([
            Paragraph(escape(str(signer.get("position_name") or "-")), signature_style),
            signature_image_block(scan_path, stamp_path),
            Paragraph(f"<u>{escape(str(signer.get('full_name') or '-'))}</u>", signature_name_style),
            Paragraph(escape(str(signer.get("rank_nip") or "-")), signature_style),
        ])
        return lines

    report_total = len(entries)
    recap_block = [
        Paragraph("Rekapitulasi", recap_title_style),
        Paragraph("Sisa bulan Lalu :", recap_style),
        Paragraph(f"Masuk Bulan laporan : {report_total}", recap_style),
        Paragraph(f"Jumlah : {report_total}", recap_style),
        Paragraph("Diselesaikan :", recap_style),
        Paragraph("Sisa Bulan Laporan :", recap_style),
    ]
    signature_table = Table(
        [[
            signature_flowables(kajari, "Mengetahui", include_stamp=use_digital_stamp),
            recap_block,
            signature_flowables(kasi, "", date_text=f"Singaraja, {format_indonesian_date(signature_date)}"),
        ]],
        colWidths=[10.2 * cm, 5.0 * cm, 10.2 * cm],
        hAlign="CENTER",
    )
    signature_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))

    def draw_register_header(pdf_canvas, _document):
        pdf_canvas.saveState()
        header_x = _document.leftMargin
        pdf_canvas.setFillColorRGB(0, 0, 0)
        pdf_canvas.setFont(PDF_FONT_BOLD, 12)
        pdf_canvas.drawString(header_x, page_height - 2.05 * cm, "R.IN.1 REGISTER SURAT MASUK")
        pdf_canvas.drawString(header_x, page_height - 2.48 * cm, f"Bulan : {period_text}")
        pdf_canvas.restoreState()

    story = [table, Spacer(1, 0.7 * cm), KeepTogether(signature_table)]
    document.build(story, onFirstPage=draw_register_header, onLaterPages=draw_register_header)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f"RIN1-SURAT-MASUK-{date.today().strftime('%Y%m%d')}.pdf", mimetype="application/pdf")


def register_intelijen_form_data():
    organization = fetch_one("SELECT * FROM organization_settings WHERE id=1") or {}
    form = request.form
    report_date = (form.get("report_date") or date.today().isoformat()).strip()
    field_code = (form.get("field_code") or "").strip()
    time_value = (form.get("received_time") or "").strip()
    default_org_name = (organization.get("organization_name") or "Kejaksaan Negeri Buleleng").strip()
    source_name = (form.get("source_name") or default_org_name).strip()
    info_value = (form.get("information_value") or "A1").strip().upper()
    notes = (form.get("notes") or "").strip()
    remarks = (form.get("remarks") or "Arsip").strip() or "Arsip"
    data = {
        "report_date": report_date,
        "field_code": field_code if field_code in ISSUE_CODES else "",
        "received_time": time_value,
        "source_name": source_name,
        "information_value": info_value if info_value in REGISTER_INFORMATION_VALUES else "A1",
        "information_description": (form.get("information_description") or "").strip(),
        "notes": notes,
        "disposition": (form.get("disposition") or "").strip(),
        "follow_up": (form.get("follow_up") or "").strip(),
        "remarks": remarks,
        "satker_name": default_org_name,
    }
    return data


@app.route("/register-intelijen")
@login_required
def register_intelijen():
    params = []
    conditions = []
    search = request.args.get("q", "").strip()[:150]
    bidang = request.args.get("bidang", "").strip()
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    chart_month = request.args.get("chart_month", "").strip()
    try:
        page = max(1, int(request.args.get("page", "1")))
    except ValueError:
        page = 1
    per_page = 10
    if search:
        search_value = f"%{search}%"
        conditions.append(
            """(entries.field_code LIKE %s OR entries.source_name LIKE %s
                OR entries.information_value LIKE %s
                OR entries.information_description LIKE %s OR entries.notes LIKE %s
                OR entries.disposition LIKE %s OR entries.follow_up LIKE %s
                OR entries.remarks LIKE %s OR users.full_name LIKE %s)"""
        )
        params.extend([search_value] * 9)
    if bidang in ISSUE_CODES:
        conditions.append("entries.field_code=%s")
        params.append(bidang)
    if month.isdigit() and 1 <= int(month) <= 12:
        conditions.append("MONTH(entries.report_date)=%s")
        params.append(int(month))
    if year.isdigit():
        conditions.append("YEAR(entries.report_date)=%s")
        params.append(int(year))
    where_clause = f"WHERE {' AND '.join(conditions)}" if conditions else ""
    filtered_total = fetch_one(
        f"""SELECT COUNT(*) AS total
            FROM register_intelijen_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}""",
        tuple(params),
    )["total"]
    total_pages = max(1, (int(filtered_total or 0) + per_page - 1) // per_page)
    page = min(page, total_pages)
    offset = (page - 1) * per_page
    rows = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.report_date DESC, entries.received_time DESC, entries.id DESC
            LIMIT %s OFFSET %s""",
        tuple(params + [per_page, offset]),
    )
    for row in rows:
        row["received_time_text"] = format_time_value(row.get("received_time"))
    total = fetch_one("SELECT COUNT(*) AS total FROM register_intelijen_entries")["total"]
    current_year = date.today().year
    current_month = date.today().month
    total_current_year = fetch_one(
        "SELECT COUNT(*) AS total FROM register_intelijen_entries WHERE YEAR(report_date)=%s",
        (current_year,),
    )["total"]
    total_current_month = fetch_one(
        """SELECT COUNT(*) AS total FROM register_intelijen_entries
           WHERE YEAR(report_date)=%s AND MONTH(report_date)=%s""",
        (current_year, current_month),
    )["total"]
    chart_month_number = 0 if chart_month in {"", "all"} else (
        int(chart_month) if chart_month.isdigit() and 1 <= int(chart_month) <= 12 else 0
    )
    chart_conditions = ["YEAR(report_date)=%s"]
    chart_params = [current_year]
    if chart_month_number:
        chart_conditions.append("MONTH(report_date)=%s")
        chart_params.append(chart_month_number)
    chart_where = " AND ".join(chart_conditions)
    notes_chart_rows = fetch_all(
        """SELECT COALESCE(NULLIF(TRIM(notes), ''), 'Tanpa Catatan') AS note_label,
                  COUNT(*) AS total
           FROM register_intelijen_entries
           WHERE """ + chart_where + """
           GROUP BY note_label
           ORDER BY total DESC, note_label ASC
           LIMIT 10""",
        tuple(chart_params),
    )
    max_note_total = max([int(row["total"] or 0) for row in notes_chart_rows] or [0])
    years = fetch_all(
        """SELECT DISTINCT YEAR(report_date) AS report_year
           FROM register_intelijen_entries
           WHERE report_date IS NOT NULL
           ORDER BY report_year DESC"""
    )
    organization_name = (fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}).get("organization_name") or "Kejaksaan Negeri Buleleng"
    return render_template(
        "register_intelijen_list.html",
        active="register_intelijen",
        register_type="R.IN.3",
        entries=rows,
        total_entries=int(total or 0),
        filtered_total=int(filtered_total or 0),
        search=search,
        page=page,
        per_page=per_page,
        total_pages=total_pages,
        filter_bidang=bidang,
        filter_month=month,
        filter_year=year,
        current_month=date.today().month,
        current_year=date.today().year,
        today=date.today().isoformat(),
        current_month_name=MONTH_NAMES_ID[current_month],
        chart_month="all" if chart_month_number == 0 else str(chart_month_number),
        chart_month_name=MONTH_NAMES_ID[chart_month_number] if chart_month_number else "Semua Bulan",
        total_current_year=int(total_current_year or 0),
        total_current_month=int(total_current_month or 0),
        notes_chart_rows=notes_chart_rows,
        max_note_total=max_note_total,
        month_names=MONTH_NAMES_ID,
        issue_code_labels=ISSUE_CODE_LABELS,
        years=sorted(
            {int(row["report_year"]) for row in years if row["report_year"]} | {date.today().year},
            reverse=True,
        ),
        organization_name=organization_name,
    )


@app.route("/register-intelijen/create", methods=["GET", "POST"])
@login_required
def create_register_intelijen():
    organization = fetch_one("SELECT * FROM organization_settings WHERE id=1") or {}
    if request.method == "POST":
        data = register_intelijen_form_data()
        if not data["field_code"] or not data["information_description"]:
            flash("Bidang dan uraian informasi wajib diisi.", "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    """INSERT INTO register_intelijen_entries
                    (register_code, report_date, field_code, received_time, source_name,
                     information_value, information_description, notes, disposition,
                     follow_up, remarks, created_by)
                    VALUES ('R.IN.3', %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                    (
                        data["report_date"], data["field_code"], data["received_time"] or None,
                        data["source_name"], data["information_value"], data["information_description"],
                        data["notes"], data["disposition"], data["follow_up"], data["remarks"],
                        session["user_id"],
                    ),
                )
            flash("Data Register Intelijen berhasil disimpan.", "success")
            return redirect(url_for("register_intelijen"))
    else:
        data = {
            "report_date": date.today().isoformat(),
            "field_code": "",
            "received_time": "",
            "source_name": organization.get("organization_name") or "Kejaksaan Negeri Buleleng",
            "information_value": "A1",
            "information_description": "",
            "notes": "",
            "disposition": "",
            "follow_up": "",
            "remarks": "Arsip",
            "satker_name": organization.get("organization_name") or "Kejaksaan Negeri Buleleng",
        }
    return render_template(
        "register_intelijen_form.html",
        active="register_intelijen",
        register_type="R.IN.3",
        form_title="Buat Register Intelijen",
        form_description="Input data R.IN.3 Register Kerja Intelijen.",
        data=data,
        issue_code_labels=ISSUE_CODE_LABELS,
        information_values=REGISTER_INFORMATION_VALUES,
        entry=None,
    )


@app.route("/register-intelijen/<int:entry_id>/edit", methods=["GET", "POST"])
@login_required
def edit_register_intelijen(entry_id):
    entry = fetch_one("SELECT * FROM register_intelijen_entries WHERE id=%s", (entry_id,))
    if not entry:
        flash("Data register tidak ditemukan.", "error")
        return redirect(url_for("register_intelijen"))
    if session.get("role") != "admin" and entry["created_by"] != session["user_id"]:
        flash("Anda tidak dapat mengubah data register ini.", "error")
        return redirect(url_for("register_intelijen"))
    if request.method == "POST":
        data = register_intelijen_form_data()
        if not data["field_code"] or not data["information_description"]:
            flash("Bidang dan uraian informasi wajib diisi.", "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    """UPDATE register_intelijen_entries
                       SET report_date=%s, field_code=%s, received_time=%s, source_name=%s,
                           information_value=%s, information_description=%s, notes=%s,
                           disposition=%s, follow_up=%s, remarks=%s
                       WHERE id=%s""",
                    (
                        data["report_date"], data["field_code"], data["received_time"] or None,
                        data["source_name"], data["information_value"], data["information_description"],
                        data["notes"], data["disposition"], data["follow_up"], data["remarks"], entry_id,
                    ),
                )
            flash("Data Register Intelijen berhasil diperbarui.", "success")
            return redirect(url_for("register_intelijen"))
    else:
        data = {
            "report_date": entry["report_date"].isoformat() if entry.get("report_date") else date.today().isoformat(),
            "field_code": entry.get("field_code") or "",
            "received_time": format_time_value(entry.get("received_time")).replace(".", ":") if entry.get("received_time") else "",
            "source_name": entry.get("source_name") or "",
            "information_value": entry.get("information_value") or "A1",
            "information_description": entry.get("information_description") or "",
            "notes": entry.get("notes") or "",
            "disposition": entry.get("disposition") or "",
            "follow_up": entry.get("follow_up") or "",
            "remarks": entry.get("remarks") or "Arsip",
            "satker_name": (fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}).get("organization_name") or "Kejaksaan Negeri Buleleng",
        }
    return render_template(
        "register_intelijen_form.html",
        active="register_intelijen",
        register_type="R.IN.3",
        form_title="Edit Register Intelijen",
        form_description="Perbarui data R.IN.3 Register Kerja Intelijen.",
        data=data,
        issue_code_labels=ISSUE_CODE_LABELS,
        information_values=REGISTER_INFORMATION_VALUES,
        entry=entry,
    )


@app.post("/register-intelijen/<int:entry_id>/delete")
@login_required
def delete_register_intelijen(entry_id):
    entry = fetch_one("SELECT id,created_by FROM register_intelijen_entries WHERE id=%s", (entry_id,))
    if not entry:
        flash("Data register tidak ditemukan.", "error")
    elif session.get("role") != "admin" and entry["created_by"] != session["user_id"]:
        flash("Anda tidak dapat menghapus data register ini.", "error")
    else:
        with get_db().cursor() as cursor:
            cursor.execute("DELETE FROM register_intelijen_entries WHERE id=%s", (entry_id,))
        flash("Data Register Intelijen berhasil dihapus.", "success")
    return redirect(url_for("register_intelijen"))


def register_intelijen_rin5_form_data():
    organization = fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}
    form = request.form
    product_date = (form.get("intelligence_product_date") or date.today().isoformat()).strip()
    field_code = (form.get("field_code") or "").strip()
    return {
        "satker_name": organization.get("organization_name") or "Kejaksaan Negeri Buleleng",
        "intelligence_product_type": (form.get("intelligence_product_type") or "").strip(),
        "intelligence_product_number": (form.get("intelligence_product_number") or "").strip(),
        "intelligence_product_date": product_date,
        "field_code": field_code if field_code in ISSUE_CODES else "",
        "subject": (form.get("subject") or "").strip(),
        "leader_disposition": (form.get("leader_disposition") or "").strip(),
        "remarks": (form.get("remarks") or "").strip(),
    }


@app.route("/register-intelijen/rin5")
@login_required
def register_intelijen_rin5():
    params = []
    conditions = []
    search = request.args.get("q", "").strip()[:150]
    bidang = request.args.get("bidang", "").strip()
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    try:
        page = max(1, int(request.args.get("page", "1")))
    except ValueError:
        page = 1
    per_page = 10
    if search:
        search_value = f"%{search}%"
        conditions.append(
            """(entries.intelligence_product_type LIKE %s
                OR entries.intelligence_product_number LIKE %s
                OR entries.field_code LIKE %s OR entries.subject LIKE %s
                OR entries.leader_disposition LIKE %s OR entries.remarks LIKE %s
                OR users.full_name LIKE %s)"""
        )
        params.extend([search_value] * 7)
    if bidang in ISSUE_CODES:
        conditions.append("entries.field_code=%s")
        params.append(bidang)
    if month.isdigit() and 1 <= int(month) <= 12:
        conditions.append("MONTH(entries.intelligence_product_date)=%s")
        params.append(int(month))
    if year.isdigit():
        conditions.append("YEAR(entries.intelligence_product_date)=%s")
        params.append(int(year))
    where_clause = f"WHERE {' AND '.join(conditions)}" if conditions else ""
    current_year = date.today().year
    current_month = date.today().month
    chart_month = request.args.get("chart_month", "all").strip() or "all"
    chart_conditions = ["YEAR(intelligence_product_date)=%s"]
    chart_params = [current_year]
    if chart_month.isdigit() and 1 <= int(chart_month) <= 12:
        chart_conditions.append("MONTH(intelligence_product_date)=%s")
        chart_params.append(int(chart_month))
        chart_month_name = MONTH_NAMES_ID[int(chart_month)]
    else:
        chart_month = "all"
        chart_month_name = "Semua bulan"
    product_chart_rows = fetch_all(
        f"""SELECT COALESCE(NULLIF(TRIM(intelligence_product_type), ''), 'Tanpa Jenis') AS product_label,
                   COUNT(*) AS total
            FROM register_intelijen_rin5_entries
            WHERE {' AND '.join(chart_conditions)}
            GROUP BY product_label
            ORDER BY total DESC, product_label ASC
            LIMIT 12""",
        tuple(chart_params),
    )
    max_product_total = max([int(row["total"] or 0) for row in product_chart_rows] or [0])
    total_current_year = int(fetch_one(
        """SELECT COUNT(*) AS total FROM register_intelijen_rin5_entries
           WHERE YEAR(intelligence_product_date)=%s""",
        (current_year,),
    )["total"] or 0)
    total_current_month = int(fetch_one(
        """SELECT COUNT(*) AS total FROM register_intelijen_rin5_entries
           WHERE YEAR(intelligence_product_date)=%s AND MONTH(intelligence_product_date)=%s""",
        (current_year, current_month),
    )["total"] or 0)
    filtered_total = int(fetch_one(
        f"""SELECT COUNT(*) AS total
            FROM register_intelijen_rin5_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}""",
        tuple(params),
    )["total"] or 0)
    total_pages = max(1, (filtered_total + per_page - 1) // per_page)
    page = min(page, total_pages)
    rows = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_rin5_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.intelligence_product_date DESC, entries.id DESC
            LIMIT %s OFFSET %s""",
        tuple(params + [per_page, (page - 1) * per_page]),
    )
    total = fetch_one("SELECT COUNT(*) AS total FROM register_intelijen_rin5_entries")["total"]
    years = fetch_all(
        """SELECT DISTINCT YEAR(intelligence_product_date) AS report_year
           FROM register_intelijen_rin5_entries
           WHERE intelligence_product_date IS NOT NULL
           ORDER BY report_year DESC"""
    )
    organization_name = (fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}).get("organization_name") or "Kejaksaan Negeri Buleleng"
    return render_template(
        "register_intelijen_rin5_list.html",
        active="register_intelijen_rin5",
        entries=rows,
        total_entries=int(total or 0),
        filtered_total=filtered_total,
        search=search,
        page=page,
        per_page=per_page,
        total_pages=total_pages,
        filter_bidang=bidang,
        filter_month=month,
        filter_year=year,
        month_names=MONTH_NAMES_ID,
        issue_code_labels=ISSUE_CODE_LABELS,
        organization_name=organization_name,
        current_year=current_year,
        current_month=current_month,
        current_month_name=MONTH_NAMES_ID[current_month],
        chart_month=chart_month,
        chart_month_name=chart_month_name,
        product_chart_rows=product_chart_rows,
        max_product_total=max_product_total,
        total_current_year=total_current_year,
        total_current_month=total_current_month,
        today=date.today().isoformat(),
        years=sorted(
            {int(row["report_year"]) for row in years if row["report_year"]} | {date.today().year},
            reverse=True,
        ),
    )


@app.get("/register-intelijen/rin5/export-excel")
@login_required
def export_register_intelijen_rin5_excel():
    params = []
    conditions = []
    search = request.args.get("q", "").strip()[:150]
    bidang = request.args.get("bidang", "").strip()
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    if search:
        search_value = f"%{search}%"
        conditions.append(
            """(entries.intelligence_product_type LIKE %s
                OR entries.intelligence_product_number LIKE %s
                OR entries.field_code LIKE %s OR entries.subject LIKE %s
                OR entries.leader_disposition LIKE %s OR entries.remarks LIKE %s
                OR users.full_name LIKE %s)"""
        )
        params.extend([search_value] * 7)
    if bidang in ISSUE_CODES:
        conditions.append("entries.field_code=%s")
        params.append(bidang)
    if month.isdigit() and 1 <= int(month) <= 12:
        conditions.append("MONTH(entries.intelligence_product_date)=%s")
        params.append(int(month))
    if year.isdigit():
        conditions.append("YEAR(entries.intelligence_product_date)=%s")
        params.append(int(year))
    where_clause = f"WHERE {' AND '.join(conditions)}" if conditions else ""
    entries = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_rin5_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.intelligence_product_date ASC, entries.id ASC""",
        tuple(params),
    )
    organization_name = (fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}).get("organization_name") or "Kejaksaan Negeri Buleleng"
    headers = [
        "No", "Nama Satker", "Jenis Produk Intelijen", "Nomor Produk Intelijen",
        "Tanggal Produk Intelijen", "Subdit/Seksi/Subseksi", "Perihal",
        "Disposisi Pimpinan", "Keterangan",
    ]
    rows = []
    for index, entry in enumerate(entries, 1):
        product_date = entry.get("intelligence_product_date")
        if isinstance(product_date, datetime):
            product_date = product_date.date()
        product_date_text = product_date.strftime("%d/%m/%Y") if isinstance(product_date, date) else "-"
        field_text = f"{entry.get('field_code') or '-'} - {ISSUE_CODE_LABELS.get(entry.get('field_code'), '-')}"
        rows.append([
            index,
            organization_name,
            entry.get("intelligence_product_type") or "",
            entry.get("intelligence_product_number") or "",
            product_date_text,
            field_text,
            entry.get("subject") or "",
            entry.get("leader_disposition") or "",
            entry.get("remarks") or "",
        ])
    workbook = make_xlsx(
        headers,
        rows,
        "R.IN.5",
        title="R.IN.5 REGISTER PRODUK INTELIJEN",
        period_text=register_export_period_text(),
    )
    filename = f"RIN5-REGISTER-PRODUK-INTELIJEN-{date.today().strftime('%Y%m%d')}.xlsx"
    return send_file(
        workbook,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.get("/register-intelijen/export-excel")
@login_required
def export_register_intelijen_excel():
    params = []
    conditions = []
    search = request.args.get("q", "").strip()[:150]
    bidang = request.args.get("bidang", "").strip()
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    if search:
        search_value = f"%{search}%"
        conditions.append(
            """(entries.field_code LIKE %s OR entries.source_name LIKE %s
                OR entries.information_value LIKE %s
                OR entries.information_description LIKE %s OR entries.notes LIKE %s
                OR entries.disposition LIKE %s OR entries.follow_up LIKE %s
                OR entries.remarks LIKE %s OR users.full_name LIKE %s)"""
        )
        params.extend([search_value] * 9)
    if bidang in ISSUE_CODES:
        conditions.append("entries.field_code=%s")
        params.append(bidang)
    if month.isdigit() and 1 <= int(month) <= 12:
        conditions.append("MONTH(entries.report_date)=%s")
        params.append(int(month))
    if year.isdigit():
        conditions.append("YEAR(entries.report_date)=%s")
        params.append(int(year))
    where_clause = f"WHERE {' AND '.join(conditions)}" if conditions else ""
    entries = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.report_date ASC, entries.received_time ASC, entries.id ASC""",
        tuple(params),
    )
    organization_name = (fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}).get("organization_name") or "Kejaksaan Negeri Buleleng"
    headers = [
        "No", "Nama Satker", "Tanggal Laporan", "Bidang", "Waktu Diterima",
        "Sumber/Bapul", "Nilai Data/Informasi", "Uraian Informasi",
        "Catatan", "Disposisi/Tindakan", "Tindaklanjut", "Keterangan",
    ]
    rows = []
    for index, entry in enumerate(entries, 1):
        report_date = entry.get("report_date")
        if isinstance(report_date, datetime):
            report_date = report_date.date()
        report_date_text = report_date.strftime("%d/%m/%Y") if isinstance(report_date, date) else "-"
        field_text = f"{entry.get('field_code') or '-'} - {ISSUE_CODE_LABELS.get(entry.get('field_code'), '-')}"
        rows.append([
            index,
            organization_name,
            report_date_text,
            field_text,
            format_time_value(entry.get("received_time")),
            entry.get("source_name") or "",
            entry.get("information_value") or "",
            entry.get("information_description") or "",
            entry.get("notes") or "",
            entry.get("disposition") or "",
            entry.get("follow_up") or "",
            entry.get("remarks") or "",
        ])
    workbook = make_xlsx(
        headers,
        rows,
        "R.IN.3",
        title="R.IN.3 REGISTER KERJA INTELIJEN",
        period_text=register_export_period_text(),
    )
    filename = f"RIN3-REGISTER-KERJA-INTELIJEN-{date.today().strftime('%Y%m%d')}.xlsx"
    return send_file(
        workbook,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.get("/register-intelijen/rin5/export-pdf")
@login_required
def export_register_intelijen_rin5_pdf():
    params = []
    conditions = []
    search = request.args.get("q", "").strip()[:150]
    bidang = request.args.get("bidang", "").strip()
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    use_scan_signature = request.args.get("use_scan_signature") == "1"
    use_digital_stamp = request.args.get("use_digital_stamp") == "1"
    use_acting_kajari = request.args.get("use_acting_kajari") == "1"
    use_acting_kasi = request.args.get("use_acting_kasi") == "1"
    signature_date_raw = request.args.get("signature_date", "").strip()
    try:
        signature_date = datetime.strptime(signature_date_raw, "%Y-%m-%d").date() if signature_date_raw else date.today()
    except ValueError:
        signature_date = date.today()
    if search:
        search_value = f"%{search}%"
        conditions.append(
            """(entries.intelligence_product_type LIKE %s
                OR entries.intelligence_product_number LIKE %s
                OR entries.field_code LIKE %s OR entries.subject LIKE %s
                OR entries.leader_disposition LIKE %s OR entries.remarks LIKE %s
                OR users.full_name LIKE %s)"""
        )
        params.extend([search_value] * 7)
    if bidang in ISSUE_CODES:
        conditions.append("entries.field_code=%s")
        params.append(bidang)
    if month.isdigit() and 1 <= int(month) <= 12:
        conditions.append("MONTH(entries.intelligence_product_date)=%s")
        params.append(int(month))
    if year.isdigit():
        conditions.append("YEAR(entries.intelligence_product_date)=%s")
        params.append(int(year))
    where_clause = f"WHERE {' AND '.join(conditions)}" if conditions else ""
    rows = fetch_all(
        f"""SELECT entries.*, users.full_name AS creator_full_name
            FROM register_intelijen_rin5_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.intelligence_product_date ASC, entries.id ASC""",
        tuple(params),
    )
    organization = fetch_one("SELECT * FROM organization_settings WHERE id=1") or {}
    organization_name = organization.get("organization_name") or "Kejaksaan Negeri Buleleng"
    kajari = fetch_one("SELECT * FROM signatories WHERE position_code='kajari'") or {}
    kasi = fetch_one("SELECT * FROM signatories WHERE position_code='kasi_intel'") or {}

    def acting_signer(base_signer, prefix, enabled):
        if not enabled:
            return base_signer
        acting_type = str(request.args.get(f"{prefix}_type") or "").strip().lower()
        acting_label = "Plt." if acting_type == "plt" else "Plh." if acting_type == "plh" else ""
        base_position = str(base_signer.get("position_name") or "").strip()
        signer_name = str(request.args.get(f"{prefix}_name") or base_signer.get("full_name") or "").strip()
        signer_position_detail = str(request.args.get(f"{prefix}_position") or "").strip()
        signer_nip = str(request.args.get(f"{prefix}_nip") or "").strip()
        rank_nip = signer_position_detail
        if signer_nip:
            rank_nip = f"{rank_nip} NIP. {signer_nip}".strip()
        acting_position = f"{acting_label} {base_position}".strip() if acting_label else base_position
        result = dict(base_signer)
        result.update(
            full_name=signer_name or "-",
            position_name=acting_position or "-",
            rank_nip=rank_nip or "-",
            signature_image=None,
        )
        return result

    kajari = acting_signer(kajari, "acting_kajari", use_acting_kajari)
    kasi = acting_signer(kasi, "acting_kasi", use_acting_kasi)

    buffer = BytesIO()
    folio_landscape = landscape((8.5 * inch, 13 * inch))
    page_width, page_height = folio_landscape
    document = SimpleDocTemplate(
        buffer,
        pagesize=folio_landscape,
        leftMargin=1 * cm,
        rightMargin=1 * cm,
        topMargin=3 * cm,
        bottomMargin=1.1 * cm,
        allowSplitting=True,
    )
    product_dates = []
    for row in rows:
        product_date = row.get("intelligence_product_date")
        if isinstance(product_date, datetime):
            product_date = product_date.date()
        if isinstance(product_date, date):
            product_dates.append(product_date)
    if product_dates:
        first_date, last_date = min(product_dates), max(product_dates)
        if first_date.year == last_date.year and first_date.month == last_date.month:
            period_text = f"{MONTH_NAMES_ID[first_date.month]} {first_date.year}"
        elif first_date.year == last_date.year:
            period_text = f"{MONTH_NAMES_ID[first_date.month]} - {MONTH_NAMES_ID[last_date.month]} {first_date.year}"
        else:
            period_text = f"{MONTH_NAMES_ID[first_date.month]} {first_date.year} - {MONTH_NAMES_ID[last_date.month]} {last_date.year}"
    elif month.isdigit() and year.isdigit() and 1 <= int(month) <= 12:
        period_text = f"{MONTH_NAMES_ID[int(month)]} {year}"
    else:
        period_text = year if year.isdigit() else "-"

    cell_style = ParagraphStyle(
        "Rin5Cell",
        fontName=PDF_FONT_NAME,
        fontSize=8.5,
        leading=10.2,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=0,
        splitLongWords=True,
    )
    center_style = ParagraphStyle("Rin5CellCenter", parent=cell_style, alignment=TA_CENTER)
    header_style = ParagraphStyle(
        "Rin5Header",
        parent=cell_style,
        fontName=PDF_FONT_BOLD,
        alignment=TA_CENTER,
    )
    signature_style = ParagraphStyle(
        "Rin5Signature",
        fontName=PDF_FONT_NAME,
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        spaceBefore=0,
        spaceAfter=0,
    )
    signature_name_style = ParagraphStyle(
        "Rin5SignatureName",
        parent=signature_style,
        fontName=PDF_FONT_BOLD,
    )
    recap_style = ParagraphStyle(
        "Rin5Recap",
        fontName=PDF_FONT_NAME,
        fontSize=10,
        leading=12,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=0,
    )
    recap_title_style = ParagraphStyle(
        "Rin5RecapTitle",
        parent=recap_style,
        fontName=PDF_FONT_BOLD,
    )

    def rin5_cell(value, style=cell_style):
        text = re.sub(r"\s+", " ", str(value or "-").strip()) or "-"
        return Paragraph(escape(text), style)

    def safe_upload_path(base_dir, filename):
        if not filename:
            return None
        path = (base_dir / filename).resolve()
        try:
            path.relative_to(base_dir.resolve())
        except ValueError:
            return None
        return path if path.is_file() else None

    def pdf_image(path, max_width, max_height):
        if not path:
            return Spacer(1, max_height)
        try:
            with PILImage.open(path) as image:
                width, height = image.size
        except Exception:
            return Spacer(1, max_height)
        if not width or not height:
            return Spacer(1, max_height)
        ratio = min(max_width / width, max_height / height)
        return Image(str(path), width=width * ratio, height=height * ratio)

    def signature_image_block(scan_path, stamp_path=None):
        flowables = []
        if stamp_path:
            flowables.append(pdf_image(stamp_path, 2.35 * cm, 2.35 * cm))
        if scan_path:
            flowables.append(pdf_image(scan_path, 6.25 * cm, 2.75 * cm))
        if not flowables:
            flowables = [Spacer(1, 2.75 * cm)]
        col_widths = ([2.65 * cm] if stamp_path else []) + ([6.9 * cm] if scan_path else [])
        if not col_widths:
            col_widths = [6.9 * cm]
        image_block = Table([flowables], colWidths=col_widths, hAlign="CENTER")
        image_block.setStyle(TableStyle([
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return image_block

    def signature_flowables(signer, prefix_text, date_text=None, include_stamp=False):
        signer_name = str(signer.get("full_name") or "-")
        signer_position = str(signer.get("position_name") or "-")
        signer_rank = str(signer.get("rank_nip") or "-")
        scan_path = safe_upload_path(SIGNATORY_UPLOAD_DIR, signer.get("signature_image")) if use_scan_signature else None
        stamp_path = safe_upload_path(ORGANIZATION_UPLOAD_DIR, organization.get("digital_stamp")) if include_stamp else None
        lines = []
        if date_text:
            lines.append(Paragraph(escape(date_text), signature_style))
        lines.extend([
            Paragraph(prefix_text, signature_style),
            Paragraph(escape(signer_position), signature_style),
            signature_image_block(scan_path, stamp_path),
            Paragraph(f"<u>{escape(signer_name)}</u>", signature_name_style),
            Paragraph(escape(signer_rank), signature_style),
        ])
        return lines

    table_data = [[
        rin5_cell("No", header_style),
        rin5_cell("Nama Satker", header_style),
        rin5_cell("Jenis Produk Intelijen", header_style),
        rin5_cell("Nomor Produk Intelijen", header_style),
        rin5_cell("Tanggal Produk Intelijen", header_style),
        rin5_cell("Subdit/Seksi/Subseksi", header_style),
        rin5_cell("Perihal", header_style),
        rin5_cell("Disposisi Pimpinan", header_style),
        rin5_cell("Keterangan", header_style),
    ]]
    for index, entry in enumerate(rows, 1):
        product_date = entry.get("intelligence_product_date")
        if isinstance(product_date, datetime):
            product_date = product_date.date()
        date_text = product_date.strftime("%d/%m/%Y") if isinstance(product_date, date) else "-"
        field_text = f"{entry.get('field_code') or '-'}\n{ISSUE_CODE_LABELS.get(entry.get('field_code'), '-')}"
        table_data.append([
            rin5_cell(index, center_style),
            rin5_cell(organization_name),
            rin5_cell(entry.get("intelligence_product_type")),
            rin5_cell(entry.get("intelligence_product_number")),
            rin5_cell(date_text, center_style),
            rin5_cell(field_text),
            rin5_cell(entry.get("subject")),
            rin5_cell(entry.get("leader_disposition")),
            rin5_cell(entry.get("remarks"), center_style),
        ])
    if len(table_data) == 1:
        for nihil_row in nihil_table_rows_values(9):
            table_data.append([rin5_cell(value, header_style if value else cell_style) for value in nihil_row])
    table = LongTable(
        table_data,
        colWidths=[0.8 * cm, 2.7 * cm, 3.1 * cm, 3.1 * cm, 2.0 * cm, 4.0 * cm, 6.6 * cm, 6.0 * cm, 2.5 * cm],
        repeatRows=1,
        hAlign="LEFT",
        splitByRow=True,
        splitInRow=False,
    )
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, "#000000"),
        ("BACKGROUND", (0, 0), (-1, 0), "#E8EDF3"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("VALIGN", (0, 0), (-1, 0), "MIDDLE"),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2.5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2.5),
        ("TOPPADDING", (0, 0), (-1, -1), 3.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
    ]))

    report_total = len(rows)
    recap_block = [
        Paragraph("Rekapitulasi", recap_title_style),
        Paragraph("Sisa bulan Lalu :", recap_style),
        Paragraph(f"Masuk Bulan laporan : {report_total}", recap_style),
        Paragraph(f"Jumlah : {report_total}", recap_style),
        Paragraph("Diselesaikan :", recap_style),
        Paragraph("Sisa Bulan Laporan :", recap_style),
    ]
    signature_table = Table(
        [[
            signature_flowables(kajari, "Mengetahui", include_stamp=use_digital_stamp),
            recap_block,
            signature_flowables(kasi, "", date_text=f"Singaraja, {format_indonesian_date(signature_date)}"),
        ]],
        colWidths=[10.2 * cm, 5.0 * cm, 10.2 * cm],
        hAlign="CENTER",
    )
    signature_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))

    def draw_register_header(pdf_canvas, _document):
        pdf_canvas.saveState()
        header_x = _document.leftMargin
        pdf_canvas.setFillColorRGB(0, 0, 0)
        pdf_canvas.setFont(PDF_FONT_BOLD, 12)
        pdf_canvas.drawString(header_x, page_height - 2.05 * cm, "R.IN.5 REGISTER PRODUK INTELIJEN")
        pdf_canvas.drawString(header_x, page_height - 2.48 * cm, f"Bulan : {period_text}")
        pdf_canvas.restoreState()

    story = [
        table,
        Spacer(1, 0.7 * cm),
        KeepTogether(signature_table),
    ]
    document.build(story, onFirstPage=draw_register_header, onLaterPages=draw_register_header)
    buffer.seek(0)
    filename = f"RIN5-REGISTER-PRODUK-INTELIJEN-{date.today().strftime('%Y%m%d')}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype="application/pdf")


@app.route("/register-intelijen/rin5/create", methods=["GET", "POST"])
@login_required
def create_register_intelijen_rin5():
    if request.method == "POST":
        data = register_intelijen_rin5_form_data()
        if (not data["intelligence_product_type"] or not data["intelligence_product_number"]
                or not data["field_code"] or not data["subject"]):
            flash("Jenis produk, nomor produk, Subdit/Seksi/Subseksi, dan perihal wajib diisi.", "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    """INSERT INTO register_intelijen_rin5_entries
                       (intelligence_product_type, intelligence_product_number,
                        intelligence_product_date, field_code, subject,
                        leader_disposition, remarks, created_by)
                       VALUES (%s,%s,%s,%s,%s,%s,%s,%s)""",
                    (
                        data["intelligence_product_type"], data["intelligence_product_number"],
                        data["intelligence_product_date"], data["field_code"], data["subject"],
                        data["leader_disposition"], data["remarks"], session["user_id"],
                    ),
                )
            flash("Data R.IN.5 berhasil disimpan.", "success")
            return redirect(url_for("register_intelijen_rin5"))
    else:
        organization = fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}
        data = {
            "satker_name": organization.get("organization_name") or "Kejaksaan Negeri Buleleng",
            "intelligence_product_type": "",
            "intelligence_product_number": "",
            "intelligence_product_date": date.today().isoformat(),
            "field_code": "",
            "subject": "",
            "leader_disposition": "",
            "remarks": "",
        }
    return render_template(
        "register_intelijen_rin5_form.html",
        active="register_intelijen_rin5",
        form_title="Buat Register Produk Intelijen R.IN.5",
        form_description="Input data produk intelijen untuk register R.IN.5.",
        data=data,
        issue_code_labels=ISSUE_CODE_LABELS,
        entry=None,
    )


@app.route("/register-intelijen/rin5/<int:entry_id>/edit", methods=["GET", "POST"])
@login_required
def edit_register_intelijen_rin5(entry_id):
    entry = fetch_one("SELECT * FROM register_intelijen_rin5_entries WHERE id=%s", (entry_id,))
    if not entry:
        flash("Data R.IN.5 tidak ditemukan.", "error")
        return redirect(url_for("register_intelijen_rin5"))
    if session.get("role") != "admin" and entry["created_by"] != session["user_id"]:
        flash("Anda tidak dapat mengubah data R.IN.5 ini.", "error")
        return redirect(url_for("register_intelijen_rin5"))
    if request.method == "POST":
        data = register_intelijen_rin5_form_data()
        if (not data["intelligence_product_type"] or not data["intelligence_product_number"]
                or not data["field_code"] or not data["subject"]):
            flash("Jenis produk, nomor produk, Subdit/Seksi/Subseksi, dan perihal wajib diisi.", "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    """UPDATE register_intelijen_rin5_entries
                       SET intelligence_product_type=%s, intelligence_product_number=%s,
                           intelligence_product_date=%s, field_code=%s, subject=%s,
                           leader_disposition=%s, remarks=%s
                       WHERE id=%s""",
                    (
                        data["intelligence_product_type"], data["intelligence_product_number"],
                        data["intelligence_product_date"], data["field_code"], data["subject"],
                        data["leader_disposition"], data["remarks"], entry_id,
                    ),
                )
            flash("Data R.IN.5 berhasil diperbarui.", "success")
            return redirect(url_for("register_intelijen_rin5"))
    else:
        organization = fetch_one("SELECT organization_name FROM organization_settings WHERE id=1") or {}
        data = {
            "satker_name": organization.get("organization_name") or "Kejaksaan Negeri Buleleng",
            "intelligence_product_type": entry.get("intelligence_product_type") or "",
            "intelligence_product_number": entry.get("intelligence_product_number") or "",
            "intelligence_product_date": entry["intelligence_product_date"].isoformat() if entry.get("intelligence_product_date") else date.today().isoformat(),
            "field_code": entry.get("field_code") or "",
            "subject": entry.get("subject") or "",
            "leader_disposition": entry.get("leader_disposition") or "",
            "remarks": entry.get("remarks") or "",
        }
    return render_template(
        "register_intelijen_rin5_form.html",
        active="register_intelijen_rin5",
        form_title="Edit Register Produk Intelijen R.IN.5",
        form_description="Perbarui data produk intelijen pada register R.IN.5.",
        data=data,
        issue_code_labels=ISSUE_CODE_LABELS,
        entry=entry,
    )


@app.post("/register-intelijen/rin5/<int:entry_id>/delete")
@login_required
def delete_register_intelijen_rin5(entry_id):
    entry = fetch_one("SELECT id,created_by FROM register_intelijen_rin5_entries WHERE id=%s", (entry_id,))
    if not entry:
        flash("Data R.IN.5 tidak ditemukan.", "error")
    elif session.get("role") != "admin" and entry["created_by"] != session["user_id"]:
        flash("Anda tidak dapat menghapus data R.IN.5 ini.", "error")
    else:
        with get_db().cursor() as cursor:
            cursor.execute("DELETE FROM register_intelijen_rin5_entries WHERE id=%s", (entry_id,))
        flash("Data R.IN.5 berhasil dihapus.", "success")
    return redirect(url_for("register_intelijen_rin5"))


def register_information_sentence(content):
    soup = BeautifulSoup(content or "", "html.parser")
    first_block = soup.find("li") or soup.find(["p", "div"])
    text = (first_block.get_text(" ", strip=True) if first_block else
            soup.get_text(" ", strip=True))
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(
        r"^\s*(?:(?:\(?\d+\)?|[IVXLCDM]+)[.)\-:]|[•●▪\-–—])\s*",
        "",
        text,
        flags=re.IGNORECASE,
    )
    if not text:
        return ""

    protected = re.sub(
        r"\b(?:PT|CV|No|Dr|Ir|Sdr|Tn|Ny|Jl|Kab|Kec|Prov|S\.H|M\.H|S\.E|S\.Sos)\.",
        lambda match: match.group(0).replace(".", "\u2024"),
        text,
        flags=re.IGNORECASE,
    )
    sentence_end = re.search(r"\.(?=\s|$)", protected)
    if sentence_end:
        protected = protected[:sentence_end.end()]
    return protected.replace("\u2024", ".").strip()


@app.post("/reports/<report_type>/<int:report_id>/register-intelijen")
@login_required
def register_report_intelijen(report_type, report_id):
    if report_type == "lapinhar":
        report = accessible_lapinhar(report_id)
        document_label = "Laporan Informasi Harian"
        return_endpoint = "lapinhar"
    elif report_type == "lapinsus":
        report = accessible_lapinsus(report_id)
        document_label = "Laporan Informasi Khusus"
        return_endpoint = "lapinsus"
    else:
        abort(404)

    if report is None:
        return jsonify(message="Laporan tidak ditemukan atau tidak dapat Anda akses."), 403
    if report.get("status") == "draft":
        return jsonify(message=f"Lengkapi dan simpan {document_label.upper()} sebelum diregister."), 409
    if not report.get("report_date") or report.get("issue_code") not in ISSUE_CODES:
        return jsonify(message="Tanggal dan bidang laporan harus tersedia sebelum diregister."), 422

    information_description = register_information_sentence(report.get("facts"))
    if not information_description:
        return jsonify(message="Informasi yang diperoleh masih kosong."), 422

    duplicate = fetch_one(
        """SELECT id FROM register_intelijen_entries
           WHERE source_report_type=%s AND source_report_id=%s""",
        (report_type, report_id),
    )
    if duplicate:
        return jsonify(
            message=f"{document_label.upper()} ini sudah masuk Register Intelijen.",
            already_registered=True,
            register_id=duplicate["id"],
        ), 409

    payload = request.get_json(silent=True) or {}
    received_time = str(payload.get("received_time") or datetime.now().strftime("%H:%M")).strip()
    try:
        received_time = datetime.strptime(received_time, "%H:%M").strftime("%H:%M:%S")
    except ValueError:
        return jsonify(message="Waktu diterima tidak valid."), 422

    information_value = str(payload.get("information_value") or "A1").strip().upper()
    if information_value not in REGISTER_INFORMATION_VALUES:
        return jsonify(message="Nilai data/informasi tidak valid."), 422

    disposition = str(payload.get("disposition") or "TL KE KEJATI").strip()[:2000]
    follow_up = str(
        payload.get("follow_up") or "-SEGERA TL TERUSKAN KE KEJATI -ARSIPKAN"
    ).strip()[:2000]
    remarks = str(payload.get("remarks") or "Arsip").strip()[:255]
    if not disposition or not follow_up or not remarks:
        return jsonify(message="Disposisi, tindak lanjut, dan keterangan wajib diisi."), 422

    try:
        with get_db().cursor() as cursor:
            cursor.execute(
                """INSERT INTO register_intelijen_entries
                   (register_code,source_report_type,source_report_id,report_date,field_code,
                    received_time,source_name,information_value,information_description,
                    notes,disposition,follow_up,remarks,created_by)
                   VALUES ('R.IN.3',%s,%s,%s,%s,%s,'Kejari Buleleng',%s,%s,%s,%s,%s,%s,%s)""",
                (
                    report_type,
                    report_id,
                    report["report_date"],
                    report["issue_code"],
                    received_time,
                    information_value,
                    information_description,
                    document_label,
                    disposition,
                    follow_up,
                    remarks,
                    session["user_id"],
                ),
            )
            register_id = cursor.lastrowid
            cursor.execute(
                """INSERT IGNORE INTO register_intelijen_rin5_entries
                   (source_report_type,source_report_id,intelligence_product_type,
                    intelligence_product_number,intelligence_product_date,field_code,
                    subject,leader_disposition,remarks,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                (
                    report_type,
                    report_id,
                    document_label,
                    report.get("report_number") or "-",
                    report["report_date"],
                    report["issue_code"],
                    report.get("title") or "-",
                    disposition,
                    remarks,
                    session["user_id"],
                ),
            )
    except pymysql.err.IntegrityError:
        existing = fetch_one(
            """SELECT id FROM register_intelijen_entries
               WHERE source_report_type=%s AND source_report_id=%s""",
            (report_type, report_id),
        )
        return jsonify(
            message=f"{document_label.upper()} ini sudah masuk Register Intelijen.",
            already_registered=True,
            register_id=existing["id"] if existing else None,
        ), 409

    return jsonify(
        message=f"{document_label.upper()} berhasil dimasukkan ke R.IN.3 dan R.IN.5.",
        registered=True,
        register_id=register_id,
        redirect_url=url_for(return_endpoint),
    )


@app.get("/register-intelijen/export-pdf")
@login_required
def export_register_intelijen_pdf():
    params = []
    conditions = []
    search = request.args.get("q", "").strip()[:150]
    bidang = request.args.get("bidang", "").strip()
    month = request.args.get("month", "").strip()
    year = request.args.get("year", "").strip()
    use_scan_signature = request.args.get("use_scan_signature") == "1"
    use_digital_stamp = request.args.get("use_digital_stamp") == "1"
    use_acting_kajari = request.args.get("use_acting_kajari") == "1"
    use_acting_kasi = request.args.get("use_acting_kasi") == "1"
    signature_date_raw = request.args.get("signature_date", "").strip()
    try:
        signature_date = datetime.strptime(signature_date_raw, "%Y-%m-%d").date() if signature_date_raw else date.today()
    except ValueError:
        signature_date = date.today()
    if search:
        search_value = f"%{search}%"
        conditions.append(
            """(entries.field_code LIKE %s OR entries.source_name LIKE %s
                OR entries.information_value LIKE %s
                OR entries.information_description LIKE %s OR entries.notes LIKE %s
                OR entries.disposition LIKE %s OR entries.follow_up LIKE %s
                OR entries.remarks LIKE %s OR users.full_name LIKE %s)"""
        )
        params.extend([search_value] * 9)
    if bidang in ISSUE_CODES:
        conditions.append("entries.field_code=%s")
        params.append(bidang)
    if month.isdigit() and 1 <= int(month) <= 12:
        conditions.append("MONTH(entries.report_date)=%s")
        params.append(int(month))
    if year.isdigit():
        conditions.append("YEAR(entries.report_date)=%s")
        params.append(int(year))
    where_clause = f"WHERE {' AND '.join(conditions)}" if conditions else ""
    rows = fetch_all(
        f"""SELECT entries.* FROM register_intelijen_entries entries
            JOIN users ON users.id=entries.created_by
            {where_clause}
            ORDER BY entries.report_date ASC, entries.received_time ASC, entries.id ASC""",
        tuple(params),
    )
    organization = fetch_one("SELECT * FROM organization_settings WHERE id=1") or {}
    kajari = fetch_one("SELECT * FROM signatories WHERE position_code='kajari'") or {}
    kasi = fetch_one("SELECT * FROM signatories WHERE position_code='kasi_intel'") or {}

    def acting_signer(base_signer, prefix, enabled):
        if not enabled:
            return base_signer
        acting_type = str(request.args.get(f"{prefix}_type") or "").strip().lower()
        acting_label = "Plt." if acting_type == "plt" else "Plh." if acting_type == "plh" else ""
        base_position = str(base_signer.get("position_name") or "").strip()
        signer_name = str(request.args.get(f"{prefix}_name") or base_signer.get("full_name") or "").strip()
        signer_position_detail = str(request.args.get(f"{prefix}_position") or "").strip()
        signer_nip = str(request.args.get(f"{prefix}_nip") or "").strip()
        rank_nip = signer_position_detail
        if signer_nip:
            rank_nip = f"{rank_nip} NIP. {signer_nip}".strip()
        acting_position = f"{acting_label} {base_position}".strip() if acting_label else base_position
        result = dict(base_signer)
        result.update(
            full_name=signer_name or "-",
            position_name=acting_position or "-",
            rank_nip=rank_nip or "-",
            signature_image=None,
        )
        return result

    kajari = acting_signer(kajari, "acting_kajari", use_acting_kajari)
    kasi = acting_signer(kasi, "acting_kasi", use_acting_kasi)

    buffer = BytesIO()
    folio_landscape = landscape((8.5 * inch, 13 * inch))
    page_width, page_height = folio_landscape
    document = SimpleDocTemplate(
        buffer,
        pagesize=folio_landscape,
        leftMargin=1 * cm,
        rightMargin=1 * cm,
        topMargin=3 * cm,
        bottomMargin=1.1 * cm,
        allowSplitting=True,
    )
    satker = organization.get("organization_name") or "Kejaksaan Negeri Buleleng"
    report_dates = []
    for row in rows:
        report_date = row.get("report_date")
        if isinstance(report_date, datetime):
            report_date = report_date.date()
        if isinstance(report_date, date):
            report_dates.append(report_date)
    if report_dates:
        first_date, last_date = min(report_dates), max(report_dates)
        if first_date.year == last_date.year and first_date.month == last_date.month:
            period_text = f"{MONTH_NAMES_ID[first_date.month]} {first_date.year}"
        elif first_date.year == last_date.year:
            period_text = (
                f"{MONTH_NAMES_ID[first_date.month]} - "
                f"{MONTH_NAMES_ID[last_date.month]} {first_date.year}"
            )
        else:
            period_text = (
                f"{MONTH_NAMES_ID[first_date.month]} {first_date.year} - "
                f"{MONTH_NAMES_ID[last_date.month]} {last_date.year}"
            )
    elif month.isdigit() and year.isdigit() and 1 <= int(month) <= 12:
        period_text = f"{MONTH_NAMES_ID[int(month)]} {year}"
    else:
        period_text = year if year.isdigit() else "-"

    body_style = ParagraphStyle(
        "RegisterCell",
        fontName=PDF_FONT_NAME,
        fontSize=7.5,
        leading=9,
        spaceBefore=0,
        spaceAfter=0,
        alignment=TA_LEFT,
        splitLongWords=True,
    )
    center_style = ParagraphStyle(
        "RegisterCellCenter",
        parent=body_style,
        alignment=TA_CENTER,
    )
    header_style = ParagraphStyle(
        "RegisterHeader",
        parent=center_style,
        fontName=PDF_FONT_BOLD,
        fontSize=7.5,
        leading=8.8,
    )
    signature_style = ParagraphStyle(
        "RegisterSignature",
        fontName=PDF_FONT_NAME,
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        spaceBefore=0,
        spaceAfter=0,
    )
    title_style = ParagraphStyle(
        "RegisterTitle",
        fontName=PDF_FONT_BOLD,
        fontSize=12,
        leading=14,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=0,
    )
    period_style = ParagraphStyle(
        "RegisterPeriod",
        fontName=PDF_FONT_BOLD,
        fontSize=12,
        leading=14,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=0,
    )
    signature_name_style = ParagraphStyle(
        "RegisterSignatureName",
        parent=signature_style,
        fontName=PDF_FONT_BOLD,
    )
    recap_style = ParagraphStyle(
        "RegisterRecap",
        fontName=PDF_FONT_NAME,
        fontSize=10,
        leading=12,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=0,
    )
    recap_title_style = ParagraphStyle(
        "RegisterRecapTitle",
        parent=recap_style,
        fontName=PDF_FONT_BOLD,
    )

    def cell_paragraph(value, style=body_style):
        text = re.sub(r"\s+", " ", str(value or "-").strip()) or "-"
        return Paragraph(escape(text), style)

    def safe_upload_path(base_dir, filename):
        if not filename:
            return None
        path = (base_dir / filename).resolve()
        try:
            path.relative_to(base_dir.resolve())
        except ValueError:
            return None
        return path if path.is_file() else None

    def pdf_image(path, max_width, max_height):
        if not path:
            return Spacer(1, max_height)
        try:
            with PILImage.open(path) as image:
                width, height = image.size
        except Exception:
            return Spacer(1, max_height)
        if not width or not height:
            return Spacer(1, max_height)
        ratio = min(max_width / width, max_height / height)
        return Image(str(path), width=width * ratio, height=height * ratio)

    def signature_image_block(scan_path, stamp_path=None):
        flowables = []
        if stamp_path:
            flowables.append(pdf_image(stamp_path, 2.35 * cm, 2.35 * cm))
        if scan_path:
            flowables.append(pdf_image(scan_path, 6.25 * cm, 2.75 * cm))
        if not flowables:
            flowables = [Spacer(1, 2.75 * cm)]
        col_widths = ([2.65 * cm] if stamp_path else []) + ([6.9 * cm] if scan_path else [])
        if not col_widths:
            col_widths = [6.9 * cm]
        image_block = Table(
            [flowables],
            colWidths=col_widths,
            hAlign="CENTER",
        )
        image_block.setStyle(
            TableStyle(
                [
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ]
            )
        )
        return image_block

    def signature_flowables(signer, prefix_text, date_text=None, include_stamp=False):
        signer_name = str(signer.get("full_name") or "-")
        signer_position = str(signer.get("position_name") or "-")
        signer_rank = str(signer.get("rank_nip") or "-")
        scan_path = safe_upload_path(SIGNATORY_UPLOAD_DIR, signer.get("signature_image")) if use_scan_signature else None
        stamp_path = safe_upload_path(ORGANIZATION_UPLOAD_DIR, organization.get("digital_stamp")) if include_stamp else None
        image_block = signature_image_block(scan_path, stamp_path)
        lines = []
        if date_text:
            lines.append(Paragraph(escape(date_text), signature_style))
        lines.extend([
            Paragraph(prefix_text, signature_style),
            Paragraph(escape(signer_position), signature_style),
            image_block,
            Paragraph(f"<u>{escape(signer_name)}</u>", signature_name_style),
            Paragraph(escape(signer_rank), signature_style),
        ])
        return lines

    headers = [
        "No",
        "Nama Satker",
        "Tanggal<br/>Laporan",
        "Bidang",
        "Waktu<br/>Diterima",
        "Sumber/<br/>Bapul",
        "Nilai Data/<br/>Informasi",
        "Uraian Informasi",
        "Catatan",
        "Disposisi/<br/>Tindakan",
        "Tindak Lanjut",
        "Keterangan",
    ]
    table_data = [[Paragraph(text, header_style) for text in headers]]
    for index, row in enumerate(rows, 1):
        report_date = row.get("report_date")
        if isinstance(report_date, (date, datetime)):
            report_date = report_date.strftime("%d/%m/%Y")
        table_data.append(
            [
                cell_paragraph(index, center_style),
                cell_paragraph(satker),
                cell_paragraph(report_date, center_style),
                cell_paragraph(issue_code_label(row.get("field_code"))),
                cell_paragraph(
                    f"{format_time_value(row.get('received_time'))} Wita"
                    if row.get("received_time")
                    else "-",
                    center_style,
                ),
                cell_paragraph(row.get("source_name")),
                cell_paragraph(row.get("information_value"), center_style),
                cell_paragraph(row.get("information_description")),
                cell_paragraph(row.get("notes")),
                cell_paragraph(row.get("disposition")),
                cell_paragraph(row.get("follow_up")),
                cell_paragraph(row.get("remarks"), center_style),
            ]
        )
    if len(table_data) == 1:
        for nihil_row in nihil_table_rows_values(len(headers)):
            table_data.append([cell_paragraph(value, header_style if value else cell_style) for value in nihil_row])

    column_widths = [
        0.8 * cm,
        2.7 * cm,
        1.9 * cm,
        3.55 * cm,
        1.55 * cm,
        2.15 * cm,
        1.35 * cm,
        6.7 * cm,
        2.35 * cm,
        2.85 * cm,
        2.85 * cm,
        1.55 * cm,
    ]
    register_table = LongTable(
        table_data,
        colWidths=column_widths,
        repeatRows=1,
        hAlign="LEFT",
        splitByRow=True,
        splitInRow=False,
    )
    register_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, "#000000"),
                ("BACKGROUND", (0, 0), (-1, 0), "#E8EDF3"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("VALIGN", (0, 0), (-1, 0), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2.5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2.5),
                ("TOPPADDING", (0, 0), (-1, -1), 3.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
            ]
        )
    )

    report_total = len(rows)
    recap_block = [
        Paragraph("Rekapitulasi", recap_title_style),
        Paragraph("Sisa bulan Lalu :", recap_style),
        Paragraph(f"Masuk Bulan laporan : {report_total}", recap_style),
        Paragraph(f"Jumlah : {report_total}", recap_style),
        Paragraph("Diselesaikan :", recap_style),
        Paragraph("Sisa Bulan Laporan :", recap_style),
    ]
    signature_table = Table(
        [
            [
                signature_flowables(
                    kajari,
                    "Mengetahui",
                    include_stamp=use_digital_stamp,
                ),
                recap_block,
                signature_flowables(
                    kasi,
                    "",
                    date_text=f"Singaraja, {format_indonesian_date(signature_date)}",
                ),
            ]
        ],
        colWidths=[10.2 * cm, 5.0 * cm, 10.2 * cm],
        hAlign="CENTER",
    )
    signature_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )

    def draw_register_header(pdf_canvas, _document):
        pdf_canvas.saveState()
        header_x = _document.leftMargin
        pdf_canvas.setFillColorRGB(0, 0, 0)
        pdf_canvas.setFont(PDF_FONT_BOLD, 12)
        pdf_canvas.drawString(header_x, page_height - 2.05 * cm, "R.IN.3 REGISTER KERJA INTELIJEN")
        pdf_canvas.drawString(header_x, page_height - 2.48 * cm, f"Bulan : {period_text}")
        pdf_canvas.restoreState()

    story = [
        register_table,
        Spacer(1, 0.7 * cm),
        KeepTogether(signature_table),
    ]
    document.build(story, onFirstPage=draw_register_header, onLaterPages=draw_register_header)
    buffer.seek(0)
    filename = f"RIN3-REGISTER-KERJA-INTELIJEN-{date.today().strftime('%Y%m%d')}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype="application/pdf")


@app.route("/lapinhar")
@login_required
def lapinhar():
    conditions = ["reports.report_type = 'lapinhar'"]
    params = []
    current_year = date.today().year
    chart_month = request.args.get("chart_month", "all").strip() or "all"
    base_where = "reports.report_type = 'lapinhar' AND YEAR(reports.report_date) = %s"
    summary = fetch_one(
        f"""SELECT COUNT(*) AS total,
            SUM(reports.lapinsus_status='belum') AS lapinsus_belum,
            SUM(reports.inteliz_status='belum') AS inteliz_belum,
            SUM(reports.status<>'draft' AND NOT EXISTS(
                SELECT 1 FROM register_intelijen_entries register_entries
                WHERE register_entries.source_report_type='lapinhar'
                AND register_entries.source_report_id=reports.id
            )) AS register_belum
            FROM reports WHERE {base_where}""", (current_year,)
    )
    report_summary = {
        "total": int(summary["total"] or 0),
        "lapinsus_belum": int(summary["lapinsus_belum"] or 0),
        "inteliz_belum": int(summary["inteliz_belum"] or 0),
        "register_belum": int(summary["register_belum"] or 0),
    }
    issue_chart_conditions = [
        "report_type='lapinhar'",
        "YEAR(report_date)=%s",
        "issue_code IS NOT NULL",
        "issue_code <> ''",
    ]
    issue_chart_params = [current_year]
    if chart_month.isdigit() and 1 <= int(chart_month) <= 12:
        issue_chart_conditions.append("MONTH(report_date)=%s")
        issue_chart_params.append(int(chart_month))
        chart_month_name = MONTH_NAMES_ID[int(chart_month)]
    else:
        chart_month = "all"
        chart_month_name = "Semua bulan"
    issue_chart_source = fetch_all(
        f"""SELECT issue_code, COUNT(*) AS total
           FROM reports
           WHERE {' AND '.join(issue_chart_conditions)}
           GROUP BY issue_code""",
        tuple(issue_chart_params),
    )
    issue_totals = {
        str(row["issue_code"] or "").strip(): int(row["total"] or 0)
        for row in issue_chart_source
    }
    issue_chart_groups = [
        {"group": "Ds", "codes": ["Ds.1", "Ds.2", "Ds.3"]},
        {"group": "Dip", "codes": ["Dip.1", "Dip.2", "Dip.3", "Dip.4"]},
        {"group": "Dsb", "codes": ["Dsb.1", "Dsb.2", "Dsb.3", "Dsb.4"]},
        {"group": "Dek", "codes": ["Dek.1", "Dek.2", "Dek.3", "Dek.4"]},
        {"group": "Dpp", "codes": ["Dpp.1", "Dpp.2", "Dpp.3", "Dpp.4"]},
        {"group": "Dti", "codes": ["Dti.1", "Dti.2", "Dti.3"]},
    ]
    max_issue_chart_total = max(
        [issue_totals.get(code, 0) for group in issue_chart_groups for code in group["codes"]] or [0]
    )
    issue_chart_axis_max = max(1, max_issue_chart_total)
    issue_chart_axis_ticks = sorted({0, issue_chart_axis_max // 2, issue_chart_axis_max})
    issue_chart_colors = {
        "Ds": "#14b8a6",
        "Dip": "#3b82f6",
        "Dsb": "#8b5cf6",
        "Dek": "#f59e0b",
        "Dpp": "#f43f5e",
        "Dti": "#22c55e",
    }
    issue_chart_dataset = {"labels": [], "totals": [], "colors": [], "groups": []}
    for group in issue_chart_groups:
        group["bars"] = [
            {
                "code": code,
                "label": ISSUE_CODE_LABELS.get(code, code),
                "total": issue_totals.get(code, 0),
            }
            for code in group["codes"]
        ]
        start_index = len(issue_chart_dataset["labels"])
        for bar in group["bars"]:
            issue_chart_dataset["labels"].append(bar["code"])
            issue_chart_dataset["totals"].append(bar["total"])
            issue_chart_dataset["colors"].append(issue_chart_colors[group["group"]])
        end_index = len(issue_chart_dataset["labels"]) - 1
        issue_chart_dataset["groups"].append(
            {"name": group["group"], "start": start_index, "end": end_index}
        )

    search = request.args.get("q", "").strip()[:150]
    if search:
        search_columns = ["reports.report_number LIKE %s", "reports.title LIKE %s"]
        term = f"%{search}%"
        params.extend([term, term])
        search_columns.append("users.full_name LIKE %s")
        params.append(term)
        conditions.append(f"({' OR '.join(search_columns)})")
    where_clause = " AND ".join(conditions)
    try:
        page = max(1, int(request.args.get("page", "1")))
    except ValueError:
        page = 1
    per_page = 10
    filtered_total = fetch_one(
        f"""SELECT COUNT(*) AS total FROM reports
            JOIN users ON users.id=reports.created_by WHERE {where_clause}""", tuple(params)
    )["total"]
    total_pages = max(1, (int(filtered_total) + per_page - 1) // per_page)
    page = min(page, total_pages)
    offset = (page - 1) * per_page
    reports = fetch_all(
        "SELECT reports.id, reports.report_number, reports.title, reports.status, reports.created_by, reports.inteliz_status, "
        "reports.lapinsus_status, reports.created_at, users.full_name AS creator_full_name, "
        "EXISTS(SELECT 1 FROM register_intelijen_entries register_entries "
        "WHERE register_entries.source_report_type='lapinhar' "
        "AND register_entries.source_report_id=reports.id) AS is_registered, "
        "(SELECT COUNT(*) FROM report_attachments WHERE report_attachments.report_id=reports.id) AS attachment_count "
        "FROM reports JOIN users ON users.id = reports.created_by "
        f"WHERE {where_clause} ORDER BY reports.created_at DESC LIMIT %s OFFSET %s",
        tuple(params + [per_page, offset]),
    )
    draft_choice = None
    try:
        draft_choice_id = int(request.args.get("choose_draft", "0"))
    except ValueError:
        draft_choice_id = 0
    if draft_choice_id:
        draft_choice = fetch_one(
            """SELECT id,report_number,title,created_at FROM lapinhar_reports
               WHERE id=%s AND created_by=%s AND status='draft'""",
            (draft_choice_id, session["user_id"]),
        )
    return render_template("reports.html", report_type="LAPINHAR", reports=reports,
                            report_summary=report_summary, active="lapinhar",
                            issue_chart_groups=issue_chart_groups,
                            max_issue_chart_total=max_issue_chart_total,
                            issue_chart_axis_max=issue_chart_axis_max,
                            issue_chart_axis_ticks=issue_chart_axis_ticks,
                            issue_chart_dataset=issue_chart_dataset,
                            current_year=current_year,
                            chart_month=chart_month,
                            chart_month_name=chart_month_name,
                            month_names=MONTH_NAMES_ID,
                            search=search, page=page, total_pages=total_pages,
                            filtered_total=int(filtered_total), per_page=per_page,
                            draft_choice=draft_choice,
                            auto_connect=request.args.get("connect_inteliz") == "1",
                            next_url=url_for("lapinhar"))


def normalize_whatsapp_number(value):
    """Convert common Indonesian phone formats into the international wa.me format."""
    number = re.sub(r"\D", "", value or "")
    if number.startswith("00"):
        number = number[2:]
    elif number.startswith("0"):
        number = "62" + number[1:]
    elif number.startswith("8"):
        number = "62" + number
    return number


def whatsapp_plain_text(content):
    """Turn editor HTML into readable WhatsApp text while preserving list numbering."""
    if not content:
        return "-"
    if "<" not in content:
        return content.strip() or "-"
    soup = BeautifulSoup(content, "html.parser")
    lines = []

    def clean_text(value):
        return re.sub(r"\s+", " ", value or "").strip()

    def render_list(list_node, level=0):
        for index, item in enumerate(list_node.find_all("li", recursive=False), 1):
            nested = item.find(["ol", "ul"], recursive=False)
            text_parts = [str(child) for child in item.contents if child is not nested]
            item_text = clean_text(BeautifulSoup("".join(text_parts), "html.parser").get_text(" ", strip=True))
            marker = f"{index}." if list_node.name == "ol" else "•"
            if item_text:
                lines.append(f"{'   ' * level}{marker} {item_text}")
            if nested:
                render_list(nested, level + 1)

    for node in soup.find_all(["p", "ol", "ul"], recursive=False):
        if node.name in {"ol", "ul"}:
            render_list(node)
        else:
            text = clean_text(node.get_text(" ", strip=True))
            if text:
                lines.append(text)
    if not lines:
        fallback = [clean_text(part) for part in soup.get_text("\n").splitlines()]
        lines.extend(part for part in fallback if part)
    return "\n".join(lines).strip() or "-"


@app.get("/lapinhar/<int:report_id>/whatsapp")
@login_required
def send_lapinhar_whatsapp(report_id):
    report = accessible_lapinhar(report_id)
    if report is None:
        flash("Laporan tidak ditemukan atau tidak dapat Anda akses.", "error")
        return redirect(url_for("lapinhar"))
    if report.get("status") == "draft":
        flash("Lengkapi dan simpan LAPINHAR sebelum mengirim WhatsApp.", "warning")
        return redirect(url_for("lapinhar"))
    setting = fetch_one(
        "SELECT contact_name, phone_number FROM whatsapp_user_settings WHERE user_id=%s",
        (session["user_id"],),
    )
    if not setting:
        flash("Simpan nama dan nomor tujuan WhatsApp terlebih dahulu.", "error")
        return redirect(url_for("integration_settings", _anchor="whatsapp"))
    phone = normalize_whatsapp_number(setting["phone_number"])
    if not re.fullmatch(r"\d{9,15}", phone):
        flash("Nomor WhatsApp pada konfigurasi tidak valid.", "error")
        return redirect(url_for("integration_settings", _anchor="whatsapp"))
    organization_name = (report.get("organization") or "KEJAKSAAN NEGERI BULELENG").title()
    sender_name = f"Kepala Seksi Intelijen {organization_name}"
    recipient_name = setting["contact_name"].rstrip(". ") + "."
    message = "\n".join([
        "*Laporan Harian*", "", "Kepada Yth.", recipient_name, "",
        f"Dari : {sender_name}", "", f"Perihal : {report['title']}", "",
        "*I. INFORMASI YANG DIPEROLEH*",
        whatsapp_plain_text(report["facts"]), "", "*II. SUMBER INFORMASI*",
        whatsapp_plain_text(report["source_name"]), "", "*III. TREN PERKEMBANGAN / PERKIRAAN*",
        whatsapp_plain_text(report["analysis"]), "", "*IV. PENDAPAT / SARAN*",
        whatsapp_plain_text(report["recommendation"]), "", "Terimakasih🙏",
    ])
    attachment_rows = fetch_all(
        """SELECT attachment_group, sort_order, image_path
           FROM report_attachments WHERE report_id=%s
           ORDER BY attachment_group, sort_order, id""",
        (report_id,),
    ) if request.args.get("attachments") == "1" else []
    downloads = ([url_for("download_all_lapinhar_attachments", report_id=report_id)]
                 if attachment_rows else [])
    whatsapp_url = f"https://web.whatsapp.com/send?phone={phone}&text={quote(message)}"
    return render_template("whatsapp_launch.html", whatsapp_url=whatsapp_url,
                           downloads=downloads, attachment_count=len(attachment_rows),
                           uses_archive=len(attachment_rows) > 1)


@app.get("/lapinsus/<int:report_id>/whatsapp")
@login_required
def send_lapinsus_whatsapp(report_id):
    report = accessible_lapinsus(report_id)
    if report is None:
        flash("LAPINSUS tidak ditemukan atau tidak dapat Anda akses.", "error")
        return redirect(url_for("lapinsus"))
    if report.get("status") == "draft":
        flash("Lengkapi dan simpan LAPINSUS sebelum mengirim WhatsApp.", "warning")
        return redirect(url_for("lapinsus"))
    setting = fetch_one(
        "SELECT contact_name, phone_number FROM whatsapp_user_settings WHERE user_id=%s",
        (session["user_id"],),
    )
    if not setting:
        flash("Simpan nama dan nomor tujuan WhatsApp terlebih dahulu.", "error")
        return redirect(url_for("integration_settings", _anchor="whatsapp"))
    phone = normalize_whatsapp_number(setting["phone_number"])
    if not re.fullmatch(r"\d{9,15}", phone):
        flash("Nomor WhatsApp pada konfigurasi tidak valid.", "error")
        return redirect(url_for("integration_settings", _anchor="whatsapp"))
    organization_name = (report.get("organization") or "KEJAKSAAN NEGERI BULELENG").title()
    sender_name = f"Kepala {organization_name}"
    acting_type = str(report.get("acting_officer_type") or "").strip().lower()
    if acting_type in {"plh", "plt"}:
        sender_name = f"{acting_type.title()}. Kepala {organization_name}"
    recipient_name = setting["contact_name"].rstrip(". ") + "."
    message = "\n".join([
        "*Laporan Khusus*", "", "Kepada Yth.", recipient_name, "",
        f"Dari : {sender_name}", "", f"Perihal : {report['title']}", "",
        "*I. INFORMASI YANG DIPEROLEH*",
        whatsapp_plain_text(report["facts"]), "", "*II. SUMBER INFORMASI*",
        whatsapp_plain_text(report["source_name"]), "", "*III. TREN PERKEMBANGAN / PERKIRAAN*",
        whatsapp_plain_text(report["analysis"]), "", "*IV. PENDAPAT / SARAN*",
        whatsapp_plain_text(report["recommendation"]), "", "Terimakasih🙏",
    ])
    attachment_rows = fetch_all(
        """SELECT attachment_group, sort_order, image_path
           FROM report_attachments WHERE report_id=%s
           ORDER BY attachment_group, sort_order, id""",
        (report_id,),
    ) if request.args.get("attachments") == "1" else []
    downloads = ([url_for("download_all_lapinsus_attachments", report_id=report_id)]
                 if attachment_rows else [])
    whatsapp_url = f"https://web.whatsapp.com/send?phone={phone}&text={quote(message)}"
    return render_template("whatsapp_launch.html", whatsapp_url=whatsapp_url,
                           downloads=downloads, attachment_count=len(attachment_rows),
                           uses_archive=len(attachment_rows) > 1)


@app.route("/lapinsus")
@login_required
def lapinsus():
    conditions = ["reports.report_type = 'lapinsus'"]
    params = []
    current_year = date.today().year
    chart_month = request.args.get("chart_month", "all").strip() or "all"
    summary_where = "reports.report_type = 'lapinsus' AND YEAR(reports.report_date) = %s"
    summary = fetch_one(
        f"""SELECT COUNT(*) total,
            SUM(inteliz_status='belum') inteliz_belum,
            SUM(sipede_status='belum') sipede_belum,
            SUM(reports.status<>'draft' AND NOT EXISTS(
                SELECT 1 FROM register_intelijen_entries register_entries
                WHERE register_entries.source_report_type='lapinsus'
                AND register_entries.source_report_id=reports.id
            )) register_belum
            FROM lapinsus_reports reports WHERE {summary_where}""", (current_year,),
    )
    report_summary = {
        "total": int(summary["total"] or 0),
        "inteliz_belum": int(summary["inteliz_belum"] or 0),
        "sipede_belum": int(summary["sipede_belum"] or 0),
        "register_belum": int(summary["register_belum"] or 0),
    }
    issue_chart_conditions = [
        "report_type='lapinsus'",
        "YEAR(report_date)=%s",
        "issue_code IS NOT NULL",
        "issue_code <> ''",
    ]
    issue_chart_params = [current_year]
    if chart_month.isdigit() and 1 <= int(chart_month) <= 12:
        issue_chart_conditions.append("MONTH(report_date)=%s")
        issue_chart_params.append(int(chart_month))
        chart_month_name = MONTH_NAMES_ID[int(chart_month)]
    else:
        chart_month = "all"
        chart_month_name = "Semua bulan"
    issue_chart_source = fetch_all(
        f"""SELECT issue_code, COUNT(*) AS total
            FROM reports
            WHERE {' AND '.join(issue_chart_conditions)}
            GROUP BY issue_code""",
        tuple(issue_chart_params),
    )
    issue_totals = {
        str(row["issue_code"] or "").strip(): int(row["total"] or 0)
        for row in issue_chart_source
    }
    issue_chart_groups = [
        {"group": "Ds", "codes": ["Ds.1", "Ds.2", "Ds.3"]},
        {"group": "Dip", "codes": ["Dip.1", "Dip.2", "Dip.3", "Dip.4"]},
        {"group": "Dsb", "codes": ["Dsb.1", "Dsb.2", "Dsb.3", "Dsb.4"]},
        {"group": "Dek", "codes": ["Dek.1", "Dek.2", "Dek.3", "Dek.4"]},
        {"group": "Dpp", "codes": ["Dpp.1", "Dpp.2", "Dpp.3", "Dpp.4"]},
        {"group": "Dti", "codes": ["Dti.1", "Dti.2", "Dti.3"]},
    ]
    max_issue_chart_total = max(
        [issue_totals.get(code, 0) for group in issue_chart_groups for code in group["codes"]] or [0]
    )
    issue_chart_axis_max = max(1, max_issue_chart_total)
    issue_chart_axis_ticks = sorted({0, issue_chart_axis_max // 2, issue_chart_axis_max})
    issue_chart_colors = {
        "Ds": "#14b8a6",
        "Dip": "#3b82f6",
        "Dsb": "#8b5cf6",
        "Dek": "#f59e0b",
        "Dpp": "#f43f5e",
        "Dti": "#22c55e",
    }
    issue_chart_dataset = {"labels": [], "totals": [], "colors": [], "groups": []}
    for group in issue_chart_groups:
        start_index = len(issue_chart_dataset["labels"])
        for code in group["codes"]:
            issue_chart_dataset["labels"].append(code)
            issue_chart_dataset["totals"].append(issue_totals.get(code, 0))
            issue_chart_dataset["colors"].append(issue_chart_colors[group["group"]])
        issue_chart_dataset["groups"].append(
            {"name": group["group"], "start": start_index, "end": len(issue_chart_dataset["labels"]) - 1}
        )
    search = request.args.get("q", "").strip()[:150]
    if search:
        columns = ["reports.report_number LIKE %s", "reports.title LIKE %s"]
        term = f"%{search}%"
        params.extend([term, term])
        columns.append("users.full_name LIKE %s")
        params.append(term)
        conditions.append(f"({' OR '.join(columns)})")
    where_clause = " AND ".join(conditions)
    try:
        page = max(1, int(request.args.get("page", "1")))
    except ValueError:
        page = 1
    per_page = 10
    filtered_total = int(fetch_one(
        f"SELECT COUNT(*) total FROM reports JOIN users ON users.id=reports.created_by WHERE {where_clause}",
        tuple(params),
    )["total"] or 0)
    total_pages = max(1, (filtered_total + per_page - 1) // per_page)
    page = min(page, total_pages)
    reports = fetch_all(
        "SELECT reports.id, reports.report_number, reports.title, reports.created_by, reports.status, reports.inteliz_status, "
        "reports.sipede_status, reports.created_at, users.full_name creator_full_name, "
        "EXISTS(SELECT 1 FROM register_intelijen_entries register_entries "
        "WHERE register_entries.source_report_type='lapinsus' "
        "AND register_entries.source_report_id=reports.id) AS is_registered, "
        "(SELECT COUNT(*) FROM report_attachments WHERE report_attachments.report_id=reports.id) AS attachment_count "
        "FROM reports JOIN users ON users.id=reports.created_by "
        f"WHERE {where_clause} ORDER BY reports.created_at DESC LIMIT %s OFFSET %s",
        tuple(params + [per_page, (page - 1) * per_page]),
    )
    draft_choice = None
    try:
        draft_choice_id = int(request.args.get("choose_draft", "0"))
    except ValueError:
        draft_choice_id = 0
    if draft_choice_id:
        draft_choice = fetch_one(
            """SELECT id,report_number,title,created_at FROM lapinsus_reports
               WHERE id=%s AND created_by=%s AND status='draft'""",
            (draft_choice_id, session["user_id"]),
        )
    return render_template("reports.html", report_type="LAPINSUS", reports=reports, active="lapinsus",
                            report_summary=report_summary,
                            current_year=current_year,
                            chart_month=chart_month,
                            chart_month_name=chart_month_name,
                            month_names=MONTH_NAMES_ID,
                            issue_chart_dataset=issue_chart_dataset,
                            issue_chart_axis_max=issue_chart_axis_max,
                            issue_chart_axis_ticks=issue_chart_axis_ticks,
                            search=search, page=page, total_pages=total_pages,
                            filtered_total=filtered_total, per_page=per_page,
                            draft_choice=draft_choice,
                            auto_connect=request.args.get("connect_inteliz") == "1",
                            next_url=url_for("lapinsus"))


def accessible_lapinsus(report_id):
    report = fetch_one("SELECT * FROM reports WHERE id=%s AND report_type='lapinsus'", (report_id,))
    if report and (session.get("role") == "admin" or report["created_by"] == session.get("user_id")):
        return report
    return None


def lapinsus_sequence_number(report_number):
    match = re.match(r"^R-LIK-(\d+)[A-Z]*/", report_number or "", re.IGNORECASE)
    return int(match.group(1)) if match else 0


def lapinsus_sequence_label(report_number):
    match = re.match(r"^R-LIK-(\d+[A-Z]*)/", report_number or "", re.IGNORECASE)
    return match.group(1).upper() if match else ""


def compose_lapinsus_number(sequence_number, institution_code, issue_code, report_date):
    try:
        selected_date = date.fromisoformat(report_date)
    except (TypeError, ValueError):
        return ""
    if not sequence_number or issue_code not in ISSUE_CODES:
        return ""
    return f"R-LIK-{sequence_number}/{institution_code}/{issue_code}/{selected_date.month:02d}/{selected_date.year}"


def reserve_lapinsus_number(document_year=None):
    document_year = int(document_year or date.today().year)
    connection = get_db()
    token = uuid.uuid4().hex
    try:
        connection.begin()
        with connection.cursor() as cursor:
            cursor.execute(
                """SELECT id,sequence_number FROM released_document_numbers
                   WHERE document_type='lapinsus' AND document_year=%s
                   ORDER BY sequence_number LIMIT 1 FOR UPDATE""", (document_year,))
            released = cursor.fetchone()
            if released:
                sequence_number = released["sequence_number"]
                cursor.execute("DELETE FROM released_document_numbers WHERE id=%s", (released["id"],))
            else:
                cursor.execute("SELECT next_number FROM document_counters WHERE document_type='lapinsus' AND document_year=%s FOR UPDATE", (document_year,))
                counter = cursor.fetchone()
                if counter is None:
                    sequence_number = 1
                    cursor.execute("INSERT INTO document_counters (document_type,document_year,next_number) VALUES ('lapinsus',%s,2)", (document_year,))
                else:
                    sequence_number = counter["next_number"]
                    cursor.execute("UPDATE document_counters SET next_number=next_number+1 WHERE document_type='lapinsus' AND document_year=%s", (document_year,))
            cursor.execute(
                "INSERT INTO document_number_reservations "
                "(reservation_token,document_type,document_year,sequence_number,created_by) VALUES (%s,'lapinsus',%s,%s,%s)",
                (token, document_year, sequence_number, session["user_id"]),
            )
        connection.commit()
        return {"reservation_token": token, "sequence_number": sequence_number, "document_year": document_year}
    except Exception:
        connection.rollback()
        raise


def lapinsus_number_available(report_number, reservation_token="", report_id=None):
    duplicate = fetch_one("SELECT id FROM reports WHERE report_number=%s", (report_number,))
    if duplicate and duplicate["id"] != report_id:
        return False, "Nomor surat sudah digunakan laporan lain. Silakan reload nomor surat."
    if report_id:
        return True, "Nomor surat tersedia."
    reservation = fetch_one(
        "SELECT sequence_number,document_year FROM document_number_reservations WHERE reservation_token=%s "
        "AND document_type='lapinsus' AND created_by=%s AND status='reserved'",
        (reservation_token, session["user_id"]),
    )
    if (not reservation or reservation["sequence_number"] != lapinsus_sequence_number(report_number)
            or reservation["document_year"] != report_number_year(report_number)):
        return False, "Reservasi nomor surat tidak berlaku. Silakan reload nomor surat."
    return True, "Nomor surat tersedia."


@app.post("/lapinsus/check-number")
@login_required
def check_lapinsus_number():
    payload = request.get_json(silent=True) or {}
    try:
        report_id = int(payload["report_id"]) if payload.get("report_id") else None
    except (TypeError, ValueError):
        report_id = None
    if report_id and accessible_lapinsus(report_id) is None:
        abort(404)
    available, message = lapinsus_number_available(
        str(payload.get("report_number", "")).strip(),
        str(payload.get("reservation_token", "")).strip(), report_id,
    )
    return jsonify(available=available, message=message), 200 if available else 409


@app.post("/lapinsus/reload-number")
@login_required
def reload_lapinsus_number():
    payload = request.get_json(silent=True) or {}
    reservation = reserve_lapinsus_number(payload.get("document_year"))
    return jsonify(**reservation)


def lapinhar_form_data():
    fields = ("report_number", "subject", "report_date", "recipient", "sender", "classification", "category_id",
              "attachment", "organization", "information", "sources", "trends", "suggestions",
              "city", "creator_position", "creator_name", "creator_rank_nip",
              "auth_position", "auth_name", "auth_rank_nip", "information_spacing",
              "sources_spacing", "trends_spacing", "suggestions_spacing")
    data = {field: request.form.get(field, "").strip() for field in fields}
    organization = fetch_one("SELECT * FROM organization_settings WHERE id = 1") or {
        "organization_name": "KEJAKSAAN NEGERI BULELENG", "address": "", "phone": "", "website": ""
    }
    data["organization"] = organization["organization_name"]
    data["institution_code"] = organization.get("institution_code") or "N.1.11"
    data["recipient"] = f"YTH. KEPALA {organization['organization_name']}"
    data["sender"] = "KASI INTELIJEN"
    data["classification"] = "RAHASIA"
    issue_code = request.form.get("issue_code", "").strip()
    data["issue_code"] = issue_code if issue_code in ISSUE_CODES else ""
    try:
        category_id = int(data["category_id"])
        data["category_id"] = category_id if 1 <= category_id <= 74 else None
    except (TypeError, ValueError):
        data["category_id"] = None
    for field in ("information_spacing", "sources_spacing", "trends_spacing", "suggestions_spacing"):
        data[field] = "1.5"
    data["city"] = os.getenv("SIGNING_CITY", "Singaraja")
    kasi = fetch_one("SELECT * FROM signatories WHERE position_code = 'kasi_intel'")
    subsection_code = request.form.get("subsection_signer", "kasubsi_1")
    if subsection_code not in {"kasi_intel", "kasubsi_1", "kasubsi_2"}:
        subsection_code = "kasubsi_1"
    data["subsection_code"] = subsection_code
    subsection = fetch_one(
        "SELECT * FROM signatories WHERE position_code = %s", (subsection_code,)
    )
    data["show_authentication"] = subsection_code != "kasi_intel"
    data["use_scanned_signatures"] = (
        1 if request.form.get("use_scanned_signatures") == "1" else 0
    )
    data["use_digital_stamp"] = (
        1 if request.form.get("use_digital_stamp") == "1" else 0
    )
    letter_signature_type = request.form.get("letter_signature_type", "tte").strip()
    data["letter_signature_type"] = (
        letter_signature_type
        if letter_signature_type in {"tte", "scan", "empty"} else "tte"
    )
    data["letter_use_digital_stamp"] = (
        1 if request.form.get("letter_use_digital_stamp") == "1" else 0
    )
    data["sipede_manual"] = request.form.get("sipede_manual") == "1"
    data["sipede_number"] = request.form.get("sipede_number", "").strip()
    if not data["sipede_manual"] and not data["sipede_number"]:
        data["sipede_number"] = "-"
    acting_officer_type = request.form.get("acting_officer_type", "").strip().lower()
    data["acting_officer_type"] = (
        acting_officer_type if acting_officer_type in {"plh", "plt"} else None
    )
    data["acting_officer_name"] = request.form.get("acting_officer_name", "").strip()
    data["acting_officer_position"] = (
        "Kepala Kejaksaan Negeri Buleleng"
        if data["acting_officer_type"] else ""
    )
    data["acting_officer_rank"] = request.form.get("acting_officer_rank", "").strip()
    data["acting_officer_nip"] = request.form.get("acting_officer_nip", "").strip()
    report_acting_enabled = request.form.get("report_acting_enabled") == "1"
    data["report_acting_enabled"] = report_acting_enabled
    report_acting_type = request.form.get("report_acting_type", "").strip().lower()
    data["report_acting_type"] = (
        report_acting_type if report_acting_enabled and report_acting_type in {"plh", "plt"} else None
    )
    data["report_acting_name"] = (
        request.form.get("report_acting_name", "").strip()
        if data["report_acting_type"] else ""
    )
    data["report_acting_position"] = (
        request.form.get("report_acting_position", "").strip()
        if data["report_acting_type"] else ""
    )
    data["report_acting_nip"] = (
        request.form.get("report_acting_nip", "").strip()
        if data["report_acting_type"] else ""
    )
    if kasi:
        data.update(auth_position=kasi["position_name"], auth_name=kasi["full_name"], auth_rank_nip=kasi["rank_nip"])
    if subsection:
        data.update(creator_position=subsection["position_name"], creator_name=subsection["full_name"],
                    creator_rank_nip=subsection["rank_nip"])
    else:
        data.update(creator_position="", creator_name="", creator_rank_nip="")
    if data["report_acting_type"]:
        prefix = "Plh." if data["report_acting_type"] == "plh" else "Plt."
        nip = data["report_acting_nip"]
        nip_text = nip if "NIP" in nip.upper() else f"NIP. {nip}"
        data.update(
            auth_position=f"{prefix} Kepala Seksi Intelijen",
            auth_name=data["report_acting_name"],
            auth_rank_nip=f"{data['report_acting_position']} {nip_text}",
        )
        if data["subsection_code"] == "kasi_intel":
            data.update(
                creator_position=f"{prefix} Kepala Seksi Intelijen",
                creator_name=data["report_acting_name"],
                creator_rank_nip=f"{data['report_acting_position']} {nip_text}",
            )
    return data


def reserve_lapinhar_number(document_year=None):
    document_year = int(document_year or date.today().year)
    connection = get_db()
    token = uuid.uuid4().hex
    try:
        connection.begin()
        with connection.cursor() as cursor:
            cursor.execute(
                """SELECT id,sequence_number FROM released_document_numbers
                   WHERE document_type=%s AND document_year=%s
                   ORDER BY sequence_number LIMIT 1 FOR UPDATE""", ("lapinhar", document_year))
            released = cursor.fetchone()
            if released:
                sequence_number = released["sequence_number"]
                cursor.execute("DELETE FROM released_document_numbers WHERE id=%s", (released["id"],))
            else:
                cursor.execute("SELECT next_number FROM document_counters WHERE document_type=%s AND document_year=%s FOR UPDATE", ("lapinhar", document_year))
                counter = cursor.fetchone()
                if counter is None:
                    cursor.execute("INSERT INTO document_counters (document_type,document_year,next_number) VALUES (%s,%s,2)",
                                   ("lapinhar", document_year))
                    sequence_number = 1
                else:
                    sequence_number = counter["next_number"]
                    cursor.execute("UPDATE document_counters SET next_number=next_number+1 WHERE document_type=%s AND document_year=%s",
                                   ("lapinhar", document_year))
            cursor.execute(
                """INSERT INTO document_number_reservations
                   (reservation_token, document_type, document_year, sequence_number, created_by)
                   VALUES (%s, 'lapinhar', %s, %s, %s)""",
                (token, document_year, sequence_number, session["user_id"]),
            )
        connection.commit()
        return {"reservation_token": token, "sequence_number": sequence_number, "document_year": document_year}
    except Exception:
        connection.rollback()
        raise


def accessible_lapinhar(report_id):
    report = fetch_one("SELECT * FROM reports WHERE id = %s AND report_type = 'lapinhar'", (report_id,))
    if report and (session.get("role") == "admin" or report["created_by"] == session.get("user_id")):
        return report
    return None


def report_sequence_number(report_number):
    match = re.match(r"^R-LIH-(\d+)[A-Z]*/", report_number or "", re.IGNORECASE)
    return int(match.group(1)) if match else 0


def report_sequence_label(report_number):
    match = re.match(r"^R-LIH-(\d+[A-Z]*)/", report_number or "", re.IGNORECASE)
    return match.group(1).upper() if match else ""


def report_number_year(report_number):
    match = re.search(r"/(\d{4})$", report_number or "")
    return int(match.group(1)) if match else 0


def sequence_suffix_rank(suffix):
    rank = 0
    for character in (suffix or "").upper():
        rank = rank * 26 + (ord(character) - ord("A") + 1)
    return rank


def sequence_suffix_from_rank(rank):
    result = ""
    while rank > 0:
        rank, remainder = divmod(rank - 1, 26)
        result = chr(ord("A") + remainder) + result
    return result


def sequence_label_parts(report_number, document_type):
    prefix = r"R-LIH" if document_type == "lapinhar" else r"R-LIK"
    match = re.match(rf"^{prefix}-(\d+)([A-Z]*)/", report_number or "", re.IGNORECASE)
    return (int(match.group(1)), match.group(2).upper()) if match else (0, "")


def next_backdated_sequence_label(document_type, selected_date, report_id, cursor=None):
    table_name = "lapinhar_reports" if document_type == "lapinhar" else "lapinsus_reports"
    def query_all(query, params):
        if cursor is None:
            return fetch_all(query, params)
        cursor.execute(query, params)
        return cursor.fetchall()

    def query_one(query, params):
        if cursor is None:
            return fetch_one(query, params)
        cursor.execute(query, params)
        return cursor.fetchone()

    rows = query_all(
        f"""SELECT report_number FROM {table_name}
            WHERE report_date=%s AND id<>%s ORDER BY id""",
        (selected_date, report_id),
    )
    booked = query_all(
        """SELECT sequence_label AS report_number
           FROM backdated_number_reservations
           WHERE document_type=%s AND report_date=%s AND report_id<>%s""",
        (document_type, selected_date, report_id),
    )
    parsed = [sequence_label_parts(row["report_number"], document_type) for row in rows]
    parsed.extend(
        (int(match.group(1)), match.group(2).upper())
        for row in booked
        if (match := re.match(r"^(\d+)([A-Z]*)$", row["report_number"] or "", re.IGNORECASE))
    )
    parsed = [item for item in parsed if item[0]]
    if parsed:
        base_number = max(item[0] for item in parsed)
        highest_suffix = max(
            (sequence_suffix_rank(suffix) for number, suffix in parsed if number == base_number),
            default=0,
        )
    else:
        previous = query_one(
            f"""SELECT report_number FROM {table_name}
                WHERE report_date<%s AND id<>%s
                ORDER BY report_date DESC,id DESC LIMIT 1""",
            (selected_date, report_id),
        )
        base_number, _suffix = sequence_label_parts(
            previous["report_number"] if previous else "", document_type
        )
        if not base_number:
            base_number = 1
        highest_suffix = 0
    return f"{base_number}{sequence_suffix_from_rank(highest_suffix + 1)}"


def reserve_backdated_number(document_type, report_id, selected_date):
    if document_type not in {"lapinhar", "lapinsus"}:
        raise ValueError("Jenis laporan tidak valid.")
    connection = get_db()
    token = uuid.uuid4().hex
    lock_name = f"backdated:{document_type}:{selected_date.isoformat()}"
    lock_acquired = False
    try:
        connection.begin()
        with connection.cursor() as cursor:
            cursor.execute("SELECT GET_LOCK(%s, 5) AS acquired", (lock_name,))
            lock_acquired = bool(cursor.fetchone()["acquired"])
            if not lock_acquired:
                raise RuntimeError("Nomor sedang diproses pengguna lain. Silakan coba lagi.")
            cursor.execute(
                "DELETE FROM backdated_number_reservations "
                "WHERE reserved_at < CURRENT_TIMESTAMP - INTERVAL 2 HOUR"
            )
            cursor.execute(
                "DELETE FROM backdated_number_reservations "
                "WHERE document_type=%s AND report_id=%s",
                (document_type, report_id),
            )
            sequence_label = next_backdated_sequence_label(
                document_type, selected_date, report_id, cursor
            )
            cursor.execute(
                """INSERT INTO backdated_number_reservations
                   (reservation_token,document_type,report_id,report_date,
                    sequence_label,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s)""",
                (token, document_type, report_id, selected_date,
                 sequence_label, session["user_id"]),
            )
        connection.commit()
        return {"reservation_token": token, "sequence_label": sequence_label}
    except Exception:
        connection.rollback()
        raise
    finally:
        if lock_acquired:
            with connection.cursor() as cursor:
                cursor.execute("SELECT RELEASE_LOCK(%s)", (lock_name,))


@app.post("/reports/backdated-number/reserve")
@login_required
def reserve_backdated_number_endpoint():
    payload = request.get_json(silent=True) or {}
    document_type = str(payload.get("document_kind", "")).strip().lower()
    try:
        report_id = int(payload.get("report_id"))
        selected_date = date.fromisoformat(str(payload.get("report_date", "")))
    except (TypeError, ValueError):
        return jsonify(message="Tanggal atau laporan tidak valid."), 400
    report = (accessible_lapinhar(report_id) if document_type == "lapinhar"
              else accessible_lapinsus(report_id) if document_type == "lapinsus" else None)
    if report is None:
        abort(404)
    if not report.get("report_date") or selected_date >= report["report_date"]:
        return jsonify(message="Reservasi khusus hanya untuk tanggal mundur."), 400
    try:
        return jsonify(**reserve_backdated_number(document_type, report_id, selected_date))
    except RuntimeError as error:
        return jsonify(message=str(error)), 409


@app.post("/reports/backdated-number/cancel")
@login_required
def cancel_backdated_number_endpoint():
    payload = request.get_json(silent=True) or request.form
    token = str(payload.get("reservation_token", "")).strip()
    if token:
        with get_db().cursor() as cursor:
            cursor.execute(
                "DELETE FROM backdated_number_reservations "
                "WHERE reservation_token=%s AND created_by=%s",
                (token, session["user_id"]),
            )
    return ("", 204)


def release_original_document_number(report_id, document_type):
    connection = get_db()
    try:
        connection.begin()
        with connection.cursor() as cursor:
            cursor.execute(
                """SELECT id,document_year,sequence_number
                   FROM document_number_reservations
                   WHERE report_id=%s AND document_type=%s AND status='used'
                   FOR UPDATE""",
                (report_id, document_type),
            )
            reservation = cursor.fetchone()
            if not reservation:
                connection.commit()
                return
            cursor.execute(
                """INSERT IGNORE INTO released_document_numbers
                   (document_type,document_year,sequence_number)
                   VALUES (%s,%s,%s)""",
                (document_type, reservation["document_year"], reservation["sequence_number"]),
            )
            cursor.execute(
                "DELETE FROM document_number_reservations WHERE id=%s",
                (reservation["id"],),
            )
        connection.commit()
    except Exception:
        connection.rollback()
        raise


def compose_lapinhar_number(sequence_number, institution_code, issue_code, report_date):
    try:
        selected_date = date.fromisoformat(report_date)
    except (TypeError, ValueError):
        return ""
    if not sequence_number or issue_code not in ISSUE_CODES:
        return ""
    return f"R-LIH-{sequence_number}/{institution_code}/{issue_code}/{selected_date.month:02d}/{selected_date.year}"


def lapinhar_number_available(report_number, reservation_token="", report_id=None):
    duplicate = fetch_one("SELECT id FROM reports WHERE report_number=%s", (report_number,))
    if duplicate and duplicate["id"] != report_id:
        return False, "Nomor surat sudah digunakan laporan lain. Silakan reload nomor surat."
    if report_id:
        return True, "Nomor surat tersedia."
    reservation = fetch_one(
        """SELECT sequence_number,document_year FROM document_number_reservations
           WHERE reservation_token=%s AND document_type='lapinhar'
           AND created_by=%s AND status='reserved'""",
        (reservation_token, session["user_id"]),
    )
    if (not reservation or reservation["sequence_number"] != report_sequence_number(report_number)
            or reservation["document_year"] != report_number_year(report_number)):
        return False, "Reservasi nomor surat tidak berlaku. Silakan reload nomor surat."
    return True, "Nomor surat tersedia."


@app.post("/lapinhar/check-number")
@login_required
def check_lapinhar_number():
    payload = request.get_json(silent=True) or {}
    try:
        report_id = int(payload["report_id"]) if payload.get("report_id") else None
    except (TypeError, ValueError):
        report_id = None
    if report_id and accessible_lapinhar(report_id) is None:
        abort(404)
    available, message = lapinhar_number_available(
        str(payload.get("report_number", "")).strip(),
        str(payload.get("reservation_token", "")).strip(), report_id,
    )
    return jsonify(available=available, message=message), 200 if available else 409


@app.post("/lapinhar/reload-number")
@login_required
def reload_lapinhar_number():
    payload = request.get_json(silent=True) or {}
    reservation = reserve_lapinhar_number(payload.get("document_year"))
    return jsonify(**reservation)


def normalize_attachment_image(content):
    try:
        with PILImage.open(BytesIO(content)) as source_image:
            normalized = ImageOps.exif_transpose(source_image)
            if normalized.mode not in ("RGB", "L"):
                background = PILImage.new("RGB", normalized.size, "white")
                alpha_image = normalized.convert("RGBA")
                background.paste(alpha_image, mask=alpha_image.getchannel("A"))
                normalized = background
            elif normalized.mode == "L":
                normalized = normalized.convert("RGB")
            else:
                normalized = normalized.copy()
    except Exception:
        return None

    width, height = normalized.size
    longest_side = max(width, height)
    if longest_side > MAX_ATTACHMENT_DIMENSION:
        resize_ratio = MAX_ATTACHMENT_DIMENSION / float(longest_side)
        normalized = normalized.resize(
            (max(1, int(width * resize_ratio)), max(1, int(height * resize_ratio))),
            PILImage.Resampling.LANCZOS,
        )

    working = normalized
    quality_steps = (82, 74, 68, 60, 54, 48, 42, 36, 30)
    for _attempt in range(6):
        for quality in quality_steps:
            output = BytesIO()
            working.save(output, format="JPEG", quality=quality, optimize=True)
            compressed = output.getvalue()
            if len(compressed) <= MAX_ATTACHMENT_SIZE_BYTES:
                return {
                    "content": compressed,
                    "width": working.width,
                    "height": working.height,
                    "extension": ".jpg",
                }
        reduced_width = max(1, int(working.width * 0.88))
        reduced_height = max(1, int(working.height * 0.88))
        if (reduced_width, reduced_height) == working.size:
            break
        working = working.resize((reduced_width, reduced_height), PILImage.Resampling.LANCZOS)

    output = BytesIO()
    working.save(output, format="JPEG", quality=30, optimize=True)
    return {
        "content": output.getvalue(),
        "width": working.width,
        "height": working.height,
        "extension": ".jpg",
    }


def collect_attachments():
    attachments = []
    keys = sorted((key for key in request.files if key.startswith("attachment_images_")),
                  key=lambda key: int(key.rsplit("_", 1)[-1]))
    for key in keys:
        group_number = int(key.rsplit("_", 1)[-1])
        images = []
        for uploaded in request.files.getlist(key)[:2]:
            if not uploaded.filename or uploaded.mimetype not in ALLOWED_IMAGE_TYPES:
                continue
            content = uploaded.read()
            normalized_image = normalize_attachment_image(content)
            if not normalized_image:
                continue
            uploaded.seek(0)
            image_index = len(images) + 1
            scale = max(40, min(100, request.form.get(f"attachment_scale_{group_number}_{image_index}", 100, type=int)))
            original_name = Path(secure_filename(uploaded.filename)).stem or f"lampiran_{group_number}_{image_index}"
            images.append({
                "filename": f"{original_name}{normalized_image['extension']}",
                "content": normalized_image["content"],
                "width": normalized_image["width"],
                "height": normalized_image["height"],
                "scale": scale,
            })
        if images:
            attachments.append({
                "group": group_number,
                "title": "",
                "layout": request.form.get(f"attachment_layout_{group_number}", "auto"),
                "images": images,
            })
    return attachments


def attachment_label(count):
    """Build the memo label from attachment pages that actually contain photos."""
    if not count:
        return "-"
    number_words = {1: "satu", 2: "dua", 3: "tiga", 4: "empat", 5: "lima",
                    6: "enam", 7: "tujuh", 8: "delapan", 9: "sembilan", 10: "sepuluh"}
    return f"{count} ({number_words.get(count, str(count))}) Lembar"


def save_report_attachments(report_id, attachments):
    report_dir = UPLOAD_DIR / str(report_id)
    report_dir.mkdir(parents=True, exist_ok=True)
    with get_db().cursor() as cursor:
        for attachment in attachments:
            for image_index, image in enumerate(attachment["images"], 1):
                filename = f"{attachment['group']}_{image_index}_{uuid.uuid4().hex[:8]}_{image['filename']}"
                path = report_dir / filename
                path.write_bytes(image["content"])
                cursor.execute(
                    """INSERT INTO report_attachments
                    (report_id, attachment_group, image_path, caption, sort_order)
                    VALUES (%s, %s, %s, %s, %s)""",
                    (report_id, attachment["group"], str(path.relative_to(Path(__file__).parent)),
                     attachment["title"], image_index),
                )


def clone_report_attachments(source_report_id, target_report_id):
    """Copy every attachment file and row without modifying the LAPINHAR originals."""
    if fetch_one("SELECT id FROM report_attachments WHERE report_id=%s LIMIT 1", (target_report_id,)):
        return 0
    rows = fetch_all(
        "SELECT attachment_group,image_path,caption,sort_order FROM report_attachments "
        "WHERE report_id=%s ORDER BY attachment_group,sort_order,id", (source_report_id,),
    )
    source_dir = (UPLOAD_DIR / str(source_report_id)).resolve()
    target_dir = (UPLOAD_DIR / str(target_report_id)).resolve()
    target_dir.mkdir(parents=True, exist_ok=True)
    copied = []
    try:
        for row in rows:
            source_path = (Path(__file__).resolve().parent / row["image_path"]).resolve()
            if not source_path.is_relative_to(source_dir) or not source_path.is_file():
                continue
            target_name = f"{row['attachment_group']}_{row['sort_order']}_{uuid.uuid4().hex[:8]}_{source_path.name}"
            target_path = target_dir / target_name
            shutil.copy2(source_path, target_path)
            copied.append((row, target_path))
        if copied:
            with get_db().cursor() as cursor:
                cursor.executemany(
                    "INSERT INTO report_attachments (report_id,attachment_group,image_path,caption,sort_order) "
                    "VALUES (%s,%s,%s,%s,%s)",
                    [(target_report_id, row["attachment_group"],
                      str(path.relative_to(Path(__file__).resolve().parent)), row["caption"], row["sort_order"])
                     for row, path in copied],
                )
        return len(copied)
    except Exception:
        for _row, path in copied:
            try:
                path.unlink()
            except OSError:
                pass
        raise


def remove_attachment_files(rows, report_id):
    """Remove known attachment files without allowing OneDrive locks to break report updates."""
    upload_root = UPLOAD_DIR.resolve()
    all_removed = True
    for row in rows:
        path = (Path(__file__).parent / row["image_path"]).resolve()
        if not path.is_relative_to(upload_root) or not path.exists():
            continue
        try:
            path.chmod(stat.S_IWRITE)
            path.unlink()
        except OSError:
            all_removed = False
            app.logger.warning("Berkas lampiran sedang terkunci dan belum dapat dihapus: %s", path)
    report_dir = (UPLOAD_DIR / str(report_id)).resolve()
    if report_dir.is_relative_to(upload_root) and report_dir.exists():
        try:
            report_dir.rmdir()
        except OSError:
            pass
    return all_removed


def existing_report_files(report_id):
    """Snapshot regular files in one report folder before replacement/deletion."""
    upload_root = UPLOAD_DIR.resolve()
    report_dir = (UPLOAD_DIR / str(report_id)).resolve()
    if not report_dir.is_relative_to(upload_root) or not report_dir.is_dir():
        return []
    return [
        {"image_path": str(path.relative_to(Path(__file__).parent))}
        for path in report_dir.iterdir()
        if path.is_file()
    ]


def attachment_layout(attachment):
    requested = attachment.get("layout")
    if requested in {"side", "stack"}:
        return requested
    images = attachment["images"]
    return "side" if len(images) == 2 and all(image["height"] > image["width"] for image in images) else "stack"


def create_empty_report_draft(document_type, reservation):
    if document_type not in {"lapinhar", "lapinsus"}:
        raise ValueError("Jenis draft laporan tidak valid.")
    organization = fetch_one("SELECT * FROM organization_settings WHERE id=1") or {
        "organization_name": "KEJAKSAAN NEGERI BULELENG", "institution_code": "N.1.11",
    }
    signatories = {
        row["position_code"]: row for row in
        fetch_all("SELECT * FROM signatories WHERE position_code IN ('kasi_intel','kasubsi_1')")
    }
    kasi, creator = signatories.get("kasi_intel") or {}, signatories.get("kasubsi_1") or {}
    selected_date = date.today()
    prefix = "R-LIH" if document_type == "lapinhar" else "R-LIK"
    provisional_number = (
        f"{prefix}-{reservation['sequence_number']}/"
        f"{organization.get('institution_code') or 'N.1.11'}/-/"
        f"{selected_date.month:02d}/{selected_date.year}"
    )
    table_name = "lapinhar_reports" if document_type == "lapinhar" else "lapinsus_reports"
    with get_db().cursor() as cursor:
        cursor.execute(
            f"""INSERT INTO {table_name}
            (report_type,report_number,title,report_date,facts,source_name,analysis,recommendation,
             recipient,sender_name,classification,category_id,issue_code,attachment,organization,
             creator_position,creator_name,creator_rank_nip,auth_position,auth_name,auth_rank_nip,
             information_spacing,sources_spacing,trends_spacing,suggestions_spacing,status,created_by)
            VALUES (%s,%s,'',%s,'','','','',%s,'KASI INTELIJEN','RAHASIA',NULL,NULL,'-',%s,
                    %s,%s,%s,%s,%s,%s,'1.5','1.5','1.5','1.5','draft',%s)""",
            (
                document_type, provisional_number, selected_date,
                f"YTH. KEPALA {organization['organization_name']}", organization["organization_name"],
                creator.get("position_name", ""), creator.get("full_name", ""),
                creator.get("rank_nip", ""), kasi.get("position_name", ""),
                kasi.get("full_name", ""), kasi.get("rank_nip", ""), session["user_id"],
            ),
        )
        report_id = cursor.lastrowid
        cursor.execute(
            """UPDATE document_number_reservations
               SET status='used',report_id=%s,used_at=CURRENT_TIMESTAMP
               WHERE reservation_token=%s AND document_type=%s
               AND created_by=%s AND status='reserved'""",
            (report_id, reservation["reservation_token"], document_type, session["user_id"]),
        )
    return report_id


@app.route("/lapinhar/create", methods=["GET", "POST"])
@login_required
def create_lapinhar():
    if request.method == "GET":
        if request.args.get("new") != "1":
            existing_draft = fetch_one(
                """SELECT id FROM lapinhar_reports
                   WHERE created_by=%s AND status='draft'
                   ORDER BY created_at DESC,id DESC LIMIT 1""",
                (session["user_id"],),
            )
            if existing_draft:
                return redirect(url_for("lapinhar", choose_draft=existing_draft["id"]))
        reservation = reserve_lapinhar_number()
        report_id = create_empty_report_draft("lapinhar", reservation)
        return redirect(url_for("edit_lapinhar", report_id=report_id))
    if request.method == "POST":
        data = lapinhar_form_data()
        attachments = collect_attachments()
        data["attachment"] = attachment_label(len(attachments))
        reservation_token = request.form.get("number_reservation_token", "")
        reservation = fetch_one(
            """SELECT * FROM document_number_reservations
               WHERE reservation_token = %s AND document_type = 'lapinhar'
               AND created_by = %s AND status = 'reserved'""",
            (reservation_token, session["user_id"]),
        )
        data["report_number"] = compose_lapinhar_number(
            reservation["sequence_number"] if reservation else 0,
            data["institution_code"], data["issue_code"], data["report_date"],
        )
        number_available, number_message = lapinhar_number_available(
            data["report_number"], reservation_token
        )
        if (not data["report_number"] or not data["subject"] or not data["information"]
                or not data["category_id"] or not data["issue_code"] or not reservation):
            flash("Nomor reservasi, nomor permasalahan, kategori, perihal, dan informasi wajib diisi.", "error")
        elif not number_available:
            flash(number_message, "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    """INSERT INTO lapinhar_reports
                    (report_type, report_number, title, report_date, facts, source_name, analysis, recommendation,
                     recipient, sender_name, classification, category_id, issue_code, attachment, organization, creator_position,
                     creator_name, creator_rank_nip, auth_position, auth_name, auth_rank_nip,
                     information_spacing, sources_spacing, trends_spacing, suggestions_spacing,
                     status, created_by)
                    VALUES ('lapinhar', %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                            %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 'draft', %s)""",
                    (data["report_number"], data["subject"], data["report_date"] or None,
                     data["information"], data["sources"], data["trends"], data["suggestions"],
                     data["recipient"], data["sender"], data["classification"], data["category_id"],
                     data["issue_code"], data["attachment"],
                     data["organization"], data["creator_position"], data["creator_name"],
                     data["creator_rank_nip"], data["auth_position"], data["auth_name"], data["auth_rank_nip"],
                     data["information_spacing"] or "1.15", data["sources_spacing"] or "1.15",
                     data["trends_spacing"] or "1.15", data["suggestions_spacing"] or "1.15",
                    session["user_id"]),
                )
                report_id = cursor.lastrowid
                cursor.execute(
                    """UPDATE document_number_reservations SET status = 'used', report_id = %s,
                       used_at = CURRENT_TIMESTAMP WHERE id = %s AND status = 'reserved'""",
                    (report_id, reservation["id"]),
                )
            save_report_attachments(report_id, attachments)
            flash("LAPINHAR berhasil disimpan.", "success")
            after_save = request.form.get("after_save", "")
            if after_save in {"docx", "print", "pdf", "sipede"}:
                return redirect(url_for("edit_lapinhar", report_id=report_id, after_save=after_save))
            return redirect(url_for("lapinhar"))
    signatories = {row["position_code"]: row for row in fetch_all("SELECT * FROM signatories")}
    organization = fetch_one("SELECT * FROM organization_settings WHERE id = 1") or {
        "organization_name": "KEJAKSAAN NEGERI BULELENG", "institution_code": "N.1.11",
        "address": "", "phone": "", "website": ""
    }
    number_reservation = None
    if request.method == "POST":
        number_reservation = fetch_one(
            """SELECT reservation_token, sequence_number FROM document_number_reservations
               WHERE reservation_token = %s AND created_by = %s AND status = 'reserved'""",
            (request.form.get("number_reservation_token", ""), session["user_id"]),
        )
    if number_reservation is None:
        number_reservation = reserve_lapinhar_number()
    return render_template("lapinhar_editor.html", active="lapinhar", today=date.today().isoformat(),
                           signatories=signatories, organization=organization,
                           number_reservation=number_reservation,
                           fixed_recipient=f"YTH. KEPALA {organization['organization_name']}",
                           fixed_sender="KASI INTELIJEN", report=None, existing_sources=[], existing_attachments=[])


@app.route("/lapinhar/<int:report_id>/edit", methods=["GET", "POST"])
@login_required
def edit_lapinhar(report_id):
    report = accessible_lapinhar(report_id)
    if report is None:
        flash("Laporan tidak ditemukan atau tidak dapat Anda akses.", "error")
        return redirect(url_for("lapinhar"))
    sequence_number = report_sequence_label(report["report_number"])
    if request.method == "POST":
        data = lapinhar_form_data()
        submitted_sequence = report_sequence_label(request.form.get("report_number", ""))
        try:
            selected_report_date = date.fromisoformat(data["report_date"])
        except (TypeError, ValueError):
            selected_report_date = None
        backdated = bool(
            selected_report_date and report.get("report_date")
            and selected_report_date < report["report_date"]
        )
        if backdated:
            backdated_reservation = fetch_one(
                """SELECT id,sequence_label FROM backdated_number_reservations
                   WHERE reservation_token=%s AND document_type='lapinhar'
                   AND report_id=%s AND report_date=%s AND created_by=%s""",
                (request.form.get("backdated_reservation_token", ""), report_id,
                 selected_report_date, session["user_id"]),
            )
            submitted_sequence = (
                backdated_reservation["sequence_label"] if backdated_reservation else ""
            )
        else:
            backdated_reservation = None
        data["report_number"] = compose_lapinhar_number(
            submitted_sequence or sequence_number, data["institution_code"], data["issue_code"], data["report_date"],
        )
        attachments = collect_attachments()
        old_attachment_files = existing_report_files(report_id) if attachments else []
        old_attachment_rows = fetch_all(
            "SELECT id FROM report_attachments WHERE report_id=%s", (report_id,)
        ) if attachments else []
        number_available, number_message = lapinhar_number_available(
            data["report_number"], report_id=report_id
        )
        if backdated and not backdated_reservation:
            flash("Nomor tanggal mundur belum dibooking. Pilih kembali tanggal laporan.", "error")
        elif (data["report_acting_enabled"] and not all((
                data["report_acting_type"],
                data["report_acting_name"], data["report_acting_position"],
                data["report_acting_nip"]))):
            flash("Jenis, nama, jabatan, dan NIP PLT/PLH penandatangan wajib diisi.", "error")
        elif (not data["report_number"] or not data["subject"] or not data["information"]
                or not data["category_id"] or not data["issue_code"]):
            flash("Nomor permasalahan, kategori, perihal, dan informasi wajib diisi.", "error")
        elif not number_available:
            flash(number_message, "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    """UPDATE lapinhar_reports SET report_number=%s, title=%s, report_date=%s, facts=%s,
                       source_name=%s, analysis=%s, recommendation=%s, recipient=%s, sender_name=%s,
                       classification='RAHASIA', category_id=%s, issue_code=%s, attachment=%s,
                       organization=%s, creator_position=%s, creator_name=%s, creator_rank_nip=%s,
                       auth_position=%s, auth_name=%s, auth_rank_nip=%s,
                       use_scanned_signatures=%s,use_digital_stamp=%s,
                       report_acting_type=%s,report_acting_name=%s,
                       report_acting_position=%s,report_acting_nip=%s,
                       information_spacing='1.5',
                       sources_spacing='1.5', trends_spacing='1.5', suggestions_spacing='1.5',
                       status='selesai'
                       WHERE id=%s""",
                    (data["report_number"], data["subject"], data["report_date"] or None,
                     data["information"], data["sources"], data["trends"], data["suggestions"],
                     data["recipient"], data["sender"], data["category_id"], data["issue_code"],
                     attachment_label(len(attachments)) if attachments else report["attachment"],
                     data["organization"], data["creator_position"], data["creator_name"],
                     data["creator_rank_nip"], data["auth_position"], data["auth_name"],
                     data["auth_rank_nip"], data["use_scanned_signatures"],
                     data["use_digital_stamp"], data["report_acting_type"],
                     data["report_acting_name"], data["report_acting_position"],
                     data["report_acting_nip"], report_id),
                )
                reservation_token = request.form.get("number_reservation_token", "")
                if reservation_token:
                    cursor.execute(
                        "UPDATE document_number_reservations SET status='used',report_id=%s,used_at=CURRENT_TIMESTAMP "
                        "WHERE reservation_token=%s AND document_type='lapinhar' AND created_by=%s AND status='reserved'",
                        (report_id, reservation_token, session["user_id"]),
                    )
            if backdated:
                release_original_document_number(report_id, "lapinhar")
                with get_db().cursor() as cursor:
                    cursor.execute(
                        "DELETE FROM backdated_number_reservations WHERE id=%s",
                        (backdated_reservation["id"],),
                    )
            if attachments:
                save_report_attachments(report_id, attachments)
                if old_attachment_rows:
                    with get_db().cursor() as cursor:
                        cursor.executemany(
                            "DELETE FROM report_attachments WHERE id=%s",
                            [(row["id"],) for row in old_attachment_rows],
                        )
                remove_attachment_files(old_attachment_files, report_id)
            flash("LAPINHAR berhasil diperbarui.", "success")
            after_save = request.form.get("after_save", "")
            if after_save in {"docx", "print", "pdf", "sipede"}:
                return redirect(url_for("edit_lapinhar", report_id=report_id, after_save=after_save))
            return redirect(url_for("lapinhar"))

    signatories = {row["position_code"]: row for row in fetch_all("SELECT * FROM signatories")}
    organization = fetch_one("SELECT * FROM organization_settings WHERE id=1")
    existing_sources = []
    for block in rich_text_blocks(report["source_name"] or ""):
        if block["text"].strip().lower() != "intelijen kejaksaan negeri buleleng":
            existing_sources.append(block["text"].strip())
    existing_attachments = fetch_all(
        "SELECT attachment_group, image_path, sort_order FROM report_attachments WHERE report_id=%s ORDER BY attachment_group, sort_order",
        (report_id,),
    )
    existing_attachment_groups = []
    for row in existing_attachments:
        group = next((item for item in existing_attachment_groups
                      if item["group"] == row["attachment_group"]), None)
        if group is None:
            group = {"group": row["attachment_group"], "images": []}
            existing_attachment_groups.append(group)
        group["images"].append({
            "url": url_for("lapinhar_attachment_file", report_id=report_id,
                           filename=Path(row["image_path"]).name),
            "filename": Path(row["image_path"]).name,
        })
    selected_signer = "kasi_intel"
    for code in ("kasubsi_1", "kasubsi_2"):
        signer = signatories.get(code)
        if not signer:
            continue
        report_nip = re.sub(r"\D", "", report.get("creator_rank_nip") or "")
        signer_nip = re.sub(r"\D", "", signer.get("rank_nip") or "")
        same_position = signer["position_name"] == report["creator_position"]
        same_name = bool(report.get("creator_name") and signer.get("full_name")
                         and signer["full_name"].casefold() == report["creator_name"].casefold())
        same_nip = bool(report_nip and signer_nip and report_nip == signer_nip)
        if same_position or same_name or same_nip:
            selected_signer = code
            break
    number_reservation = {"reservation_token": "", "sequence_number": sequence_number}
    return render_template(
        "lapinhar_editor.html", active="lapinhar", today=report["report_date"].isoformat() if report["report_date"] else date.today().isoformat(),
        signatories=signatories, organization=organization, number_reservation=number_reservation,
        fixed_recipient=report["recipient"], fixed_sender=report["sender_name"], report=report,
        existing_sources=existing_sources, existing_attachments=existing_attachments,
        existing_attachment_groups=existing_attachment_groups,
        selected_signer=selected_signer,
    )


@app.get("/lapinhar/<int:report_id>/attachments/<path:filename>")
@login_required
def lapinhar_attachment_file(report_id, filename):
    if accessible_lapinhar(report_id) is None or Path(filename).name != filename:
        abort(404)
    row = next((item for item in fetch_all(
        "SELECT image_path FROM report_attachments WHERE report_id=%s", (report_id,)
    ) if Path(item["image_path"]).name == filename), None)
    if row is None or Path(row["image_path"]).name != filename:
        abort(404)
    return send_from_directory(UPLOAD_DIR / str(report_id), filename)


@app.get("/lapinhar/<int:report_id>/attachments/<path:filename>/download")
@login_required
def download_lapinhar_attachment(report_id, filename):
    if accessible_lapinhar(report_id) is None or Path(filename).name != filename:
        abort(404)
    row = next((item for item in fetch_all(
        """SELECT attachment_group, sort_order, image_path FROM report_attachments
           WHERE report_id=%s""", (report_id,)
    ) if Path(item["image_path"]).name == filename), None)
    if row is None or Path(row["image_path"]).name != filename:
        abort(404)
    extension = Path(filename).suffix.lower() or ".jpg"
    download_name = f"Lampiran-{row['attachment_group']}-Foto-{row['sort_order']}{extension}"
    return send_from_directory(UPLOAD_DIR / str(report_id), filename,
                               as_attachment=True, download_name=download_name)


@app.get("/lapinhar/<int:report_id>/attachments/download-all")
@login_required
def download_all_lapinhar_attachments(report_id):
    report = accessible_lapinhar(report_id)
    if report is None:
        abort(404)
    rows = fetch_all(
        """SELECT attachment_group, sort_order, image_path FROM report_attachments
           WHERE report_id=%s ORDER BY attachment_group, sort_order, id""",
        (report_id,),
    )
    if not rows:
        abort(404)
    archive = BytesIO()
    report_dir = (UPLOAD_DIR / str(report_id)).resolve()
    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as bundle:
        for row in rows:
            source = (Path(__file__).resolve().parent / row["image_path"]).resolve()
            if not source.is_relative_to(report_dir) or not source.is_file():
                continue
            extension = source.suffix.lower() or ".jpg"
            archive_name = f"Lampiran-{row['attachment_group']}-Foto-{row['sort_order']}{extension}"
            bundle.write(source, archive_name)
    if not archive.getvalue():
        abort(404)
    archive.seek(0)
    sequence = report_sequence_number(report["report_number"]) or report_id
    return send_file(archive, as_attachment=True,
                     download_name=f"Lampiran-RLIH-{sequence}.zip",
                     mimetype="application/zip")


@app.get("/lapinsus/<int:report_id>/attachments/download-all")
@login_required
def download_all_lapinsus_attachments(report_id):
    report = accessible_lapinsus(report_id)
    if report is None:
        abort(404)
    rows = fetch_all(
        """SELECT attachment_group, sort_order, image_path FROM report_attachments
           WHERE report_id=%s ORDER BY attachment_group, sort_order, id""",
        (report_id,),
    )
    if not rows:
        abort(404)
    archive = BytesIO()
    report_dir = (UPLOAD_DIR / str(report_id)).resolve()
    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as bundle:
        for row in rows:
            source = (Path(__file__).resolve().parent / row["image_path"]).resolve()
            if not source.is_relative_to(report_dir) or not source.is_file():
                continue
            extension = source.suffix.lower() or ".jpg"
            archive_name = f"Lampiran-{row['attachment_group']}-Foto-{row['sort_order']}{extension}"
            bundle.write(source, archive_name)
    if not archive.getvalue():
        abort(404)
    archive.seek(0)
    sequence = lapinsus_sequence_label(report["report_number"]) or report_id
    return send_file(archive, as_attachment=True,
                     download_name=f"Lampiran-RLIK-{sequence}.zip",
                     mimetype="application/zip")


@app.post("/lapinhar/<int:report_id>/delete")
@login_required
def delete_lapinhar(report_id):
    report = accessible_lapinhar(report_id)
    if report is None:
        flash("Laporan tidak ditemukan atau tidak dapat Anda akses.", "error")
        return redirect(url_for("lapinhar"))
    attachment_files = existing_report_files(report_id)
    with get_db().cursor() as cursor:
        cursor.execute("DELETE FROM report_attachments WHERE report_id=%s", (report_id,))
        cursor.execute("DELETE FROM lapinhar_reports WHERE id=%s", (report_id,))
    files_removed = remove_attachment_files(attachment_files, report_id)
    flash(f"LAPINHAR {report['report_number']} berhasil dihapus permanen.", "success")
    if not files_removed:
        flash("Sebagian berkas foto sedang dikunci OneDrive dan belum dapat dibersihkan, tetapi data laporan sudah terhapus.", "warning")
    return redirect(url_for("lapinhar"))


def lapinsus_editor_context(report=None, number_reservation=None):
    signatories = {row["position_code"]: row for row in fetch_all("SELECT * FROM signatories")}
    organization = fetch_one("SELECT * FROM organization_settings WHERE id=1") or {
        "organization_name": "KEJAKSAAN NEGERI BULELENG", "institution_code": "N.1.11",
        "address": "", "phone": "", "website": "",
    }
    existing_sources, existing_attachments, groups = [], [], []
    selected_signer = "kasubsi_1"
    if report:
        for block in rich_text_blocks(report["source_name"] or ""):
            if block["text"].strip().lower() != "intelijen kejaksaan negeri buleleng":
                existing_sources.append(block["text"].strip())
        existing_attachments = fetch_all(
            "SELECT attachment_group,image_path,sort_order FROM report_attachments "
            "WHERE report_id=%s ORDER BY attachment_group,sort_order", (report["id"],),
        )
        for row in existing_attachments:
            group = next((item for item in groups if item["group"] == row["attachment_group"]), None)
            if group is None:
                group = {"group": row["attachment_group"], "images": []}
                groups.append(group)
            group["images"].append({
                "url": url_for("lapinsus_attachment_file", report_id=report["id"],
                               filename=Path(row["image_path"]).name),
                "filename": Path(row["image_path"]).name,
            })
        if report.get("report_acting_type"):
            selected_signer = "kasi_intel"
        else:
            for code in ("kasi_intel", "kasubsi_1", "kasubsi_2"):
                signer = signatories.get(code)
                if signer and (signer.get("position_name") == report.get("creator_position")
                               or (signer.get("full_name") or "").casefold() == (report.get("creator_name") or "").casefold()):
                    selected_signer = code
                    break
    return dict(
        active="lapinsus", document_kind="lapinsus", document_label="LAPINSUS",
        report_heading="LAPORAN INFORMASI KHUSUS", report_code="L. IN.2",
        list_url=url_for("lapinsus"), check_number_url=url_for("check_lapinsus_number"),
        reload_number_url=url_for("reload_lapinsus_number"),
        today=(report["report_date"].isoformat() if report and report["report_date"] else date.today().isoformat()),
        signatories=signatories, organization=organization, number_reservation=number_reservation,
        fixed_recipient=(report["recipient"] if report else f"YTH. KEPALA {organization['organization_name']}"),
        fixed_sender=(report["sender_name"] if report else "KASI INTELIJEN"), report=report,
        existing_sources=existing_sources, existing_attachments=existing_attachments,
        existing_attachment_groups=groups, selected_signer=selected_signer,
    )


def insert_lapinsus(data, reservation, attachments):
    with get_db().cursor() as cursor:
        cursor.execute(
            """INSERT INTO lapinsus_reports
            (report_type,report_number,title,report_date,facts,source_name,analysis,recommendation,
             recipient,sender_name,classification,category_id,issue_code,attachment,organization,
             creator_position,creator_name,creator_rank_nip,auth_position,auth_name,auth_rank_nip,
             information_spacing,sources_spacing,trends_spacing,suggestions_spacing,status,created_by)
            VALUES ('lapinsus',%s,%s,%s,%s,%s,%s,%s,%s,%s,'RAHASIA',%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,
                    '1.5','1.5','1.5','1.5','draft',%s)""",
            (data["report_number"], data["subject"], data["report_date"] or None,
             data["information"], data["sources"], data["trends"], data["suggestions"],
             data["recipient"], data["sender"], data["category_id"], data["issue_code"],
             attachment_label(len(attachments)), data["organization"], data["creator_position"],
             data["creator_name"], data["creator_rank_nip"], data["auth_position"], data["auth_name"],
             data["auth_rank_nip"], session["user_id"]),
        )
        report_id = cursor.lastrowid
        cursor.execute("UPDATE document_number_reservations SET status='used',report_id=%s,used_at=CURRENT_TIMESTAMP "
                       "WHERE id=%s AND status='reserved'", (report_id, reservation["id"]))
    save_report_attachments(report_id, attachments)
    return report_id


@app.post("/lapinhar/<int:report_id>/lapinsuskan")
@login_required
def convert_lapinhar_to_lapinsus(report_id):
    source = accessible_lapinhar(report_id)
    if source is None:
        flash("LAPINHAR tidak ditemukan atau tidak dapat Anda akses.", "error")
        return redirect(url_for("lapinhar"))
    if source.get("status") == "draft":
        flash("Lengkapi dan simpan LAPINHAR sebelum dijadikan LAPINSUS.", "warning")
        return redirect(url_for("lapinhar"))
    existing = fetch_one("SELECT id FROM lapinsus_reports WHERE source_lapinhar_id=%s", (report_id,))
    if existing:
        flash("LAPINHAR tersebut sudah pernah dijadikan LAPINSUS.", "warning")
        return redirect(url_for("edit_lapinsus", report_id=existing["id"]))
    organization = fetch_one("SELECT institution_code FROM organization_settings WHERE id=1") or {"institution_code": "N.1.11"}
    reservation = reserve_lapinsus_number(source["report_date"].year if source.get("report_date") else date.today().year)
    report_number = compose_lapinsus_number(
        reservation["sequence_number"], organization["institution_code"],
        source.get("issue_code") or "", source["report_date"].isoformat() if source.get("report_date") else "",
    )
    if not report_number:
        flash("Tanggal atau nomor permasalahan LAPINHAR belum lengkap.", "error")
        return redirect(url_for("edit_lapinhar", report_id=report_id))
    with get_db().cursor() as cursor:
        cursor.execute(
            """INSERT INTO lapinsus_reports
            (report_type,report_number,title,report_date,source_name,facts,analysis,recommendation,
             recipient,sender_name,classification,category_id,issue_code,attachment,organization,
             creator_position,creator_name,creator_rank_nip,auth_position,auth_name,auth_rank_nip,
             information_spacing,sources_spacing,trends_spacing,suggestions_spacing,status,created_by,source_lapinhar_id)
            VALUES ('lapinsus',%s,%s,%s,%s,%s,%s,%s,%s,%s,'RAHASIA',%s,%s,'-',%s,%s,%s,%s,%s,%s,%s,
                    '1.5','1.5','1.5','1.5','draft',%s,%s)""",
            (report_number, source["title"], source["report_date"], source["source_name"], source["facts"],
             source["analysis"], source["recommendation"], source["recipient"], source["sender_name"],
             source["category_id"], source["issue_code"], source["organization"], source["creator_position"],
             source["creator_name"], source["creator_rank_nip"], source["auth_position"], source["auth_name"],
             source["auth_rank_nip"], session["user_id"], report_id),
        )
        lapinsus_id = cursor.lastrowid
        cursor.execute("UPDATE lapinhar_reports SET lapinsus_status='sudah' WHERE id=%s", (report_id,))
        cursor.execute("UPDATE document_number_reservations SET status='used',report_id=%s,used_at=CURRENT_TIMESTAMP "
                       "WHERE reservation_token=%s", (lapinsus_id, reservation["reservation_token"]))
    attachment_count = clone_report_attachments(report_id, lapinsus_id)
    if attachment_count:
        with get_db().cursor() as cursor:
            cursor.execute("UPDATE lapinsus_reports SET attachment=%s WHERE id=%s",
                           (attachment_label(len({row['attachment_group'] for row in fetch_all('SELECT attachment_group FROM report_attachments WHERE report_id=%s', (lapinsus_id,))})), lapinsus_id))
    flash(f"Draft LAPINSUS berhasil dibuat dari LAPINHAR. {attachment_count} foto lampiran ikut disalin.", "success")
    return redirect(url_for("edit_lapinsus", report_id=lapinsus_id))


@app.route("/lapinsus/create", methods=["GET", "POST"])
@login_required
def create_lapinsus():
    if request.method == "GET":
        if request.args.get("new") != "1":
            existing_draft = fetch_one(
                """SELECT id FROM lapinsus_reports
                   WHERE created_by=%s AND status='draft'
                   ORDER BY created_at DESC,id DESC LIMIT 1""",
                (session["user_id"],),
            )
            if existing_draft:
                return redirect(url_for("lapinsus", choose_draft=existing_draft["id"]))
        reservation = reserve_lapinsus_number()
        report_id = create_empty_report_draft("lapinsus", reservation)
        return redirect(url_for("edit_lapinsus", report_id=report_id))
    reservation = None
    if request.method == "POST":
        data = lapinhar_form_data()
        attachments = collect_attachments()
        token = request.form.get("number_reservation_token", "")
        reservation = fetch_one(
            "SELECT * FROM document_number_reservations WHERE reservation_token=%s "
            "AND document_type='lapinsus' AND created_by=%s AND status='reserved'",
            (token, session["user_id"]),
        )
        data["report_number"] = compose_lapinsus_number(
            reservation["sequence_number"] if reservation else 0, data["institution_code"],
            data["issue_code"], data["report_date"],
        )
        available, message = lapinsus_number_available(data["report_number"], token)
        if not all((data["report_number"], data["subject"], data["information"],
                    data["category_id"], data["issue_code"], reservation)):
            flash("Nomor reservasi, nomor permasalahan, kategori, perihal, dan informasi wajib diisi.", "error")
        elif not available:
            flash(message, "error")
        else:
            report_id = insert_lapinsus(data, reservation, attachments)
            flash("LAPINSUS berhasil disimpan.", "success")
            after_save = request.form.get("after_save", "")
            if after_save in {"docx", "print", "pdf", "sipede"}:
                return redirect(url_for("edit_lapinsus", report_id=report_id, after_save=after_save))
            return redirect(url_for("lapinsus"))
    if reservation is None:
        reservation = reserve_lapinsus_number()
    return render_template("lapinhar_editor.html", **lapinsus_editor_context(number_reservation=reservation))


@app.route("/lapinsus/<int:report_id>/edit", methods=["GET", "POST"])
@login_required
def edit_lapinsus(report_id):
    report = accessible_lapinsus(report_id)
    if report is None:
        flash("LAPINSUS tidak ditemukan atau tidak dapat Anda akses.", "error")
        return redirect(url_for("lapinsus"))
    sequence = lapinsus_sequence_label(report["report_number"])
    if request.method == "POST":
        data = lapinhar_form_data()
        submitted_sequence = lapinsus_sequence_label(request.form.get("report_number", ""))
        try:
            selected_report_date = date.fromisoformat(data["report_date"])
        except (TypeError, ValueError):
            selected_report_date = None
        backdated = bool(
            selected_report_date and report.get("report_date")
            and selected_report_date < report["report_date"]
        )
        if backdated:
            backdated_reservation = fetch_one(
                """SELECT id,sequence_label FROM backdated_number_reservations
                   WHERE reservation_token=%s AND document_type='lapinsus'
                   AND report_id=%s AND report_date=%s AND created_by=%s""",
                (request.form.get("backdated_reservation_token", ""), report_id,
                 selected_report_date, session["user_id"]),
            )
            submitted_sequence = (
                backdated_reservation["sequence_label"] if backdated_reservation else ""
            )
        else:
            backdated_reservation = None
        data["report_number"] = compose_lapinsus_number(
            submitted_sequence or sequence, data["institution_code"], data["issue_code"], data["report_date"])
        attachments = collect_attachments()
        old_files = existing_report_files(report_id) if attachments else []
        old_rows = fetch_all("SELECT id FROM report_attachments WHERE report_id=%s", (report_id,)) if attachments else []
        available, message = lapinsus_number_available(data["report_number"], report_id=report_id)
        if backdated and not backdated_reservation:
            flash("Nomor tanggal mundur belum dibooking. Pilih kembali tanggal laporan.", "error")
        elif (data["acting_officer_type"] and not all((
                data["acting_officer_name"], data["acting_officer_rank"],
                data["acting_officer_nip"]))):
            flash("Nama, pangkat, dan NIP pejabat PLH/PLT wajib diisi.", "error")
        elif (data["report_acting_enabled"] and not all((
                data["report_acting_type"],
                data["report_acting_name"], data["report_acting_position"],
                data["report_acting_nip"]))):
            flash("Jenis, nama, jabatan, dan NIP PLT/PLH penandatangan wajib diisi.", "error")
        elif data["sipede_manual"] and not data["sipede_number"]:
            flash("Nomor Sipede manual wajib diisi.", "error")
        elif not all((data["report_number"], data["subject"], data["information"], data["category_id"], data["issue_code"])):
            flash("Nomor permasalahan, kategori, perihal, dan informasi wajib diisi.", "error")
        elif not available:
            flash(message, "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    """UPDATE lapinsus_reports SET report_number=%s,title=%s,report_date=%s,facts=%s,source_name=%s,
                    analysis=%s,recommendation=%s,recipient=%s,sender_name=%s,classification='RAHASIA',
                    category_id=%s,issue_code=%s,attachment=%s,organization=%s,creator_position=%s,
                    creator_name=%s,creator_rank_nip=%s,auth_position=%s,auth_name=%s,auth_rank_nip=%s,
                    use_scanned_signatures=%s,use_digital_stamp=%s,letter_signature_type=%s,
                    report_acting_type=%s,report_acting_name=%s,
                    report_acting_position=%s,report_acting_nip=%s,
                    letter_use_digital_stamp=%s,
                    acting_officer_type=%s,acting_officer_name=%s,
                    acting_officer_position=%s,acting_officer_rank=%s,acting_officer_nip=%s,
                    sipede_number=%s,sipede_status=%s,
                    information_spacing='1.5',sources_spacing='1.5',trends_spacing='1.5',suggestions_spacing='1.5',
                    status='selesai'
                    WHERE id=%s""",
                    (data["report_number"], data["subject"], data["report_date"] or None, data["information"],
                     data["sources"], data["trends"], data["suggestions"], data["recipient"], data["sender"],
                     data["category_id"], data["issue_code"], attachment_label(len(attachments)) if attachments else report["attachment"],
                     data["organization"], data["creator_position"], data["creator_name"], data["creator_rank_nip"],
                     data["auth_position"], data["auth_name"], data["auth_rank_nip"],
                     data["use_scanned_signatures"], data["use_digital_stamp"],
                     data["letter_signature_type"], data["report_acting_type"],
                     data["report_acting_name"], data["report_acting_position"],
                     data["report_acting_nip"], data["letter_use_digital_stamp"],
                     data["acting_officer_type"], data["acting_officer_name"],
                     data["acting_officer_position"], data["acting_officer_rank"],
                     data["acting_officer_nip"],
                     data["sipede_number"] if data["sipede_manual"] else (report["sipede_number"] or None),
                     "manual" if data["sipede_manual"] else ("belum" if report["sipede_status"] == "manual" else report["sipede_status"]),
                     report_id),
                )
                reservation_token = request.form.get("number_reservation_token", "")
                if reservation_token:
                    cursor.execute(
                        "UPDATE document_number_reservations SET status='used',report_id=%s,used_at=CURRENT_TIMESTAMP "
                        "WHERE reservation_token=%s AND document_type='lapinsus' AND created_by=%s AND status='reserved'",
                        (report_id, reservation_token, session["user_id"]),
                    )
            if backdated:
                release_original_document_number(report_id, "lapinsus")
                with get_db().cursor() as cursor:
                    cursor.execute(
                        "DELETE FROM backdated_number_reservations WHERE id=%s",
                        (backdated_reservation["id"],),
                    )
            if attachments:
                save_report_attachments(report_id, attachments)
                if old_rows:
                    with get_db().cursor() as cursor:
                        cursor.executemany("DELETE FROM report_attachments WHERE id=%s", [(row["id"],) for row in old_rows])
                remove_attachment_files(old_files, report_id)
            flash("LAPINSUS berhasil diperbarui.", "success")
            after_save = request.form.get("after_save", "")
            if after_save == "sipede" and data["sipede_manual"]:
                after_save = ""
            if after_save in {"docx", "print", "pdf", "sipede"}:
                return redirect(url_for("edit_lapinsus", report_id=report_id, after_save=after_save))
            return redirect(url_for("lapinsus"))
        report = accessible_lapinsus(report_id)
    reservation = {"reservation_token": "", "sequence_number": sequence}
    return render_template("lapinhar_editor.html", **lapinsus_editor_context(report, reservation))


@app.get("/lapinsus/<int:report_id>/attachments/<path:filename>")
@login_required
def lapinsus_attachment_file(report_id, filename):
    if accessible_lapinsus(report_id) is None or Path(filename).name != filename:
        abort(404)
    row = next((item for item in fetch_all("SELECT image_path FROM report_attachments WHERE report_id=%s", (report_id,))
                if Path(item["image_path"]).name == filename), None)
    if row is None:
        abort(404)
    return send_from_directory(UPLOAD_DIR / str(report_id), filename)


@app.post("/lapinsus/<int:report_id>/delete")
@login_required
def delete_lapinsus(report_id):
    report = accessible_lapinsus(report_id)
    if report is None:
        flash("LAPINSUS tidak ditemukan atau tidak dapat Anda akses.", "error")
        return redirect(url_for("lapinsus"))
    files = existing_report_files(report_id)
    with get_db().cursor() as cursor:
        cursor.execute("DELETE FROM report_attachments WHERE report_id=%s", (report_id,))
        cursor.execute("DELETE FROM lapinsus_reports WHERE id=%s", (report_id,))
    remove_attachment_files(files, report_id)
    flash(f"LAPINSUS {report['report_number']} berhasil dihapus.", "success")
    return redirect(url_for("lapinsus"))


@app.post("/lapinsus/<int:report_id>/upload-sipede")
@login_required
def upload_lapinsus_sipede(report_id):
    report = accessible_lapinsus(report_id)
    if report is None:
        return jsonify(message="LAPINSUS tidak ditemukan atau bukan milik Anda."), 403
    setting = fetch_one("""SELECT sipede_username,sipede_password_encrypted,
                        session_data_encrypted,connected_at
                        FROM sipede_user_settings WHERE user_id=%s""",
                        (session["user_id"],))
    if not setting:
        return jsonify(message="Konfigurasi Sipede belum tersedia. Isi username dan password pada Konfigurasi Integrasi."), 409
    if report.get("sipede_status") == "sudah":
        return jsonify(message="LAPINSUS sudah diupload ke Sipede."), 200
    if report.get("sipede_status") == "manual":
        return jsonify(message="Nomor SIPede diinput manual. Upload ke SIPede dinonaktifkan."), 400
    if not setting.get("session_data_encrypted") or not setting.get("connected_at"):
        return jsonify(message="Login SIPede diperlukan.", requires_sipede_login=True), 401
    session_data = get_sipede_session_data(session["user_id"])
    if not session_data:
        return jsonify(message="Login SIPede diperlukan.", requires_sipede_login=True), 401
    sipede_stage = "membuka form Surat Keluar"
    try:
        http_session = requests.Session()
        for cookie in session_data.get("cookies", []):
            http_session.cookies.set(
                cookie.get("name", ""), cookie.get("value", ""),
                domain=cookie.get("domain") or "sipede.kejaksaan.go.id",
                path=cookie.get("path") or "/",
            )
        response = http_session.get(
            SIPEDE_SURATKELUAR_CREATE_URL,
            headers={"Accept": "text/html,application/xhtml+xml", "Referer": f"{SIPEDE_BASE_URL}/"},
            timeout=60,
        )
        create_context = sipede_create_context_from_html(response.text, response.url)
        if "/login" in urlparse(response.url).path.lower() or not create_context:
            with get_db().cursor() as cursor:
                cursor.execute(
                    "UPDATE sipede_user_settings SET connected_at=NULL WHERE user_id=%s",
                    (session["user_id"],),
                )
            return jsonify(message="Sesi SIPede kedaluwarsa.", requires_sipede_login=True), 401
        kajari = fetch_one("SELECT full_name,position_name FROM signatories WHERE position_code='kajari'")
        kajari_name = kajari.get("full_name", "") if kajari else ""
        signatory_id = sipede_signatory_id(create_context, kajari_name)
        if not signatory_id:
            return jsonify(
                message=(f"Penandatangan {kajari_name or 'Kepala Kejaksaan Negeri'} tidak ditemukan "
                         "pada daftar SIPede. Periksa Konfigurasi Tanda Tangan."),
            ), 422
        issue_code = sipede_issue_code_value(create_context, report.get("issue_code"))
        if not issue_code:
            return jsonify(
                message=(f"Kode masalah {report.get('issue_code') or 'LAPINSUS'} tidak ditemukan "
                         "pada daftar Kode Masalah SIPede."),
            ), 422
        create_context["selected_kode_masalah"] = issue_code
        create_context["selected_penandatangan"] = str(signatory_id)

        ajax_headers = {
            "Accept": "*/*",
            "Referer": SIPEDE_SURATKELUAR_CREATE_URL,
            "X-Requested-With": "XMLHttpRequest",
        }
        csrf_token = create_context.get("csrf_token", "")
        if not csrf_token:
            return jsonify(message="Token form Surat Keluar SIPede tidak ditemukan."), 422
        temporary_number_key = f"sipede_upload_number_{report_id}"
        submitted_sipede_number = request.form.get("sipede_number", "").strip() if request.files.get("document") else ""
        temporary_sipede_number = str(session.get(temporary_number_key) or "").strip()
        if submitted_sipede_number:
            if (not temporary_sipede_number or submitted_sipede_number != temporary_sipede_number
                    or len(submitted_sipede_number) > 150
                    or re.search(r"[\r\n]", submitted_sipede_number)):
                return jsonify(message="Nomor SIPede sementara tidak valid. Ulangi proses upload."), 409
            sipede_number = temporary_sipede_number
        else:
            sipede_number = str(report.get("sipede_number") or "").strip()
        if not sipede_number or sipede_number == "-":
            sipede_stage = "memeriksa nomor otomatis"
            auto_number_response = http_session.get(
                f"{SIPEDE_BASE_URL}/suratkeluar/check-auto-number",
                params={"surat": "23"}, headers=ajax_headers, timeout=30,
            )
            auto_number_response.raise_for_status()
            sipede_stage = "mengambil nomor sequence surat"
            number_response = http_session.post(
                f"{SIPEDE_BASE_URL}/getnosurat/cekno",
                headers={**ajax_headers, "X-CSRF-TOKEN": csrf_token},
                data={
                    "_token": csrf_token, "jenis_surat": "23",
                    "tanggal": str(report.get("report_date") or date.today().isoformat()),
                    "penandatangan": str(signatory_id), "sifat_surat": "R",
                },
                timeout=30,
            )
            number_response.raise_for_status()
            sipede_number = number_response.text.strip().strip('"')
            if not sipede_number or "<html" in sipede_number.lower():
                return jsonify(message="Nomor sequence surat tidak diterima dari SIPede."), 502
            session[temporary_number_key] = sipede_number
            session.modified = True

        sipede_stage = "mengambil daftar tujuan disposisi"
        destinations_response = http_session.get(
            f"{SIPEDE_BASE_URL}/suratkeluar/get-master-tujuan-disposisi",
            headers={**ajax_headers, "X-CSRF-TOKEN": csrf_token},
            timeout=30,
        )
        destinations_response.raise_for_status()
        try:
            destinations_payload = destinations_response.json()
        except ValueError:
            return jsonify(message="Daftar tujuan SIPede tidak dapat dibaca."), 502
        destination_rows = destinations_payload.get("data", []) if isinstance(destinations_payload, dict) else []
        destinations = []
        for item in destination_rows:
            users = item.get("user") or []
            user_name = users[0].get("nama", "NO USER") if users and isinstance(users[0], dict) else "NO USER"
            destinations.append({
                "id": str(item.get("id_tujuan_penerusan", "")),
                "position": str(item.get("nama_jabatan", "")).strip(),
                "user": str(user_name).strip() or "NO USER",
            })
        destinations = [item for item in destinations if item["id"] and item["position"]]
        if not destinations:
            return jsonify(message="Daftar tujuan disposisi SIPede masih kosong."), 422

        selected_destinations = []
        if request.files.get("document"):
            try:
                selected_destinations = json.loads(request.form.get("destinations", "[]"))
            except (TypeError, ValueError):
                selected_destinations = []
            allowed_destination_ids = {item["id"] for item in destinations}
            selected_destinations = [str(item) for item in selected_destinations
                                     if str(item) in allowed_destination_ids]
            if not selected_destinations:
                return jsonify(message="Pilih minimal satu tujuan SIPede."), 422
            document = request.files["document"]
            document_bytes = document.read()
            if not document_bytes or len(document_bytes) > 18_000_000:
                return jsonify(message="Dokumen PDF LAPINSUS tidak valid atau terlalu besar."), 422
            organization = fetch_one("SELECT * FROM organization_settings WHERE id=1") or {
                "organization_name": "Kejaksaan Negeri Buleleng"
            }
            submit_data = [
                ("_token", csrf_token),
                ("suratmasuk", ""),
                ("nomor", sipede_number),
                ("tanggal", str(report.get("report_date") or date.today().isoformat())),
                ("jenis", "23"),
                ("sifat", "R"),
                ("kode_masalah", issue_code),
                ("tujuan", "Yth.\nKepala Kejaksaan Tinggi Bali\nDi - Denpasar"),
                ("dari", str((kajari or {}).get("position_name") or
                             f"Kepala {organization.get('organization_name', 'Kejaksaan Negeri Buleleng')}")),
                ("hal", str(report.get("title") or report.get("subject") or "LAPINSUS")),
                ("penandatangan", str(signatory_id)),
                ("idSurat", "125"),
                ("idSuratMasuk", ""),
                ("tujuan_surat", ",".join(selected_destinations)),
                ("submitModal", "ok"),
            ]
            tembusan_field = "tembusan[]" if "tembusan[]" in create_context.get("field_names", []) else "tembusan"
            for tembusan_id in sipede_tembusan_ids():
                submit_data.append((tembusan_field, tembusan_id))
            sipede_stage = "mengirim PDF dan data Surat Keluar"
            submit_response = http_session.post(
                create_context.get("form_action") or f"{SIPEDE_BASE_URL}/suratkeluar",
                headers={
                    "Accept": ("text/html,application/xhtml+xml,application/xml;q=0.9,"
                               "image/avif,image/webp,image/apng,*/*;q=0.8"),
                    "Accept-Language": "id,en-US;q=0.9,en;q=0.8",
                    "Cache-Control": "max-age=0",
                    "Origin": SIPEDE_BASE_URL,
                    "Priority": "u=0, i",
                    "Referer": SIPEDE_SURATKELUAR_CREATE_URL,
                    "Sec-CH-UA": '"Not;A=Brand";v="8", "Chromium";v="150", "Google Chrome";v="150"',
                    "Sec-CH-UA-Mobile": "?0",
                    "Sec-CH-UA-Platform": '"Windows"',
                    "Sec-Fetch-Dest": "document",
                    "Sec-Fetch-Mode": "navigate",
                    "Sec-Fetch-Site": "same-origin",
                    "Sec-Fetch-User": "?1",
                    "Upgrade-Insecure-Requests": "1",
                    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                                   "AppleWebKit/537.36 (KHTML, like Gecko) "
                                   "Chrome/150.0.0.0 Safari/537.36"),
                },
                data=submit_data,
                files={"upload": (document.filename or "LAPINSUS.pdf", document_bytes, "application/pdf")},
                timeout=90,
                allow_redirects=True,
            )
            submit_response.raise_for_status()
            submit_soup = BeautifulSoup(submit_response.text, "html.parser")
            returned_form = submit_soup.select_one('#formCreateEdit, form[action$="/suratkeluar"]')
            error_node = submit_soup.select_one('.alert-danger, .invalid-feedback, .text-danger')
            if returned_form and urlparse(submit_response.url).path.rstrip("/").endswith("/create"):
                error_message = error_node.get_text(" ", strip=True) if error_node else "SIPede belum menerima surat."
                app.logger.warning(
                    "SIPede menolak upload LAPINSUS %s: %s",
                    report_id, error_message,
                )
                return jsonify(message=error_message), 422
            with get_db().cursor() as cursor:
                cursor.execute(
                    "UPDATE lapinsus_reports SET sipede_number=%s,sipede_status='sudah' WHERE id=%s",
                    (sipede_number, report_id),
                )
            session.pop(temporary_number_key, None)
            session.modified = True
            save_refreshed_sipede_session(session["user_id"], session_data, http_session, create_context)
            flash("LAPINSUS berhasil dikirim ke SIPede.", "success")
            return jsonify(
                message="LAPINSUS berhasil dikirim ke SIPede.",
                uploaded=True,
                sipede_number=sipede_number,
                redirect_url=url_for("lapinsus"),
            )

        save_refreshed_sipede_session(session["user_id"], session_data, http_session, create_context)
        return jsonify(
            message="Nomor sequence surat SIPede berhasil diperoleh.",
            sipede_connected=True,
            sipede_number=sipede_number,
            signatory_id=str(signatory_id),
            signatory_name=kajari_name,
            kode_masalah=issue_code,
            destinations=destinations,
        ), 202
    except requests.RequestException as exc:
        error_message = sipede_request_error_message(exc, sipede_stage)
        app.logger.warning("%s", error_message)
        response_status = getattr(getattr(exc, "response", None), "status_code", None)
        if (sipede_stage == "membuka form Surat Keluar"
                or response_status in {401, 403, 419}):
            with get_db().cursor() as cursor:
                cursor.execute(
                    "UPDATE sipede_user_settings SET connected_at=NULL WHERE user_id=%s",
                    (session["user_id"],),
                )
            return jsonify(
                message="Sesi SIPede tidak aktif. Silakan login kembali.",
                requires_sipede_login=True,
            ), 401
        return jsonify(message=error_message), 502
    except Exception as exc:
        app.logger.exception("Upload LAPINSUS ke SIPede gagal")
        return jsonify(message=f"Upload LAPINSUS ke SIPede gagal: {exc}"), 500


def rich_text_blocks(content):
    if "<" not in content:
        return [{"text": line.strip(), "prefix": "", "align": "justify", "indent": 0,
                 "runs": [{"text": line.strip(), "bold": False, "italic": False, "underline": False}]}
                for line in content.splitlines() if line.strip()]
    soup = BeautifulSoup(content, "html.parser")
    blocks = []
    ordered_counters = {}

    def collect_runs(node, bold=False, italic=False, underline=False):
        runs = []
        for child in node.children:
            if isinstance(child, NavigableString):
                if str(child):
                    runs.append({"text": str(child), "bold": bold, "italic": italic, "underline": underline})
            elif isinstance(child, Tag) and child.name not in {"ol", "ul"}:
                runs.extend(collect_runs(child, bold or child.name in {"b", "strong"},
                                         italic or child.name in {"i", "em"},
                                         underline or child.name == "u"))
        return runs

    for element in soup.find_all(["p", "li"]):
        runs = collect_runs(element)
        text = "".join(run["text"] for run in runs).strip()
        if not text:
            continue
        classes = element.get("class", [])
        align = next((value.replace("ql-align-", "") for value in classes if value.startswith("ql-align-")), "justify")
        inline_style = element.get("style", "").replace(" ", "").lower()
        if "text-align:" in inline_style:
            align = inline_style.split("text-align:", 1)[1].split(";", 1)[0]
        indent_class = next((value for value in classes if value.startswith("ql-indent-")), "")
        indent = int(indent_class.rsplit("-", 1)[-1]) if indent_class else 0
        if "margin-left:" in inline_style:
            margin = inline_style.split("margin-left:", 1)[1].split(";", 1)[0]
            try:
                indent = max(indent, round(float(margin.replace("px", "").replace("em", "")) / (40 if "px" in margin else 1)))
            except ValueError:
                pass
        prefix = ""
        if element.name == "li":
            parent = element.parent
            list_type = element.get("data-list")
            if list_type == "ordered" or (not list_type and parent and parent.name == "ol"):
                key = id(parent)
                ordered_counters[key] = ordered_counters.get(key, 0) + 1
                prefix = f"{ordered_counters[key]}. "
            else:
                prefix = "• "
        blocks.append({"text": text, "prefix": prefix, "align": align, "indent": indent, "runs": runs})
    return blocks


def add_docx_section(document, title, content, spacing="1.15"):
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    heading.add_run(title).bold = True
    alignments = {"center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                  "justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "left": WD_ALIGN_PARAGRAPH.LEFT}
    for block in rich_text_blocks(content):
        paragraph = document.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Cm(.7 + block["indent"] * .7)
        paragraph.paragraph_format.first_line_indent = Cm(-.7 if block["prefix"] else 0)
        paragraph.paragraph_format.alignment = alignments.get(block["align"], WD_ALIGN_PARAGRAPH.JUSTIFY)
        paragraph.paragraph_format.line_spacing = float(spacing or 1.15)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        if block["prefix"]: paragraph.add_run(block["prefix"]).bold = True
        for item in block["runs"]:
            run = paragraph.add_run(item["text"])
            run.bold, run.italic, run.underline = item["bold"], item["italic"], item["underline"]
            run.font.name, run.font.size = "Times New Roman", Pt(12)


def add_docx_label(document, label, value):
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width, table.columns[1].width, table.columns[2].width = Cm(3.2), Cm(.5), Cm(13)
    cells = table.rows[0].cells
    cells[0].text, cells[1].text, cells[2].text = label, ":", value
    for cell in cells:
        cell.vertical_alignment = 1
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def lapinhar_export_basename(report_number, report_date, subject):
    """Return a safe RLIH/RLIK export filename."""
    number_text = str(report_number or "")
    is_lapinsus = number_text.startswith("R-LIK-")
    sequence = lapinsus_sequence_label(number_text) if is_lapinsus else report_sequence_label(number_text)
    prefix = "RLIK" if is_lapinsus else "RLIH"
    try:
        selected_date = date.fromisoformat(str(report_date or ""))
        months = ("Januari", "Februari", "Maret", "April", "Mei", "Juni",
                  "Juli", "Agustus", "September", "Oktober", "November", "Desember")
        date_label = f"{selected_date.day} {months[selected_date.month - 1]} {selected_date.year}"
    except ValueError:
        date_label = ""
    clean_subject = BeautifulSoup(str(subject or ""), "html.parser").get_text(" ", strip=True)
    basename = f"{prefix}-{sequence or 'Draft'} {date_label} {clean_subject}".strip()
    basename = re.sub(r'[<>:"/\\|?*\x00-\x1f]', " ", basename)
    basename = re.sub(r"\s+", " ", basename).strip(" .")
    return basename[:220].rstrip(" .") or f"{prefix}-Draft"


def inteliz_credential_cipher():
    secret = str(app.config["SECRET_KEY"]).encode("utf-8")
    key = base64.urlsafe_b64encode(hashlib.sha256(secret + b":inteliz-credentials:v1").digest())
    return Fernet(key)


def get_inteliz_credentials(user_id):
    """Used by the upcoming automation without exposing the password to templates."""
    row = fetch_one(
        "SELECT inteliz_username, inteliz_password_encrypted FROM inteliz_user_settings WHERE user_id=%s",
        (user_id,),
    )
    if not row:
        return None
    try:
        password = inteliz_credential_cipher().decrypt(
            row["inteliz_password_encrypted"].encode("ascii")
        ).decode("utf-8")
    except (InvalidToken, UnicodeError, ValueError):
        app.logger.error("Kredensial Inteliz pengguna %s tidak dapat didekripsi.", user_id)
        return None
    return {"username": row["inteliz_username"], "password": password}


def get_inteliz_session_data(user_id):
    row = fetch_one(
        "SELECT session_data_encrypted FROM inteliz_user_settings WHERE user_id=%s", (user_id,)
    )
    if not row or not row["session_data_encrypted"]:
        return None
    try:
        decrypted = inteliz_credential_cipher().decrypt(
            row["session_data_encrypted"].encode("ascii")
        ).decode("utf-8")
        return json.loads(decrypted)
    except (InvalidToken, UnicodeError, ValueError, json.JSONDecodeError):
        app.logger.error("Sesi Inteliz pengguna %s tidak dapat didekripsi.", user_id)
        return None


def html_form_value(soup, field_name):
    field = soup.select_one(f'[name="{field_name}"]')
    if field is None:
        return ""
    if field.name == "select":
        selected = field.select_one("option[selected]") or field.select_one("option[value]")
        return selected.get("value", "") if selected else ""
    return field.get("value", "")


def inteliz_signatory_value(report, soup, saved_options=None):
    options = soup.select('[name="penandatangan"] option[value]')
    valid_options = [(option.get("value", "").strip(), option.get_text(" ", strip=True))
                     for option in options if option.get("value", "").strip()]
    if not valid_options:
        valid_options = [(str(item.get("value", "")).strip(), str(item.get("label", "")).strip())
                         for item in (saved_options or []) if str(item.get("value", "")).strip()]

    # Inteliz memakai pejabat autentikasi/Kasi Intelijen sebagai penandatangan,
    # bukan Kasubsi yang menyusun laporan. NIP pada database dapat mengandung spasi.
    for source in (report.get("auth_rank_nip"), report.get("creator_rank_nip")):
        digits = re.sub(r"\D", "", source or "")
        nip = digits[-18:] if len(digits) >= 18 else ""
        if nip and (not valid_options or any(value == nip for value, _ in valid_options)):
            return nip

    signer_name = (report.get("auth_name") or report.get("creator_name") or "").casefold()
    for value, label in valid_options:
        if signer_name and signer_name in label.casefold():
            return value
    return ""


def save_refreshed_inteliz_session(user_id, session_data, http_session, csrf_token, create_key="lapinhar_create"):
    session_data["cookies"] = [
        {"name": cookie.name, "value": cookie.value, "domain": cookie.domain,
         "path": cookie.path or "/", "expires": cookie.expires or -1,
         "secure": bool(cookie.secure), "httpOnly": bool(cookie.has_nonstandard_attr("HttpOnly"))}
        for cookie in http_session.cookies
    ]
    session_data.setdefault(create_key, {})["csrf_token"] = csrf_token
    session_data["captured_at"] = int(time.time())
    encrypted = inteliz_credential_cipher().encrypt(
        json.dumps(session_data, ensure_ascii=False).encode("utf-8")
    ).decode("ascii")
    with get_db().cursor() as cursor:
        cursor.execute(
            "UPDATE inteliz_user_settings SET session_data_encrypted=%s WHERE user_id=%s",
            (encrypted, user_id),
        )


def extract_inteliz_csrf_token(soup, saved_context=None):
    token = html_form_value(soup, "_token")
    if token:
        return token
    meta = soup.select_one('meta[name="csrf-token"]')
    if meta and meta.get("content"):
        return meta.get("content", "").strip()
    if saved_context:
        return str(saved_context.get("csrf_token") or "").strip()
    return ""


@app.post("/lapinhar/<int:report_id>/sync-inteliz")
@login_required
def sync_lapinhar_inteliz(report_id):
    report = accessible_lapinhar(report_id)
    if report is None:
        flash("LAPINHAR tidak ditemukan atau tidak dapat Anda akses.", "error")
        return redirect(url_for("lapinhar"))
    if report.get("status") == "draft":
        flash("Lengkapi dan simpan LAPINHAR sebelum sinkronisasi Inteliz.", "warning")
        return redirect(url_for("lapinhar"))
    if report["inteliz_status"] == "sudah":
        flash("LAPINHAR tersebut sudah tersinkron ke Inteliz.", "success")
        return redirect(url_for("lapinhar"))

    session_data = get_inteliz_session_data(report["created_by"])
    if not session_data or not session_data.get("cookies"):
        flash("Sesi Inteliz pembuat laporan belum tersedia. Hubungkan Inteliz terlebih dahulu.", "error")
        if report["created_by"] == session["user_id"]:
            return redirect(url_for("lapinhar", connect_inteliz=1))
        return redirect(url_for("lapinhar"))

    http_session = requests.Session()
    for cookie in session_data["cookies"]:
        if cookie.get("expires", -1) not in (-1, None) and cookie["expires"] < time.time():
            continue
        http_session.cookies.set(cookie["name"], cookie["value"],
                                 domain=cookie.get("domain") or "inteliz.kejaksaan.go.id",
                                 path=cookie.get("path") or "/")
    headers = {
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "id,en-US;q=0.9,en;q=0.8",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                      "(KHTML, like Gecko) Chrome/150.0.0.0 Safari/537.36",
    }
    try:
        create_response = http_session.get(INTELIZ_LAPINHAR_CREATE_URL, headers=headers, timeout=45)
        if "/login" in create_response.url.lower():
            raise IntelizAuthenticationRequired("Sesi Inteliz sudah kedaluwarsa.")
        create_response.raise_for_status()
        soup = BeautifulSoup(create_response.text, "html.parser")
        saved_create = session_data.get("lapinhar_create", {})
        csrf_token = extract_inteliz_csrf_token(soup, saved_create)
        id_satker = html_form_value(soup, "id_satker") or saved_create.get("id_satker", "")
        signatory = inteliz_signatory_value(report, soup, saved_create.get("signatory_options"))
        if not csrf_token or not id_satker or not signatory:
            missing = []
            if not csrf_token:
                missing.append("token CSRF")
            if not id_satker:
                missing.append("ID satker")
            if not signatory:
                missing.append("penandatangan")
            app.logger.warning("Metadata create Inteliz laporan %s belum lengkap: %s; url=%s; fields=%s",
                               report_id, ", ".join(missing), create_response.url,
                               sorted({field.get("name") for field in soup.select("[name]") if field.get("name")}))
            raise RuntimeError(
                f"Data form Inteliz belum lengkap: {', '.join(missing)}. Cookie login tetap disimpan."
            )

        fields = {
            "_token": csrf_token,
            "id_satker": str(id_satker),
            "kategori_laporan": str(report["category_id"] or ""),
            "nomor": report["report_number"] or "",
            "tgl": report["report_date"].isoformat() if report["report_date"] else "",
            "info_yg_diperoleh": report["facts"] or "",
            "sumber_info": report["source_name"] or "",
            "trend_perkembangan": report["analysis"] or "",
            "pendapat_saran": report["recommendation"] or "",
            "penandatangan": signatory,
        }
        if not all(fields.values()):
            raise RuntimeError("Data LAPINHAR belum lengkap untuk dikirim ke Inteliz.")
        upload_response = http_session.post(
            "https://inteliz.kejaksaan.go.id/lapinhar",
            files={name: (None, value) for name, value in fields.items()},
            headers={**headers, "Referer": INTELIZ_LAPINHAR_CREATE_URL}, timeout=60,
        )
        if upload_response.status_code == 419 or "/login" in upload_response.url.lower():
            raise IntelizAuthenticationRequired("Sesi atau token Inteliz kedaluwarsa.")
        upload_response.raise_for_status()
        if urlparse(upload_response.url).path.rstrip("/") != "/lapinhar":
            error_soup = BeautifulSoup(upload_response.text, "html.parser")
            validation = " ".join(node.get_text(" ", strip=True) for node in
                                  error_soup.select(".invalid-feedback, .alert-danger, .text-danger") if node.get_text(strip=True))
            raise RuntimeError(validation or "Inteliz belum menerima data LAPINHAR.")
        save_refreshed_inteliz_session(report["created_by"], session_data, http_session, csrf_token, "lapinhar_create")
        with get_db().cursor() as cursor:
            cursor.execute("UPDATE lapinhar_reports SET inteliz_status='sudah' WHERE id=%s", (report_id,))
        flash("LAPINHAR berhasil disinkronkan ke Inteliz.", "success")
    except IntelizAuthenticationRequired as exc:
        with get_db().cursor() as cursor:
            cursor.execute(
                "UPDATE inteliz_user_settings SET connected_at=NULL WHERE user_id=%s",
                (report["created_by"],),
            )
        flash(f"{exc} Silakan selesaikan login Inteliz pada popup.", "error")
        if report["created_by"] == session["user_id"]:
            return redirect(url_for("lapinhar", connect_inteliz=1))
        flash("Pembuat LAPINHAR harus menghubungkan ulang akun Inteliz miliknya.", "error")
        return redirect(url_for("lapinhar"))
    except (requests.RequestException, RuntimeError) as exc:
        app.logger.warning("Sinkron Inteliz laporan %s gagal: %s", report_id, exc)
        flash(f"Sinkron Inteliz gagal: {exc}", "error")
    return redirect(url_for("lapinhar"))


@app.post("/lapinsus/<int:report_id>/sync-inteliz")
@login_required
def sync_lapinsus_inteliz(report_id):
    report = accessible_lapinsus(report_id)
    if report is None:
        flash("LAPINSUS tidak ditemukan atau tidak dapat Anda akses.", "error")
        return redirect(url_for("lapinsus"))
    if report.get("status") == "draft":
        flash("Lengkapi dan simpan LAPINSUS sebelum sinkronisasi Inteliz.", "warning")
        return redirect(url_for("lapinsus"))
    if report["inteliz_status"] == "sudah":
        flash("LAPINSUS tersebut sudah tersinkron ke Inteliz.", "success")
        return redirect(url_for("lapinsus"))

    session_data = get_inteliz_session_data(report["created_by"])
    if not session_data or not session_data.get("cookies"):
        flash("Sesi Inteliz pembuat laporan belum tersedia. Hubungkan Inteliz terlebih dahulu.", "error")
        if report["created_by"] == session["user_id"]:
            return redirect(url_for("lapinsus", connect_inteliz=1))
        return redirect(url_for("lapinsus"))

    http_session = requests.Session()
    for cookie in session_data["cookies"]:
        if cookie.get("expires", -1) not in (-1, None) and cookie["expires"] < time.time():
            continue
        http_session.cookies.set(
            cookie["name"], cookie["value"],
            domain=cookie.get("domain") or "inteliz.kejaksaan.go.id",
            path=cookie.get("path") or "/",
        )
    headers = {
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "id,en-US;q=0.9,en;q=0.8",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                      "(KHTML, like Gecko) Chrome/150.0.0.0 Safari/537.36",
    }
    try:
        create_response = http_session.get(INTELIZ_LAPINSUS_CREATE_URL, headers=headers, timeout=45)
        if "/login" in create_response.url.lower():
            raise IntelizAuthenticationRequired("Sesi Inteliz sudah kedaluwarsa.")
        create_response.raise_for_status()
        soup = BeautifulSoup(create_response.text, "html.parser")
        saved_create = session_data.get("lapinsus_create", {})
        csrf_token = extract_inteliz_csrf_token(soup, saved_create)
        id_satker = html_form_value(soup, "id_satker") or saved_create.get("id_satker", "")
        signatory = inteliz_signatory_value(report, soup, saved_create.get("signatory_options"))
        if not csrf_token or not id_satker or not signatory:
            missing = []
            if not csrf_token:
                missing.append("token CSRF")
            if not id_satker:
                missing.append("ID satker")
            if not signatory:
                missing.append("penandatangan")
            app.logger.warning("Metadata create Inteliz LAPINSUS %s belum lengkap: %s; url=%s; fields=%s",
                               report_id, ", ".join(missing), create_response.url,
                               sorted({field.get("name") for field in soup.select("[name]") if field.get("name")}))
            raise RuntimeError(
                f"Data form Inteliz belum lengkap: {', '.join(missing)}. Cookie login tetap disimpan."
            )

        fields = {
            "_token": csrf_token,
            "id_satker": str(id_satker),
            "kategori_laporan": str(report["category_id"] or ""),
            "nomor_surat": report["report_number"] or "",
            "tgl": report["report_date"].isoformat() if report["report_date"] else "",
            "informasi_yang_diperoleh": report["facts"] or "",
            "sumber_informasi": report["source_name"] or "",
            "perkembangan": report["analysis"] or "",
            "saran_tindak": report["recommendation"] or "",
            "penandatangan": signatory,
        }
        if not all(fields.values()):
            raise RuntimeError("Data LAPINSUS belum lengkap untuk dikirim ke Inteliz.")
        upload_response = http_session.post(
            "https://inteliz.kejaksaan.go.id/lapinsus",
            files={name: (None, value) for name, value in fields.items()},
            headers={**headers, "Referer": INTELIZ_LAPINSUS_CREATE_URL}, timeout=60,
        )
        if upload_response.status_code == 419 or "/login" in upload_response.url.lower():
            raise IntelizAuthenticationRequired("Sesi atau token Inteliz kedaluwarsa.")
        upload_response.raise_for_status()
        if urlparse(upload_response.url).path.rstrip("/") != "/lapinsus":
            error_soup = BeautifulSoup(upload_response.text, "html.parser")
            validation = " ".join(
                node.get_text(" ", strip=True)
                for node in error_soup.select(".invalid-feedback, .alert-danger, .text-danger")
                if node.get_text(strip=True)
            )
            raise RuntimeError(validation or "Inteliz belum menerima data LAPINSUS.")
        save_refreshed_inteliz_session(report["created_by"], session_data, http_session, csrf_token, "lapinsus_create")
        with get_db().cursor() as cursor:
            cursor.execute("UPDATE lapinsus_reports SET inteliz_status='sudah' WHERE id=%s", (report_id,))
        flash("LAPINSUS berhasil disinkronkan ke Inteliz.", "success")
    except IntelizAuthenticationRequired as exc:
        with get_db().cursor() as cursor:
            cursor.execute(
                "UPDATE inteliz_user_settings SET connected_at=NULL WHERE user_id=%s",
                (report["created_by"],),
            )
        flash(f"{exc} Silakan selesaikan login Inteliz pada popup.", "error")
        if report["created_by"] == session["user_id"]:
            return redirect(url_for("lapinsus", connect_inteliz=1))
        flash("Pembuat LAPINSUS harus menghubungkan ulang akun Inteliz miliknya.", "error")
        return redirect(url_for("lapinsus"))
    except (requests.RequestException, RuntimeError) as exc:
        app.logger.warning("Sinkron Inteliz LAPINSUS %s gagal: %s", report_id, exc)
        flash(f"Sinkron Inteliz gagal: {exc}", "error")
    return redirect(url_for("lapinsus"))


def read_inteliz_create_context(page):
    return page.evaluate("""() => {
        const form = document.querySelector('form[action*="lapinhar"]') || document.querySelector('form');
        const valueOf = (name) => {
            const field = document.querySelector(`[name="${name}"]`);
            return field ? field.value : '';
        };
        const optionsOf = (name) => Array.from(document.querySelectorAll(`[name="${name}"] option`))
            .map(option => ({ value: option.value, label: option.textContent.trim() }))
            .filter(option => option.value);
        return {
            csrf_token: valueOf('_token') || document.querySelector('meta[name="csrf-token"]')?.content || '',
            form_action: form?.action || 'https://inteliz.kejaksaan.go.id/lapinhar',
            form_method: (form?.method || 'post').toUpperCase(),
            id_satker: valueOf('id_satker'),
            category_options: optionsOf('kategori_laporan'),
            signatory_options: optionsOf('penandatangan'),
            field_names: form ? Array.from(form.elements).map(field => field.name).filter(Boolean) : []
        };
    }""")


def update_inteliz_auth(auth_id, **values):
    with INTELIZ_AUTH_LOCK:
        state = INTELIZ_AUTH_SESSIONS.get(auth_id)
        if state:
            state.update(values, updated_at=time.time())


def wait_inteliz_auth_input(auth_id, field, timeout=600):
    deadline = time.time() + timeout
    while time.time() < deadline:
        with INTELIZ_AUTH_LOCK:
            state = INTELIZ_AUTH_SESSIONS.get(auth_id)
            if not state or state.get("cancelled"):
                return None
            value = state.pop(field, None)
        if value:
            return value
        time.sleep(.25)
    return None


def visible_inteliz_otp_input(page):
    selectors = (
        'input[name="code"]', 'input[autocomplete="one-time-code"]', 'input[name="otp"]',
        'input[name="token"]', 'input[name="authenticator_code"]',
        'input[name="one_time_password"]', 'input[id*="otp" i]', 'input[id*="auth" i]',
    )
    for selector in selectors:
        locator = page.locator(selector).first
        if locator.count() and locator.is_visible():
            return locator
    return None


def persist_inteliz_browser_session(user_id, context, page, create_context=None):
    cookies = context.cookies(["https://inteliz.kejaksaan.go.id"])
    local_storage = page.evaluate(
        "Object.fromEntries(Array.from({length: localStorage.length}, (_, i) => "
        "[localStorage.key(i), localStorage.getItem(localStorage.key(i))]))"
    )
    payload = {
        "cookies": cookies,
        "local_storage": local_storage,
        "lapinhar_create": create_context or {},
        "captured_url": page.url,
        "captured_at": int(time.time()),
    }
    encrypted_session = inteliz_credential_cipher().encrypt(
        json.dumps(payload, ensure_ascii=False).encode("utf-8")
    ).decode("ascii")
    with app.app_context():
        with get_db().cursor() as cursor:
            cursor.execute(
                """UPDATE inteliz_user_settings SET session_data_encrypted=%s,
                   connected_at=CURRENT_TIMESTAMP WHERE user_id=%s""",
                (encrypted_session, user_id),
            )
    return len(cookies)


def run_inteliz_auth(auth_id, user_id, credentials):
    browser = None
    try:
        if not CHROME_EXECUTABLE.exists():
            raise RuntimeError("Google Chrome tidak ditemukan pada komputer server.")
        with sync_playwright() as playwright:
            browser = playwright.chromium.launch(
                executable_path=str(CHROME_EXECUTABLE),
                headless=os.getenv("INTELIZ_BROWSER_HEADLESS", "1") == "1",
                slow_mo=int(os.getenv("INTELIZ_BROWSER_SLOW_MO", "0")),
                args=browser_launch_args(browser_key="INTELIZ_BROWSER"),
            )
            context = browser.new_context(viewport={"width": 1280, "height": 900}, locale="id-ID")
            page = context.new_page()
            update_inteliz_auth(auth_id, status="loading", message="Membuka halaman Inteliz…")
            page.goto(INTELIZ_LOGIN_URL, wait_until="domcontentloaded", timeout=60000)
            page.locator("#username").fill(credentials["username"])
            page.locator("#password-field").fill(credentials["password"])

            authenticated = False
            for attempt in range(3):
                if page.locator("#username").count():
                    page.locator("#username").fill(credentials["username"])
                if page.locator("#password-field").count():
                    page.locator("#password-field").fill(credentials["password"])
                captcha = page.locator("#captcha-container img").first
                captcha.wait_for(state="visible", timeout=20000)
                captcha_data = base64.b64encode(captcha.screenshot(type="png")).decode("ascii")
                update_inteliz_auth(
                    auth_id, status="captcha", captcha=f"data:image/png;base64,{captcha_data}",
                    message="Masukkan kode CAPTCHA yang tampil.",
                )
                captcha_value = wait_inteliz_auth_input(auth_id, "captcha_input")
                if captcha_value is None:
                    raise RuntimeError("Sesi login dibatalkan atau kedaluwarsa.")
                page.locator('input[name="captcha"]').fill(captcha_value)
                page.locator("#kt_sign_in_submit").click()
                page.wait_for_timeout(2500)
                otp_input = visible_inteliz_otp_input(page)
                current_path = urlparse(page.url).path.rstrip("/")
                if otp_input is None and current_path == INTELIZ_2FA_PATH:
                    try:
                        page.locator('input[name="code"]').wait_for(state="visible", timeout=10000)
                    except Exception:
                        pass
                    otp_input = visible_inteliz_otp_input(page)
                if otp_input is not None:
                    break
                if "/login" not in page.url.lower() and current_path != INTELIZ_2FA_PATH:
                    authenticated = True
                    break
                update_inteliz_auth(auth_id, status="loading", captcha=None,
                                    message="CAPTCHA belum diterima, memuat CAPTCHA baru…")

            if not authenticated:
                otp_input = visible_inteliz_otp_input(page)
                if otp_input is None:
                    raise RuntimeError("Login belum berhasil. Periksa CAPTCHA atau kredensial Inteliz.")
                for _attempt in range(3):
                    update_inteliz_auth(auth_id, status="otp", captcha=None,
                                        message="Masukkan kode autentikator Inteliz.")
                    otp_value = wait_inteliz_auth_input(auth_id, "otp_input")
                    if otp_value is None:
                        raise RuntimeError("Sesi autentikator dibatalkan atau kedaluwarsa.")
                    otp_input.fill(otp_value)
                    otp_form = otp_input.locator("xpath=ancestor::form[1]")
                    submit = otp_form.locator('button[type="submit"], input[type="submit"]').first
                    if not submit.count():
                        submit = page.locator('button[type="submit"], input[type="submit"]').first
                    if not submit.count():
                        raise RuntimeError("Tombol verifikasi autentikator tidak ditemukan.")
                    submit.click()
                    try:
                        page.wait_for_function(
                            "() => !['/login', '/2fa/challenge'].includes(location.pathname.replace(/\\/$/, ''))",
                            timeout=30000,
                        )
                    except Exception:
                        page.wait_for_timeout(1000)
                    otp_input = visible_inteliz_otp_input(page)
                    current_path = urlparse(page.url).path.rstrip("/")
                    if (otp_input is None and "/login" not in page.url.lower()
                            and current_path != INTELIZ_2FA_PATH):
                        authenticated = True
                        break
                if not authenticated:
                    raise RuntimeError("Kode autentikator belum diterima oleh Inteliz.")

            cookie_count = persist_inteliz_browser_session(user_id, context, page)
            if not cookie_count:
                raise RuntimeError("Dashboard terbuka, tetapi cookie Inteliz tidak ditemukan.")
            update_inteliz_auth(auth_id, status="loading", captcha=None,
                                message="Cookie tersimpan. Membaca formulir Inteliz…")
            try:
                page.goto(INTELIZ_LAPINHAR_CREATE_URL, wait_until="domcontentloaded", timeout=60000)
                page.wait_for_timeout(1200)
                lapinhar_context = read_inteliz_create_context(page) if urlparse(page.url).path.rstrip("/") == "/lapinhar/create" else None

                lapinsus_context = None
                page.goto(INTELIZ_LAPINSUS_CREATE_URL, wait_until="domcontentloaded", timeout=60000)
                page.wait_for_timeout(1200)
                if urlparse(page.url).path.rstrip("/") == "/lapinsus/create":
                    lapinsus_context = read_inteliz_create_context(page)

                cookie_payload = {
                    "cookies": context.cookies(["https://inteliz.kejaksaan.go.id"]),
                    "local_storage": page.evaluate(
                        "Object.fromEntries(Array.from({length: localStorage.length}, (_, i) => "
                        "[localStorage.key(i), localStorage.getItem(localStorage.key(i))]))"
                    ),
                    "lapinhar_create": lapinhar_context or {},
                    "lapinsus_create": lapinsus_context or {},
                    "captured_url": page.url,
                    "captured_at": int(time.time()),
                }
                encrypted = inteliz_credential_cipher().encrypt(
                    json.dumps(cookie_payload, ensure_ascii=False).encode("utf-8")
                ).decode("ascii")
                with get_db().cursor() as cursor:
                    cursor.execute(
                        "UPDATE inteliz_user_settings SET session_data_encrypted=%s WHERE user_id=%s",
                        (encrypted, user_id),
                    )
            except Exception as form_exc:
                app.logger.warning("Cookie Inteliz tersimpan, metadata form belum terbaca lengkap: %s", form_exc)
            update_inteliz_auth(auth_id, status="success", captcha=None,
                                message="Inteliz berhasil terhubung. Cookie telah disimpan terenkripsi.")
    except Exception as exc:
        app.logger.warning("Login Inteliz pengguna %s gagal: %s", user_id, exc)
        update_inteliz_auth(auth_id, status="error", captcha=None, message=str(exc))
    finally:
        if browser:
            try:
                browser.close()
            except Exception:
                pass


def get_sipede_credentials(user_id):
    row = fetch_one(
        "SELECT sipede_username, sipede_password_encrypted FROM sipede_user_settings WHERE user_id=%s",
        (user_id,),
    )
    if not row or not row.get("sipede_password_encrypted"):
        return None
    try:
        password = inteliz_credential_cipher().decrypt(
            row["sipede_password_encrypted"].encode("ascii")
        ).decode("utf-8")
    except Exception:
        app.logger.error("Kredensial SIPede pengguna %s tidak dapat didekripsi.", user_id)
        return None
    return {"username": row["sipede_username"], "password": password}


def get_sipede_session_data(user_id):
    row = fetch_one(
        "SELECT session_data_encrypted FROM sipede_user_settings WHERE user_id=%s",
        (user_id,),
    )
    if not row or not row.get("session_data_encrypted"):
        return None
    try:
        decrypted = inteliz_credential_cipher().decrypt(
            row["session_data_encrypted"].encode("ascii")
        )
        return json.loads(decrypted.decode("utf-8"))
    except Exception:
        app.logger.error("Sesi SIPede pengguna %s tidak dapat didekripsi.", user_id)
        return None


def sipede_create_context_from_html(html, response_url):
    soup = BeautifulSoup(html, "html.parser")
    form = soup.select_one('#formCreateEdit, form[action$="/suratkeluar"]')
    if not form:
        return None
    token = form.select_one('input[name="_token"]') or soup.select_one('meta[name="csrf-token"]')
    def options_of(name):
        return [
            {"value": option.get("value", ""), "label": option.get_text(" ", strip=True)}
            for option in form.select(f'[name="{name}"] option[value]')
            if option.get("value")
        ]

    return {
        "csrf_token": (token.get("value") or token.get("content") or "") if token else "",
        "form_action": form.get("action") or f"{SIPEDE_BASE_URL}/suratkeluar",
        "form_method": (form.get("method") or "post").upper(),
        "form_enctype": form.get("enctype") or "multipart/form-data",
        "id_surat": (form.select_one('input[name="idSurat"]') or {}).get("value", "125"),
        "jenis_options": options_of("jenis"),
        "sifat_options": options_of("sifat"),
        "kode_masalah_options": options_of("kode_masalah"),
        "penandatangan_options": options_of("penandatangan"),
        "field_names": [field.get("name") for field in form.select("input[name],select[name],textarea[name]")],
        "captured_url": response_url,
    }


def sipede_signatory_id(create_context, configured_name):
    normalized_target = re.sub(r"[^A-Z0-9]", "", (configured_name or "").upper())
    if not normalized_target:
        return None
    for option in create_context.get("penandatangan_options", []):
        normalized_label = re.sub(r"[^A-Z0-9]", "", option.get("label", "").upper())
        if normalized_target in normalized_label:
            return option.get("value")
    return None


def sipede_issue_code_value(create_context, report_issue_code):
    target = (report_issue_code or "").strip()
    if not target:
        return None
    for option in create_context.get("kode_masalah_options", []):
        if option.get("value", "").strip().casefold() == target.casefold():
            return option.get("value", "").strip()
    normalized_target = re.sub(r"[^A-Z0-9]", "", target.upper())
    for option in create_context.get("kode_masalah_options", []):
        value = option.get("value", "").strip()
        label = option.get("label", "")
        normalized_value = re.sub(r"[^A-Z0-9]", "", value.upper())
        normalized_label = re.sub(r"[^A-Z0-9]", "", label.upper())
        if normalized_target == normalized_value or normalized_target in normalized_label:
            return value
    return None


def sipede_tembusan_text(report):
    lines = [
        "Yth. Wakil Kepala Kejaksaan Tinggi Bali;",
        "Yth. Asisten Bidang Intelijen Kejaksaan Tinggi Bali;",
        "Yth. Asisten Bidang Pengawasan Kejaksaan Tinggi Bali;",
    ]
    if report.get("acting_officer_type"):
        lines.append("Yth. Kepala Kejaksaan Negeri Buleleng (Sebagai Laporan);")
    lines.append("Arsip.")
    return "\n".join(lines)


def sipede_tembusan_ids():
    return ["603", "3309", "1230"]


def sipede_request_error_message(error, stage):
    response = getattr(error, "response", None)
    if response is None:
        detail = str(error).strip()
        if isinstance(error, requests.Timeout):
            detail = "waktu tunggu koneksi habis"
        elif isinstance(error, requests.ConnectionError):
            detail = "koneksi dari server terputus atau ditolak"
        return f"SIPede gagal saat {stage}: {detail or 'koneksi tidak tersedia'}."

    status = response.status_code
    detail = ""
    content_type = (response.headers.get("Content-Type") or "").lower()
    if "html" in content_type:
        soup = BeautifulSoup(response.text or "", "html.parser")
        error_node = soup.select_one(
            ".alert-danger, .alert-error, .invalid-feedback, .text-danger, "
            ".swal2-html-container, main"
        )
        if error_node:
            detail = error_node.get_text(" ", strip=True)
        elif soup.title:
            detail = soup.title.get_text(" ", strip=True)
    elif response.text:
        detail = response.text.strip()

    detail = re.sub(r"\s+", " ", detail)[:350]
    message = f"SIPede gagal saat {stage} (HTTP {status})"
    if detail:
        message += f": {detail}"
    return message + "."


def save_refreshed_sipede_session(user_id, session_data, http_session, create_context):
    session_data["cookies"] = [
        {
            "name": cookie.name,
            "value": cookie.value,
            "domain": cookie.domain,
            "path": cookie.path,
            "secure": cookie.secure,
        }
        for cookie in http_session.cookies
    ]
    session_data["suratkeluar_create"] = create_context or {}
    session_data["captured_url"] = (create_context or {}).get("captured_url", SIPEDE_SURATKELUAR_CREATE_URL)
    session_data["captured_at"] = int(time.time())
    encrypted = inteliz_credential_cipher().encrypt(
        json.dumps(session_data, ensure_ascii=False).encode("utf-8")
    ).decode("ascii")
    with get_db().cursor() as cursor:
        cursor.execute(
            """UPDATE sipede_user_settings SET session_data_encrypted=%s,
               connected_at=CURRENT_TIMESTAMP WHERE user_id=%s""",
            (encrypted, user_id),
        )


def update_sipede_auth(auth_id, **values):
    with SIPEDE_AUTH_LOCK:
        state = SIPEDE_AUTH_SESSIONS.get(auth_id)
        if state:
            state.update(values, updated_at=time.time())


def wait_sipede_auth_input(auth_id, field, timeout=600):
    deadline = time.time() + timeout
    while time.time() < deadline:
        with SIPEDE_AUTH_LOCK:
            state = SIPEDE_AUTH_SESSIONS.get(auth_id)
            if not state or state.get("cancelled"):
                return None
            value = state.pop(field, None)
        if value:
            return value
        time.sleep(.25)
    return None


def read_sipede_create_context(page):
    return page.evaluate("""() => {
        const form = document.querySelector('#formCreateEdit') ||
            document.querySelector('form[action$="/suratkeluar"]');
        const optionsOf = (name) => Array.from(document.querySelectorAll(`[name="${name}"] option`))
            .map(option => ({ value: option.value, label: option.textContent.trim() }))
            .filter(option => option.value);
        const valueOf = (name) => document.querySelector(`[name="${name}"]`)?.value || '';
        return {
            csrf_token: valueOf('_token') || document.querySelector('meta[name="csrf-token"]')?.content || '',
            form_action: form?.action || 'https://sipede.kejaksaan.go.id/suratkeluar',
            form_method: (form?.method || 'post').toUpperCase(),
            form_enctype: form?.enctype || 'multipart/form-data',
            id_surat: valueOf('idSurat') || '125',
            jenis_options: optionsOf('jenis'),
            sifat_options: optionsOf('sifat'),
            kode_masalah_options: optionsOf('kode_masalah'),
            penandatangan_options: optionsOf('penandatangan'),
            field_names: form ? Array.from(form.elements).map(field => field.name).filter(Boolean) : []
        };
    }""")


def persist_sipede_browser_session(user_id, context, page, create_context=None):
    cookies = context.cookies([SIPEDE_BASE_URL])
    local_storage = page.evaluate(
        "Object.fromEntries(Array.from({length: localStorage.length}, (_, i) => "
        "[localStorage.key(i), localStorage.getItem(localStorage.key(i))]))"
    )
    payload = {
        "cookies": cookies,
        "local_storage": local_storage,
        "suratkeluar_create": create_context or {},
        "captured_url": page.url,
        "captured_at": int(time.time()),
    }
    encrypted = inteliz_credential_cipher().encrypt(
        json.dumps(payload, ensure_ascii=False).encode("utf-8")
    ).decode("ascii")
    with app.app_context():
        with get_db().cursor() as cursor:
            cursor.execute(
                """UPDATE sipede_user_settings
                   SET session_data_encrypted=%s, connected_at=CURRENT_TIMESTAMP
                   WHERE user_id=%s""",
                (encrypted, user_id),
            )
    return len(cookies)


def run_sipede_auth(auth_id, user_id, credentials):
    browser = None
    try:
        if not CHROME_EXECUTABLE.exists():
            raise RuntimeError("Google Chrome tidak ditemukan pada komputer server.")
        with sync_playwright() as playwright:
            def open_sipede_browser(use_proxy=True):
                opened_browser = playwright.chromium.launch(
                    executable_path=str(CHROME_EXECUTABLE),
                    headless=os.getenv("SIPEDE_BROWSER_HEADLESS", "1") == "1",
                    slow_mo=int(os.getenv("SIPEDE_BROWSER_SLOW_MO", "250")),
                    args=browser_launch_args(
                        ["--incognito"], browser_key="SIPEDE_BROWSER", use_proxy=use_proxy
                    ),
                )
                opened_context = opened_browser.new_context(
                    viewport={"width": 1280, "height": 900}, locale="id-ID"
                )
                return opened_browser, opened_context, opened_context.new_page()

            browser, context, page = open_sipede_browser(use_proxy=True)
            update_sipede_auth(auth_id, status="loading", message="Membuka halaman SIPede…")
            try:
                page.goto(SIPEDE_LOGIN_URL, wait_until="domcontentloaded", timeout=60000)
            except Exception as goto_exc:
                goto_message = str(goto_exc)
                if ("ERR_PROXY_CONNECTION_FAILED" not in goto_message
                        and "ERR_NETWORK_ACCESS_DENIED" not in goto_message):
                    raise
                app.logger.warning("Login SIPede via proxy gagal, mencoba ulang tanpa proxy: %s", goto_exc)
                update_sipede_auth(
                    auth_id,
                    status="loading",
                    message="Proxy SIPede gagal. Mencoba ulang koneksi langsung...",
                )
                try:
                    browser.close()
                except Exception:
                    pass
                browser, context, page = open_sipede_browser(use_proxy=False)
                page.goto(SIPEDE_LOGIN_URL, wait_until="domcontentloaded", timeout=60000)

            authenticated = False
            for _attempt in range(3):
                page.locator('#username, input[name="username"]').first.fill(credentials["username"])
                page.locator('input[name="password"]').first.fill(credentials["password"])
                captcha_image = page.locator('.captcha img, img[src*="captcha" i]').first
                captcha_image.wait_for(state="visible", timeout=30000)
                captcha_data = base64.b64encode(captcha_image.screenshot(type="png")).decode("ascii")
                update_sipede_auth(
                    auth_id,
                    status="captcha",
                    captcha=f"data:image/png;base64,{captcha_data}",
                    message="Masukkan kode CAPTCHA SIPede yang tampil.",
                )
                captcha_value = wait_sipede_auth_input(auth_id, "captcha_input")
                if captcha_value is None:
                    raise RuntimeError("Sesi login SIPede dibatalkan atau kedaluwarsa.")
                page.locator('#captcha, input[name="captcha"]').first.fill(captcha_value)
                page.locator('form[action*="login-post"] button[type="submit"], button[type="submit"]').first.click()
                try:
                    page.wait_for_function(
                        "() => !location.pathname.replace(/\\/$/, '').endsWith('/login')",
                        timeout=30000,
                    )
                except Exception:
                    page.wait_for_timeout(1200)
                if "/login" not in urlparse(page.url).path.lower():
                    authenticated = True
                    break
                update_sipede_auth(
                    auth_id,
                    status="loading",
                    captcha=None,
                    message="Login belum diterima. Memuat CAPTCHA baru…",
                )

            if not authenticated:
                raise RuntimeError("Login SIPede belum berhasil. Periksa CAPTCHA, username, atau password.")
            cookie_count = persist_sipede_browser_session(user_id, context, page)
            if not cookie_count:
                raise RuntimeError("Halaman SIPede terbuka, tetapi cookie sesi tidak ditemukan.")
            update_sipede_auth(
                auth_id,
                status="loading",
                captcha=None,
                message="Cookie tersimpan. Membuka formulir Surat Keluar SIPede…",
            )
            create_context = None
            try:
                page.goto(SIPEDE_SURATKELUAR_CREATE_URL, wait_until="domcontentloaded", timeout=60000)
                page.locator('#formCreateEdit, form[action$="/suratkeluar"]').first.wait_for(
                    state="attached", timeout=30000
                )
                create_context = read_sipede_create_context(page)
                persist_sipede_browser_session(user_id, context, page, create_context)
            except Exception as form_exc:
                app.logger.warning("Cookie SIPede tersimpan, metadata form Surat Keluar belum terbaca: %s", form_exc)
            update_sipede_auth(
                auth_id,
                status="success",
                captcha=None,
                message=("SIPede berhasil terhubung. Sesi, cookie, token, dan form Surat Keluar "
                         "telah disimpan terenkripsi."),
            )
    except Exception as exc:
        app.logger.warning("Login SIPede pengguna %s gagal: %s", user_id, exc)
        update_sipede_auth(auth_id, status="error", captcha=None, message=str(exc))
    finally:
        if browser:
            try:
                browser.close()
            except Exception:
                pass


@app.post("/lapinhar/export/<file_type>")
@login_required
def export_lapinhar(file_type):
    data = lapinhar_form_data()
    document_kind = request.form.get("document_kind", "lapinhar")
    is_lapinsus = document_kind == "lapinsus"
    document_label = "LAPINSUS" if is_lapinsus else "LAPINHAR"
    document_heading = "LAPORAN INFORMASI KHUSUS" if is_lapinsus else "LAPORAN INFORMASI HARIAN"
    attachments = collect_attachments()
    data["attachment"] = attachment_label(len(attachments))
    export_name = lapinhar_export_basename(data["report_number"], data["report_date"], data["subject"])
    if file_type == "docx":
        output = BytesIO()
        document = Document()
        section = document.sections[0]
        section.page_width, section.page_height = Cm(21.59), Cm(33.02)
        section.top_margin, section.bottom_margin = Cm(2), Cm(2)
        section.left_margin, section.right_margin = Cm(2), Cm(1.5)
        normal = document.styles["Normal"]
        normal.font.name, normal.font.size = "Times New Roman", Pt(12)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing = 1.5
        organization = data["organization"] or "KEJAKSAAN NEGERI BULELENG"
        cover_org = document.add_paragraph()
        cover_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_org.add_run(organization).bold = True
        logo = document.add_picture(str(STATIC_IMG_DIR / "logo-kejaksaan.jpeg"), width=Cm(4.2))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_title = document.add_paragraph()
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_title.paragraph_format.space_before = Pt(16)
        cover_title.add_run(document_label).bold = True
        cover_subject = document.add_paragraph()
        cover_subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_subject.add_run(data["subject"]).bold = True
        secret = document.add_picture(str(STATIC_IMG_DIR / "rahasia.jpeg"), width=Cm(17.5), height=Cm(14.5))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        memo_section = document.add_section(WD_SECTION.NEW_PAGE)
        memo_section.page_width, memo_section.page_height = Cm(21.59), Cm(33.02)
        memo_section.top_margin, memo_section.bottom_margin = Cm(2), Cm(2)
        memo_section.left_margin, memo_section.right_margin = Cm(2), Cm(1.5)
        office = document.add_paragraph(organization)
        office.alignment = WD_ALIGN_PARAGRAPH.CENTER
        office.runs[0].bold, office.runs[0].underline = True, True
        document.add_paragraph("\n\n")
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("NOTA  -  DINAS")
        run.bold, run.underline = True, True
        document.add_paragraph()
        for label, value in (("K E P A D A", data["recipient"]), ("D A R I", data["sender"]),
                             ("NOMOR", data["report_number"]), ("TANGGAL", data["report_date"]),
                             ("SIFAT", data["classification"]), ("LAMPIRAN", data["attachment"]),
                             ("PERIHAL", data["subject"])):
            add_docx_label(document, label, value)
        document.add_paragraph("\n")
        signer = document.add_paragraph()
        signer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        signer.paragraph_format.left_indent = Cm(9)
        signer.add_run(f"{data['auth_position']}\n{organization}\n\n\n\n")
        signer.add_run(data["auth_name"]).bold = True
        signer.add_run(f"\n{data['auth_rank_nip']}")

        second = document.add_section(WD_SECTION.NEW_PAGE)
        second.page_width, second.page_height = Cm(21.59), Cm(33.02)
        second.top_margin, second.bottom_margin = Cm(2), Cm(2)
        second.left_margin, second.right_margin = Cm(2), Cm(1.5)
        second.header.is_linked_to_previous = False
        second.footer.is_linked_to_previous = False
        report_header = second.header.paragraphs[0]
        report_header.text = "RAHASIA"
        report_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_header.runs[0].font.name, report_header.runs[0].font.size = "Times New Roman", Pt(10)
        report_footer = second.footer.paragraphs[0]
        report_footer.text = "RAHASIA"
        report_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_footer.runs[0].font.name, report_footer.runs[0].font.size = "Times New Roman", Pt(10)
        document.add_paragraph(f"{organization}\t{'L. IN.2' if is_lapinsus else 'L. IN.1'}")
        document.add_paragraph("Copy ke 1\nDari 1 copies")
        report_title = document.add_paragraph()
        report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_title.paragraph_format.line_spacing = 1.5
        rr = report_title.add_run(document_heading)
        rr.bold, rr.underline = True, True
        number = document.add_paragraph(f"NOMOR : {data['report_number']}")
        number.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if is_lapinsus:
            subject_line = document.add_paragraph(f"PERIHAL : {data['subject']}")
            subject_line.paragraph_format.space_after = Pt(12)
        add_docx_section(document, "I.    INFORMASI YANG DIPEROLEH", data["information"], data["information_spacing"])
        add_docx_section(document, "II.   SUMBER INFORMASI", data["sources"], data["sources_spacing"])
        add_docx_section(document, "III.  TREN PERKEMBANGAN / PERKIRAAN", data["trends"], data["trends_spacing"])
        add_docx_section(document, "IV.   PENDAPAT / SARAN", data["suggestions"], data["suggestions_spacing"])
        signatures = document.add_table(rows=1, cols=2)
        signature_cells = []
        if data["show_authentication"]:
            left = signatures.rows[0].cells[0]
            left.text = f"AUTENTIKASI:\n{data['auth_position']}\n{organization}\n\n\n\n\n{data['auth_name']}\n{data['auth_rank_nip']}"
            signature_cells.append(left)
            right = signatures.rows[0].cells[1]
        else:
            right = signatures.rows[0].cells[1]
        right.text = f"{data['city']}, {data['report_date']}\nYang Membuat Laporan\n{data['creator_position']}\n\n\n\n\n{data['creator_name']}\n{data['creator_rank_nip']}"
        signature_cells.append(right)
        for cell in signature_cells:
            for paragraph in cell.paragraphs: paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for attachment in attachments:
            document.add_page_break()
            attachment_heading = document.add_paragraph(f"LAMPIRAN {document_heading}")
            attachment_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            attachment_heading.runs[0].bold = True
            document.add_paragraph(f"NOMOR : {data['report_number']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(f"TANGGAL : {data['report_date']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
            if attachment_layout(attachment) == "side" and len(attachment["images"]) == 2:
                photo_table = document.add_table(rows=1, cols=2)
                for cell, image_data in zip(photo_table.rows[0].cells, attachment["images"]):
                    ratio = image_data["width"] / image_data["height"]
                    scale = image_data["scale"] / 100
                    width = min(7.2, 18 * ratio) * scale
                    height = width / ratio
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(BytesIO(image_data["content"]), width=Cm(width), height=Cm(height))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                for image_data in attachment["images"]:
                    ratio = image_data["width"] / image_data["height"]
                    scale = image_data["scale"] / 100
                    max_width, max_height = 15.5, (9 if len(attachment["images"]) == 2 else 17)
                    width = min(max_width, max_height * ratio) * scale
                    height = width / ratio
                    document.add_picture(BytesIO(image_data["content"]), width=Cm(width), height=Cm(height))
                    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name=f"{export_name}.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if file_type == "pdf":
        output = BytesIO()
        styles = getSampleStyleSheet()
        body = ParagraphStyle("ReportBody", parent=styles["BodyText"], fontName=PDF_FONT_NAME,
                              fontSize=12, leading=18, alignment=TA_JUSTIFY, spaceBefore=0, spaceAfter=0)
        heading = ParagraphStyle("ReportHeading", parent=body, fontName=PDF_FONT_BOLD,
                                 spaceBefore=0, spaceAfter=0, keepWithNext=True)
        centered = ParagraphStyle("ReportTitle", parent=heading, fontSize=12, leading=18, alignment=TA_CENTER)
        folio = (21.59*cm, 33.02*cm)
        pdf = SimpleDocTemplate(output, pagesize=folio, rightMargin=1.5*cm, leftMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm, title=f"LAPINHAR {data['report_number']}")
        months = ("", "Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli",
                  "Agustus", "September", "Oktober", "November", "Desember")
        try:
            report_day = date.fromisoformat(data["report_date"])
            display_date = f"{report_day.day} {months[report_day.month]} {report_day.year}"
        except (TypeError, ValueError):
            display_date = data["report_date"] or "—"

        def draw_cover(canvas, document_template):
            page_width, page_height = folio
            canvas.saveState()
            canvas.setFont(PDF_FONT_BOLD, 12)
            canvas.drawCentredString(page_width / 2, page_height - 2.2 * cm, data["organization"])
            canvas.drawImage(str(STATIC_IMG_DIR / "logo-kejaksaan.jpeg"),
                             (page_width - 4.2 * cm) / 2, page_height - 7.85 * cm,
                             width=4.2 * cm, height=4.1 * cm, preserveAspectRatio=True, mask="auto")
            canvas.drawCentredString(page_width / 2, page_height - 8.9 * cm, "LAPINHAR")
            canvas.setFont(PDF_FONT_BOLD, 12)
            words, lines, current = data["subject"].split(), [], ""
            for word in words:
                candidate = f"{current} {word}".strip()
                if canvas.stringWidth(candidate, PDF_FONT_BOLD, 12) <= 17.5 * cm:
                    current = candidate
                else:
                    lines.append(current)
                    current = word
            if current:
                lines.append(current)
            for index, line in enumerate(lines):
                canvas.drawCentredString(page_width / 2, page_height - (9.55 * cm + index * .55 * cm), line)
            canvas.saveState()
            canvas.translate(page_width * .43, 11 * cm)
            canvas.rotate(56)
            canvas.drawImage(str(STATIC_IMG_DIR / "rahasia.jpeg"),
                             -8.65 * cm, -1.375 * cm, width=17.3 * cm, height=2.75 * cm,
                             preserveAspectRatio=True, mask="auto")
            canvas.restoreState()
            canvas.drawCentredString(page_width / 2, 1.35 * cm, f"{data['city']}, {display_date}")
            canvas.restoreState()

        story = [PageBreak(), Paragraph(escape(data["organization"]), centered), Spacer(1, 3)]
        office_rule = Table([[""]], colWidths=[17.5 * cm], rowHeights=[3])
        office_rule.setStyle(TableStyle([("LINEABOVE", (0, 0), (-1, -1), 1.2, "black"),
                                         ("LINEBELOW", (0, 0), (-1, -1), .6, "black")]))
        story.extend([office_rule, Spacer(1, 42), Paragraph("<u>NOTA  -  DINAS</u>", centered), Spacer(1, 25)])
        for label, value in (("K E P A D A", data["recipient"]), ("D A R I", data["sender"]),
                             ("NOMOR", data["report_number"]), ("TANGGAL", display_date),
                             ("SIFAT", data["classification"]), ("LAMPIRAN", data["attachment"]),
                             ("PERIHAL", data["subject"])):
            story.append(Table([[label, ":", Paragraph(escape(value), body)]], colWidths=[3.3*cm,.5*cm,14*cm]))
        memo_signer = Paragraph(
            f"{escape(data['auth_position'])}<br/>{escape(data['organization'])}<br/><br/><br/><br/><br/>"
            f"<u>{escape(data['auth_name'])}</u><br/>{escape(data['auth_rank_nip'])}", centered)
        story.extend([Spacer(1, 70), Table([["", memo_signer]], colWidths=[8.5*cm, 8.5*cm]), PageBreak(),
                      Table([[Paragraph(f"<b>{escape(data['organization'])}</b>", body),
                              Paragraph("<b>L. IN.1</b><br/>Copy ke 1<br/>Dari 1 copies", body)]],
                            colWidths=[13.2*cm, 3.8*cm]), Spacer(1, 18),
                      Paragraph("<u>LAPORAN INFORMASI HARIAN</u>", centered),
                      Paragraph(f"NOMOR : {escape(data['report_number'])}", centered), Spacer(1, 12)])
        for title, content, spacing in (("I.    INFORMASI YANG DIPEROLEH", data["information"], data["information_spacing"]),
                               ("II.   SUMBER INFORMASI", data["sources"], data["sources_spacing"]),
                               ("III.  TREN PERKEMBANGAN / PERKIRAAN", data["trends"], data["trends_spacing"]),
                               ("IV.   PENDAPAT / SARAN", data["suggestions"], data["suggestions_spacing"])):
            story.append(Paragraph(title, heading))
            for block in rich_text_blocks(content):
                prefix = escape(block["prefix"])
                formatted = "".join(
                    ("<b>" if run["bold"] else "") + ("<i>" if run["italic"] else "") +
                    ("<u>" if run["underline"] else "") + escape(run["text"]) +
                    ("</u>" if run["underline"] else "") + ("</i>" if run["italic"] else "") +
                    ("</b>" if run["bold"] else "") for run in block["runs"]
                )
                pdf_alignments = {"center": TA_CENTER, "justify": TA_JUSTIFY, "left": TA_LEFT, "right": TA_RIGHT}
                block_style = ParagraphStyle(
                    f"Section{len(story)}", parent=body, leading=18,
                    alignment=pdf_alignments.get(block["align"], TA_JUSTIFY),
                    leftIndent=(block["indent"] * 18) + (18 if prefix else 0),
                    firstLineIndent=-18 if prefix else 0,
                )
                story.append(Paragraph(f"{prefix}{formatted}", block_style))
        auth_signature = Paragraph(f"AUTENTIKASI:<br/>{escape(data['auth_position'])}<br/>{escape(data['organization'])}<br/><br/><br/><br/><br/><u>{escape(data['auth_name'])}</u><br/>{escape(data['auth_rank_nip'])}", centered)
        creator_signature = Paragraph(f"{escape(data['city'])}, {escape(display_date)}<br/>Yang Membuat Laporan<br/>{escape(data['creator_position'])}<br/><br/><br/><br/><br/><u>{escape(data['creator_name'])}</u><br/>{escape(data['creator_rank_nip'])}", centered)
        if data["show_authentication"]:
            signature_table = Table([[auth_signature, creator_signature]], colWidths=[8.5*cm, 8.5*cm])
        else:
            signature_table = Table([["", creator_signature]], colWidths=[8.5*cm, 8.5*cm])
        story.extend([Spacer(1, 20), signature_table])
        for attachment in attachments:
            story.extend([PageBreak(), Paragraph("<b>LAMPIRAN LAPORAN INFORMASI HARIAN</b>", centered),
                          Paragraph(f"NOMOR : {escape(data['report_number'])}", centered),
                          Paragraph(f"TANGGAL : {escape(display_date)}", centered), Spacer(1, 14)])
            if attachment_layout(attachment) == "side" and len(attachment["images"]) == 2:
                photo_cells = []
                for image_data in attachment["images"]:
                    ratio = image_data["width"] / image_data["height"]
                    scale = image_data["scale"] / 100
                    width = min(7.5 * cm, 19 * cm * ratio) * scale
                    photo_cells.append(Image(BytesIO(image_data["content"]), width=width, height=width / ratio))
                story.append(Table([photo_cells], colWidths=[8.2 * cm, 8.2 * cm]))
            else:
                for image_data in attachment["images"]:
                    ratio = image_data["width"] / image_data["height"]
                    scale = image_data["scale"] / 100
                    max_width, max_height = 16 * cm, (9 * cm if len(attachment["images"]) == 2 else 18 * cm)
                    width = min(max_width, max_height * ratio) * scale
                    picture_flowable = Image(BytesIO(image_data["content"]), width=width, height=width / ratio)
                    picture_flowable.hAlign = "CENTER"
                    story.extend([picture_flowable, Spacer(1, 8)])
        class ClassifiedCanvas(reportlab_canvas.Canvas):
            def showPage(self):
                if self.getPageNumber() >= 3:
                    self.saveState()
                    self.setFont(PDF_FONT_NAME, 10)
                    self.drawCentredString(folio[0] / 2, folio[1] - 1.1 * cm, "RAHASIA")
                    self.drawCentredString(folio[0] / 2, .75 * cm, "RAHASIA")
                    self.restoreState()
                super().showPage()

        pdf.build(story, onFirstPage=draw_cover, canvasmaker=ClassifiedCanvas)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name=f"{export_name}.pdf", mimetype="application/pdf")
    return "Format tidak didukung", 400


@app.post("/lapinhar/export-preview-pdf")
@login_required
def export_lapinhar_preview_pdf():
    payload = request.get_json(silent=True) or {}
    preview_html = payload.get("html", "")
    report_number = str(payload.get("report_number", "lapinhar"))
    report_date = str(payload.get("report_date", ""))
    subject = str(payload.get("subject", ""))
    document_kind = str(payload.get("document_kind", "lapinhar"))
    try:
        report_id = int(payload["report_id"]) if payload.get("report_id") else None
    except (TypeError, ValueError):
        report_id = None
    accessible_report = accessible_lapinsus(report_id) if document_kind == "lapinsus" else accessible_lapinhar(report_id)
    if report_id and accessible_report is None:
        abort(404)
    availability_check = lapinsus_number_available if document_kind == "lapinsus" else lapinhar_number_available
    number_available, number_message = availability_check(report_number, str(payload.get("reservation_token", "")), report_id)
    if not number_available:
        return number_message, 409
    if not preview_html or len(preview_html) > 18_000_000:
        return "Pratinjau PDF tidak valid atau terlalu besar.", 400

    soup = BeautifulSoup(preview_html, "html.parser")
    for unsafe in soup.find_all(["script", "iframe", "object", "embed"]):
        unsafe.decompose()
    for element in soup.find_all(True):
        for attribute in list(element.attrs):
            if attribute.lower().startswith("on"):
                del element.attrs[attribute]

    stylesheet = (Path(__file__).parent / "static" / "css" / "style.css").read_text(encoding="utf-8")
    printable_html = f"""<!doctype html><html><head><meta charset="utf-8">
    <style>{stylesheet}</style></head><body class="printing-preview">
    <div class="app-shell"><div class="app-main"><main class="page-content">
    <form class="report-editor"><aside class="preview-column"><div class="preview-pages">
    {str(soup)}
    </div></aside></form></main></div></div></body></html>"""

    if not CHROME_EXECUTABLE.exists():
        return "Google Chrome tidak ditemukan. Atur CHROME_EXECUTABLE pada file .env.", 500
    try:
        with sync_playwright() as playwright:
            browser = playwright.chromium.launch(
                executable_path=str(CHROME_EXECUTABLE),
                headless=True,
                args=browser_launch_args(),
            )
            page = browser.new_page(viewport={"width": 816, "height": 1248})
            page.set_content(printable_html, wait_until="networkidle")
            page.emulate_media(media="print")
            page.evaluate(
                """() => {
                const group = document.querySelector("[data-report-preview-group]");
                if (!group) return;
                const pages = () => Array.from(group.querySelectorAll("[data-report-page]"));
                const lockPage = (page) => {
                  page.style.width = "8.5in";
                  page.style.height = "13in";
                  page.style.minHeight = "13in";
                  page.style.maxHeight = "13in";
                  page.style.boxSizing = "border-box";
                  page.style.overflow = "hidden";
                  page.style.display = "flex";
                  page.style.flexDirection = "column";
                };
                pages().forEach(lockPage);
                const pageSafeBottom = (page, extraReserve = 0) => {
                  const pageBox = page.getBoundingClientRect();
                  const pageStyle = window.getComputedStyle(page);
                  const paddingBottom = Number.parseFloat(pageStyle.paddingBottom) || 0;
                  const footerReserve = Math.max(paddingBottom, 96);
                  return pageBox.bottom - footerReserve - extraReserve;
                };
                const deepestFlowBottom = (page) => {
                  const flowRoots = Array.from(page.children).filter((child) => {
                    const style = window.getComputedStyle(child);
                    return style.position !== "absolute" && !child.hidden;
                  });
                  let bottom = 0;
                  flowRoots.forEach((root) => {
                    [root, ...Array.from(root.querySelectorAll("*"))].forEach((node) => {
                      const style = window.getComputedStyle(node);
                      if (style.display === "none" || style.visibility === "hidden") return;
                      const rect = node.getBoundingClientRect();
                      if (rect.height > 0) bottom = Math.max(bottom, rect.bottom);
                    });
                  });
                  return bottom;
                };
                const overflows = (page, extraReserve = 0) =>
                  page.scrollHeight > page.clientHeight + 2 ||
                  deepestFlowBottom(page) > pageSafeBottom(page, extraReserve);
                const skeleton = pages()[0].cloneNode(true);
                skeleton.querySelector(".official-body").innerHTML = "";
                skeleton.querySelector(".signature-grid")?.remove();
                skeleton.querySelector(".report-code")?.remove();
                skeleton.querySelector(".document-header")?.remove();
                skeleton.classList.add("continuation-page");
                const ensureNext = (page) => {
                  let next = page.nextElementSibling?.matches?.("[data-report-page]")
                    ? page.nextElementSibling
                    : null;
                  if (!next) {
                    next = skeleton.cloneNode(true);
                    group.insertBefore(next, page.nextSibling);
                    lockPage(next);
                  }
                  return next;
                };
                const prependBlock = (body, title, block) => {
                  let section = body.querySelector(".paginated-section");
                  if (!section || section.querySelector("h3")?.textContent !== title) {
                    section = document.createElement("section");
                    section.className = "paginated-section";
                    const heading = document.createElement("h3");
                    heading.textContent = title;
                    section.appendChild(heading);
                    body.insertBefore(section, body.firstChild);
                  }
                  const heading = section.querySelector("h3");
                  section.insertBefore(block, heading ? heading.nextSibling : section.firstChild);
                };
                const countContentBlocks = (page) =>
                  page.querySelectorAll(".official-body .paginated-section > p, .official-body .paginated-section > ol, .official-body .paginated-section > ul").length;
                const moveLastMovable = (page) => {
                  const signature = page.querySelector(":scope > .signature-grid");
                  if (signature) {
                    ensureNext(page).appendChild(signature);
                    return true;
                  }
                  if (countContentBlocks(page) <= 1) return false;
                  const sections = Array.from(page.querySelectorAll(".official-body .paginated-section"));
                  const section = sections.reverse().find((candidate) =>
                    candidate.querySelector(":scope > p, :scope > ol, :scope > ul")
                  );
                  if (!section) return false;
                  const blocks = Array.from(section.querySelectorAll(":scope > p, :scope > ol, :scope > ul"));
                  const block = blocks[blocks.length - 1];
                  const title = section.querySelector("h3")?.textContent || "";
                  block.remove();
                  if (!section.querySelector(":scope > p, :scope > ol, :scope > ul")) section.remove();
                  prependBlock(ensureNext(page).querySelector(".official-body"), title, block);
                  return true;
                };
                let guard = 0;
                while (guard++ < 240) {
                  let changed = false;
                  for (const page of pages()) {
                    if (overflows(page, page.querySelector(":scope > .signature-grid") ? 18 : 0)) {
                      changed = moveLastMovable(page) || changed;
                    }
                  }
                  if (!changed) break;
                }
                const renderedSectionTitles = new Set();
                pages().forEach((page) => {
                  page.querySelectorAll(".official-body .paginated-section > h3").forEach((heading) => {
                    const title = heading.textContent.trim().replace(/\\s+/g, " ").toUpperCase();
                    if (renderedSectionTitles.has(title)) heading.remove();
                    else renderedSectionTitles.add(title);
                  });
                });
                pages().forEach((page, index) => { page.dataset.pageNumber = String(index + 1); });
              }"""
            )
            pdf_bytes = page.pdf(
                width="8.5in", height="13in", print_background=True,
                prefer_css_page_size=True,
                margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
            )
            browser.close()
    except Exception as exc:
        app.logger.exception("Gagal merender PDF pratinjau")
        return f"Gagal membuat PDF: {exc}", 500

    export_name = lapinhar_export_basename(report_number, report_date, subject)
    return send_file(BytesIO(pdf_bytes), as_attachment=True,
                     download_name=f"{export_name}.pdf", mimetype="application/pdf")


@app.route("/users", methods=["GET", "POST"])
@admin_required
def users():
    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        full_name = request.form.get("full_name", "").strip()
        password = request.form.get("password", "")
        role = request.form.get("role", "user")
        if not username or not full_name or len(password) < 6 or role not in {"admin", "user", "mahasiswa"}:
            flash("Lengkapi data. Kata sandi minimal 6 karakter.", "error")
        elif fetch_one("SELECT id FROM users WHERE username = %s", (username,)):
            flash("Username sudah digunakan.", "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    "INSERT INTO users (username, full_name, password_hash, role) VALUES (%s, %s, %s, %s)",
                    (username, full_name, generate_password_hash(password), role),
                )
            flash("Pengguna baru berhasil dibuat.", "success")
            return redirect(url_for("users"))
    return render_template("users.html", users=fetch_all("SELECT id, username, full_name, role, is_active, profile_photo, created_at FROM users ORDER BY created_at DESC"), active="users")


@app.get("/users/<int:user_id>/avatar")
@login_required
def user_avatar(user_id):
    user = fetch_one("SELECT profile_photo FROM users WHERE id=%s", (user_id,))
    if not user or not user.get("profile_photo"):
        abort(404)
    return send_from_directory(USER_AVATAR_UPLOAD_DIR, user["profile_photo"], max_age=0)


@app.route("/profile", methods=["GET", "POST"])
@login_required
def profile_settings():
    user = fetch_one("SELECT id,username,full_name,profile_photo FROM users WHERE id=%s", (session["user_id"],))
    if user is None:
        abort(404)
    if request.method == "POST":
        full_name = request.form.get("full_name", "").strip()
        current_password = request.form.get("current_password", "")
        new_password = request.form.get("new_password", "")
        confirm_password = request.form.get("confirm_password", "")
        try:
            profile_photo = save_user_profile_photo(request.files.get("profile_photo"), user.get("profile_photo"))
        except ValueError as exc:
            flash(str(exc), "error")
            return redirect(url_for("profile_settings"))
        if not full_name:
            flash("Nama pengguna wajib diisi.", "error")
        elif new_password and len(new_password) < 6:
            flash("Password baru minimal 6 karakter.", "error")
        elif new_password and new_password != confirm_password:
            flash("Konfirmasi password tidak sama.", "error")
        elif new_password:
            current = fetch_one("SELECT password_hash FROM users WHERE id=%s", (session["user_id"],))
            if not check_password_hash(current["password_hash"], current_password):
                flash("Password lama tidak sesuai.", "error")
                return redirect(url_for("profile_settings"))
            with get_db().cursor() as cursor:
                cursor.execute(
                    "UPDATE users SET full_name=%s, profile_photo=%s, password_hash=%s WHERE id=%s",
                    (full_name, profile_photo, generate_password_hash(new_password), session["user_id"]),
                )
            session["full_name"] = full_name
            session["profile_photo"] = profile_photo
            flash("Profil dan password berhasil diperbarui.", "success")
            return redirect(url_for("profile_settings"))
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    "UPDATE users SET full_name=%s, profile_photo=%s WHERE id=%s",
                    (full_name, profile_photo, session["user_id"]),
                )
            session["full_name"] = full_name
            session["profile_photo"] = profile_photo
            flash("Profil berhasil diperbarui.", "success")
            return redirect(url_for("profile_settings"))
    return render_template("profile_settings.html", active="profile", user=user)


@app.post("/users/<int:user_id>/update-profile")
@admin_required
def update_user_profile_admin(user_id):
    user = fetch_one("SELECT id,full_name,profile_photo FROM users WHERE id=%s", (user_id,))
    if user is None:
        abort(404)
    full_name = request.form.get("full_name", "").strip()
    new_password = request.form.get("new_password", "")
    confirm_password = request.form.get("confirm_password", "")
    try:
        profile_photo = save_user_profile_photo(request.files.get("profile_photo"), user.get("profile_photo"))
    except ValueError as exc:
        flash(str(exc), "error")
        return redirect(url_for("users"))
    if not full_name:
        flash("Nama pengguna wajib diisi.", "error")
    elif new_password and len(new_password) < 6:
        flash("Password baru minimal 6 karakter.", "error")
    elif new_password and new_password != confirm_password:
        flash("Konfirmasi password tidak sama.", "error")
    else:
        with get_db().cursor() as cursor:
            if new_password:
                cursor.execute(
                    "UPDATE users SET full_name=%s, profile_photo=%s, password_hash=%s WHERE id=%s",
                    (full_name, profile_photo, generate_password_hash(new_password), user_id),
                )
            else:
                cursor.execute(
                    "UPDATE users SET full_name=%s, profile_photo=%s WHERE id=%s",
                    (full_name, profile_photo, user_id),
                )
        if user_id == session["user_id"]:
            session["full_name"] = full_name
            session["profile_photo"] = profile_photo
        flash("Profil pengguna berhasil diperbarui.", "success")
    return redirect(url_for("users"))


@app.post("/users/<int:user_id>/toggle-active")
@admin_required
def toggle_user_active(user_id):
    user = fetch_one("SELECT id,full_name,is_active FROM users WHERE id=%s", (user_id,))
    if user is None:
        abort(404)
    if user_id == session["user_id"]:
        flash("Anda tidak dapat menonaktifkan akun sendiri.", "error")
    else:
        new_status = 0 if user["is_active"] else 1
        with get_db().cursor() as cursor:
            cursor.execute("UPDATE users SET is_active=%s WHERE id=%s", (new_status, user_id))
        flash(f"Akun {user['full_name']} berhasil {'diaktifkan' if new_status else 'dinonaktifkan'}.", "success")
    return redirect(url_for("users"))


@app.post("/users/<int:user_id>/change-password")
@admin_required
def change_user_password(user_id):
    user = fetch_one("SELECT id,full_name FROM users WHERE id=%s", (user_id,))
    if user is None:
        abort(404)
    password = request.form.get("new_password", "")
    confirmation = request.form.get("confirm_password", "")
    if len(password) < 6:
        flash("Password baru minimal 6 karakter.", "error")
    elif password != confirmation:
        flash("Konfirmasi password tidak sama.", "error")
    else:
        with get_db().cursor() as cursor:
            cursor.execute("UPDATE users SET password_hash=%s WHERE id=%s",
                           (generate_password_hash(password), user_id))
        flash(f"Password {user['full_name']} berhasil diganti.", "success")
    return redirect(url_for("users"))


@app.route("/settings/integrations", methods=["GET", "POST"])
@login_required
def integration_settings():
    user_id = session["user_id"]
    inteliz = fetch_one("SELECT inteliz_username,updated_at,connected_at FROM inteliz_user_settings WHERE user_id=%s", (user_id,))
    whatsapp = fetch_one("SELECT contact_name,phone_number,updated_at FROM whatsapp_user_settings WHERE user_id=%s", (user_id,))
    sipede = fetch_one("SELECT sipede_username,updated_at,connected_at FROM sipede_user_settings WHERE user_id=%s", (user_id,))
    if request.method == "POST":
        integration = request.form.get("integration", "")
        if integration == "inteliz":
            username = request.form.get("inteliz_username", "").strip()
            password = request.form.get("inteliz_password", "")
            existing = fetch_one("SELECT inteliz_password_encrypted FROM inteliz_user_settings WHERE user_id=%s", (user_id,))
            if not username or (not password and not existing):
                flash("Username dan password Inteliz wajib diisi pada konfigurasi pertama.", "error")
            else:
                encrypted = inteliz_credential_cipher().encrypt(password.encode()).decode() if password else existing["inteliz_password_encrypted"]
                with get_db().cursor() as cursor:
                    cursor.execute("""INSERT INTO inteliz_user_settings (user_id,inteliz_username,inteliz_password_encrypted)
                    VALUES (%s,%s,%s) ON DUPLICATE KEY UPDATE inteliz_username=VALUES(inteliz_username),
                    inteliz_password_encrypted=VALUES(inteliz_password_encrypted),updated_at=CURRENT_TIMESTAMP""",
                                   (user_id, username, encrypted))
                flash("Konfigurasi Inteliz berhasil disimpan.", "success")
                return redirect(url_for("integration_settings", _anchor="inteliz"))
        elif integration == "whatsapp":
            name = request.form.get("contact_name", "").strip()
            phone = normalize_whatsapp_number(request.form.get("phone_number", ""))
            if not name or not re.fullmatch(r"\d{9,15}", phone):
                flash("Nama tujuan atau nomor WhatsApp tidak valid.", "error")
            else:
                with get_db().cursor() as cursor:
                    cursor.execute("""INSERT INTO whatsapp_user_settings (user_id,contact_name,phone_number)
                    VALUES (%s,%s,%s) ON DUPLICATE KEY UPDATE contact_name=VALUES(contact_name),
                    phone_number=VALUES(phone_number),updated_at=CURRENT_TIMESTAMP""", (user_id, name, phone))
                flash("Konfigurasi WhatsApp berhasil disimpan.", "success")
                return redirect(url_for("integration_settings", _anchor="whatsapp"))
        elif integration == "sipede":
            username = request.form.get("sipede_username", "").strip()
            password = request.form.get("sipede_password", "")
            existing = fetch_one("SELECT sipede_password_encrypted FROM sipede_user_settings WHERE user_id=%s", (user_id,))
            if not username or (not password and not existing):
                flash("Username dan password Sipede wajib diisi pada konfigurasi pertama.", "error")
            else:
                encrypted = inteliz_credential_cipher().encrypt(password.encode()).decode() if password else existing["sipede_password_encrypted"]
                with get_db().cursor() as cursor:
                    cursor.execute("""INSERT INTO sipede_user_settings (user_id,sipede_username,sipede_password_encrypted)
                    VALUES (%s,%s,%s) ON DUPLICATE KEY UPDATE sipede_username=VALUES(sipede_username),
                    sipede_password_encrypted=VALUES(sipede_password_encrypted),
                    session_data_encrypted=NULL,connected_at=NULL,updated_at=CURRENT_TIMESTAMP""",
                                    (user_id, username, encrypted))
                flash("Konfigurasi Sipede berhasil disimpan terenkripsi.", "success")
                return redirect(url_for("integration_settings", _anchor="sipede"))
        else:
            flash("Jenis integrasi tidak valid.", "error")
    return render_template("integration_settings.html", active="integration_settings",
                           inteliz=inteliz, whatsapp=whatsapp, sipede=sipede)


@app.route("/settings/inteliz", methods=["GET", "POST"])
@login_required
def inteliz_settings():
    setting = fetch_one(
        "SELECT inteliz_username, updated_at, connected_at FROM inteliz_user_settings WHERE user_id=%s",
        (session["user_id"],),
    )
    if request.method == "POST":
        inteliz_username = request.form.get("inteliz_username", "").strip()
        inteliz_password = request.form.get("inteliz_password", "")
        existing = fetch_one(
            "SELECT inteliz_password_encrypted FROM inteliz_user_settings WHERE user_id=%s",
            (session["user_id"],),
        )
        if not inteliz_username:
            flash("Username Inteliz wajib diisi.", "error")
        elif not inteliz_password and not existing:
            flash("Password Inteliz wajib diisi saat konfigurasi pertama.", "error")
        else:
            encrypted_password = (
                inteliz_credential_cipher().encrypt(inteliz_password.encode("utf-8")).decode("ascii")
                if inteliz_password else existing["inteliz_password_encrypted"]
            )
            with get_db().cursor() as cursor:
                cursor.execute(
                    """INSERT INTO inteliz_user_settings
                       (user_id, inteliz_username, inteliz_password_encrypted)
                       VALUES (%s, %s, %s)
                       ON DUPLICATE KEY UPDATE inteliz_username=VALUES(inteliz_username),
                       inteliz_password_encrypted=VALUES(inteliz_password_encrypted),
                       updated_at=CURRENT_TIMESTAMP""",
                    (session["user_id"], inteliz_username, encrypted_password),
                )
            flash("Konfigurasi Inteliz berhasil disimpan dengan aman.", "success")
            return redirect(url_for("inteliz_settings"))
    next_url = request.args.get("next", "")
    if not next_url.startswith("/") or next_url.startswith("//"):
        next_url = ""
    return render_template("inteliz_settings.html", setting=setting, active="inteliz_settings",
                           auto_connect=request.args.get("connect") == "1", next_url=next_url)


@app.route("/settings/whatsapp", methods=["GET", "POST"])
@login_required
def whatsapp_settings():
    setting = fetch_one(
        "SELECT contact_name, phone_number, updated_at FROM whatsapp_user_settings WHERE user_id=%s",
        (session["user_id"],),
    )
    next_url = request.values.get("next", "")
    if not next_url.startswith("/") or next_url.startswith("//"):
        next_url = ""
    if request.method == "POST":
        contact_name = request.form.get("contact_name", "").strip()
        phone_number = request.form.get("phone_number", "").strip()
        normalized = normalize_whatsapp_number(phone_number)
        if not contact_name:
            flash("Nama tujuan WhatsApp wajib diisi.", "error")
        elif not re.fullmatch(r"\d{9,15}", normalized):
            flash("Nomor WhatsApp tidak valid. Contoh: 081234567890.", "error")
        else:
            with get_db().cursor() as cursor:
                cursor.execute(
                    """INSERT INTO whatsapp_user_settings (user_id, contact_name, phone_number)
                       VALUES (%s, %s, %s)
                       ON DUPLICATE KEY UPDATE contact_name=VALUES(contact_name),
                       phone_number=VALUES(phone_number), updated_at=CURRENT_TIMESTAMP""",
                    (session["user_id"], contact_name, normalized),
                )
            flash("Konfigurasi WhatsApp berhasil disimpan.", "success")
            return redirect(next_url or url_for("whatsapp_settings"))
    return render_template("whatsapp_settings.html", setting=setting,
                           active="whatsapp_settings", next_url=next_url)


@app.post("/settings/inteliz/connect/start")
@login_required
def start_inteliz_connection():
    credentials = get_inteliz_credentials(session["user_id"])
    if not credentials:
        return jsonify(message="Simpan username dan password Inteliz terlebih dahulu."), 400
    auth_id = uuid.uuid4().hex
    with INTELIZ_AUTH_LOCK:
        for state in INTELIZ_AUTH_SESSIONS.values():
            if state.get("user_id") == session["user_id"] and state.get("status") not in {"success", "error"}:
                state["cancelled"] = True
        INTELIZ_AUTH_SESSIONS[auth_id] = {
            "user_id": session["user_id"], "status": "starting", "message": "Menyiapkan browser server…",
            "captcha": None, "cancelled": False, "updated_at": time.time(),
        }
    threading.Thread(target=run_inteliz_auth, args=(auth_id, session["user_id"], credentials),
                     daemon=True, name=f"inteliz-auth-{session['user_id']}").start()
    return jsonify(auth_id=auth_id, status="starting"), 202


def owned_inteliz_auth(auth_id):
    with INTELIZ_AUTH_LOCK:
        state = INTELIZ_AUTH_SESSIONS.get(auth_id)
        if not state or state.get("user_id") != session.get("user_id"):
            return None
        return state


@app.get("/settings/inteliz/connect/<auth_id>/status")
@login_required
def inteliz_connection_status(auth_id):
    state = owned_inteliz_auth(auth_id)
    if state is None:
        abort(404)
    return jsonify(status=state["status"], message=state.get("message", ""),
                   captcha=state.get("captcha"))


@app.post("/settings/inteliz/connect/<auth_id>/captcha")
@login_required
def submit_inteliz_captcha(auth_id):
    value = str((request.get_json(silent=True) or {}).get("captcha", "")).strip()
    with INTELIZ_AUTH_LOCK:
        state = INTELIZ_AUTH_SESSIONS.get(auth_id)
        if not state or state.get("user_id") != session["user_id"]:
            abort(404)
        if state.get("status") != "captcha" or not value or len(value) > 20:
            return jsonify(message="Kode CAPTCHA tidak valid."), 400
        state["captcha_input"] = value
        state["status"] = "loading"
        state["message"] = "Memverifikasi CAPTCHA…"
    return jsonify(message="CAPTCHA dikirim.")


@app.post("/settings/inteliz/connect/<auth_id>/otp")
@login_required
def submit_inteliz_otp(auth_id):
    value = str((request.get_json(silent=True) or {}).get("otp", "")).strip()
    with INTELIZ_AUTH_LOCK:
        state = INTELIZ_AUTH_SESSIONS.get(auth_id)
        if not state or state.get("user_id") != session["user_id"]:
            abort(404)
        if state.get("status") != "otp" or not value or len(value) > 20:
            return jsonify(message="Kode autentikator tidak valid."), 400
        state["otp_input"] = value
        state["status"] = "loading"
        state["message"] = "Memverifikasi kode autentikator…"
    return jsonify(message="Kode autentikator dikirim.")


@app.post("/settings/inteliz/connect/<auth_id>/cancel")
@login_required
def cancel_inteliz_connection(auth_id):
    with INTELIZ_AUTH_LOCK:
        state = INTELIZ_AUTH_SESSIONS.get(auth_id)
        if not state or state.get("user_id") != session["user_id"]:
            abort(404)
        state["cancelled"] = True
        state["status"] = "error"
        state["message"] = "Proses login dibatalkan."
    return jsonify(message="Proses login dibatalkan.")


@app.post("/settings/sipede/connect/start")
@login_required
def start_sipede_connection():
    credentials = get_sipede_credentials(session["user_id"])
    if not credentials:
        return jsonify(message="Simpan username dan password SIPede terlebih dahulu."), 400
    auth_id = uuid.uuid4().hex
    with SIPEDE_AUTH_LOCK:
        for state in SIPEDE_AUTH_SESSIONS.values():
            if state.get("user_id") == session["user_id"] and state.get("status") not in {"success", "error"}:
                state["cancelled"] = True
        SIPEDE_AUTH_SESSIONS[auth_id] = {
            "user_id": session["user_id"],
            "status": "starting",
            "message": "Menyiapkan browser server…",
            "captcha": None,
            "cancelled": False,
            "updated_at": time.time(),
        }
    threading.Thread(
        target=run_sipede_auth,
        args=(auth_id, session["user_id"], credentials),
        daemon=True,
        name=f"sipede-auth-{session['user_id']}",
    ).start()
    return jsonify(auth_id=auth_id, status="starting"), 202


def owned_sipede_auth(auth_id):
    with SIPEDE_AUTH_LOCK:
        state = SIPEDE_AUTH_SESSIONS.get(auth_id)
        if not state or state.get("user_id") != session.get("user_id"):
            return None
        return state


@app.get("/settings/sipede/connect/<auth_id>/status")
@login_required
def sipede_connection_status(auth_id):
    state = owned_sipede_auth(auth_id)
    if state is None:
        abort(404)
    return jsonify(status=state["status"], message=state.get("message", ""),
                   captcha=state.get("captcha"))


@app.post("/settings/sipede/connect/<auth_id>/captcha")
@login_required
def submit_sipede_captcha(auth_id):
    value = str((request.get_json(silent=True) or {}).get("captcha", "")).strip()
    with SIPEDE_AUTH_LOCK:
        state = SIPEDE_AUTH_SESSIONS.get(auth_id)
        if not state or state.get("user_id") != session["user_id"]:
            abort(404)
        if state.get("status") != "captcha" or not value or len(value) > 20:
            return jsonify(message="Kode CAPTCHA tidak valid."), 400
        state["captcha_input"] = value
        state["status"] = "loading"
        state["message"] = "Memverifikasi CAPTCHA SIPede…"
    return jsonify(message="CAPTCHA dikirim.")


@app.post("/settings/sipede/connect/<auth_id>/cancel")
@login_required
def cancel_sipede_connection(auth_id):
    with SIPEDE_AUTH_LOCK:
        state = SIPEDE_AUTH_SESSIONS.get(auth_id)
        if not state or state.get("user_id") != session["user_id"]:
            abort(404)
        state["cancelled"] = True
        state["status"] = "error"
        state["message"] = "Proses login SIPede dibatalkan."
    return jsonify(message="Proses login SIPede dibatalkan.")


@app.route("/settings/signatories", methods=["GET", "POST"])
@admin_required
def signatory_settings():
    if request.method == "POST":
        code = request.form.get("position_code", "")
        if code not in {"kajari", "kasi_intel", "kasubsi_1", "kasubsi_2"}:
            flash("Jenis penandatangan tidak valid.", "error")
        else:
            position_name = request.form.get("position_name", "").strip()
            full_name = request.form.get("full_name", "").strip()
            rank_nip = request.form.get("rank_nip", "").strip()
            use_tte = 1 if code == "kajari" and request.form.get("use_tte") == "1" else 0
            current = fetch_one(
                "SELECT signature_image FROM signatories WHERE position_code=%s", (code,)
            )
            signature_image = current["signature_image"] if current else None
            uploaded_signature = request.files.get("signature_image")
            remove_signature = request.form.get("remove_signature") == "1"
            if not position_name or not full_name or not rank_nip:
                flash("Seluruh data penandatangan wajib diisi.", "error")
            elif (uploaded_signature and uploaded_signature.filename
                  and uploaded_signature.mimetype not in ALLOWED_IMAGE_TYPES):
                flash("Scan tanda tangan harus berupa JPG, PNG, atau WebP.", "error")
            else:
                old_signature = signature_image
                if remove_signature:
                    signature_image = None
                if uploaded_signature and uploaded_signature.filename:
                    extension = {
                        "image/jpeg": ".jpg",
                        "image/png": ".png",
                        "image/webp": ".webp",
                    }[uploaded_signature.mimetype]
                    SIGNATORY_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
                    signature_image = f"{code}_{uuid.uuid4().hex[:12]}{extension}"
                    uploaded_signature.save(SIGNATORY_UPLOAD_DIR / signature_image)
                with get_db().cursor() as cursor:
                    cursor.execute(
                        """UPDATE signatories SET position_name = %s, full_name = %s, rank_nip = %s,
                           use_tte = %s, signature_image = %s,
                           updated_at = CURRENT_TIMESTAMP WHERE position_code = %s""",
                        (position_name, full_name, rank_nip, use_tte, signature_image, code),
                    )
                if old_signature and old_signature != signature_image:
                    old_path = (SIGNATORY_UPLOAD_DIR / old_signature).resolve()
                    if old_path.parent == SIGNATORY_UPLOAD_DIR.resolve() and old_path.is_file():
                        try:
                            old_path.unlink()
                        except PermissionError:
                            pass
                flash("Data penandatangan berhasil diperbarui.", "success")
                return redirect(url_for("signatory_settings"))
    rows = fetch_all("SELECT * FROM signatories ORDER BY FIELD(position_code, 'kajari', 'kasi_intel', 'kasubsi_1', 'kasubsi_2')")
    return render_template("signatory_settings.html", signatories=rows, active="signatories")


@app.get("/settings/signatories/<position_code>/signature")
@login_required
def signatory_signature_file(position_code):
    if position_code not in {"kajari", "kasi_intel", "kasubsi_1", "kasubsi_2"}:
        abort(404)
    signer = fetch_one(
        "SELECT signature_image FROM signatories WHERE position_code=%s", (position_code,)
    )
    if not signer or not signer["signature_image"]:
        abort(404)
    return send_from_directory(SIGNATORY_UPLOAD_DIR, signer["signature_image"], max_age=0)


@app.route("/settings/organization", methods=["GET", "POST"])
@admin_required
def organization_settings():
    if request.method == "POST":
        organization_name = request.form.get("organization_name", "").strip()
        institution_code = request.form.get("institution_code", "").strip()
        address = request.form.get("address", "").strip()
        phone = request.form.get("phone", "").strip()
        website = request.form.get("website", "").strip()
        current = fetch_one("SELECT digital_stamp FROM organization_settings WHERE id=1")
        digital_stamp = current["digital_stamp"] if current else None
        uploaded_stamp = request.files.get("digital_stamp")
        remove_stamp = request.form.get("remove_digital_stamp") == "1"
        if not all((organization_name, institution_code, address, phone, website)):
            flash("Seluruh data instansi wajib diisi.", "error")
        elif (uploaded_stamp and uploaded_stamp.filename
              and uploaded_stamp.mimetype not in ALLOWED_IMAGE_TYPES):
            flash("Cap digital harus berupa JPG, PNG, atau WebP.", "error")
        else:
            old_stamp = digital_stamp
            if remove_stamp:
                digital_stamp = None
            if uploaded_stamp and uploaded_stamp.filename:
                extension = {
                    "image/jpeg": ".jpg",
                    "image/png": ".png",
                    "image/webp": ".webp",
                }[uploaded_stamp.mimetype]
                ORGANIZATION_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
                digital_stamp = f"cap_{uuid.uuid4().hex[:12]}{extension}"
                uploaded_stamp.save(ORGANIZATION_UPLOAD_DIR / digital_stamp)
            with get_db().cursor() as cursor:
                cursor.execute(
                    """INSERT INTO organization_settings
                       (id, organization_name, institution_code, address, phone, website, digital_stamp)
                       VALUES (1, %s, %s, %s, %s, %s, %s)
                       ON DUPLICATE KEY UPDATE organization_name = VALUES(organization_name),
                       institution_code = VALUES(institution_code), address = VALUES(address),
                       phone = VALUES(phone), website = VALUES(website),
                       digital_stamp = VALUES(digital_stamp),
                       updated_at = CURRENT_TIMESTAMP""",
                    (organization_name, institution_code, address, phone, website, digital_stamp),
                )
            if old_stamp and old_stamp != digital_stamp:
                old_path = (ORGANIZATION_UPLOAD_DIR / old_stamp).resolve()
                if old_path.parent == ORGANIZATION_UPLOAD_DIR.resolve() and old_path.is_file():
                    try:
                        old_path.unlink()
                    except PermissionError:
                        pass
            flash("Data instansi berhasil diperbarui.", "success")
            return redirect(url_for("organization_settings"))
    organization = fetch_one("SELECT * FROM organization_settings WHERE id = 1")
    return render_template("organization_settings.html", organization=organization, active="organization")


@app.get("/settings/organization/digital-stamp")
@login_required
def organization_digital_stamp_file():
    organization = fetch_one("SELECT digital_stamp FROM organization_settings WHERE id=1")
    if not organization or not organization["digital_stamp"]:
        abort(404)
    return send_from_directory(
        ORGANIZATION_UPLOAD_DIR, organization["digital_stamp"], max_age=0
    )


@app.post("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


initialize_database()

if __name__ == "__main__":
    app.run(
        host=os.getenv("FLASK_HOST", "0.0.0.0"),
        port=int(os.getenv("FLASK_PORT", "5000")),
        debug=os.getenv("FLASK_DEBUG", "0") == "1",
    )
